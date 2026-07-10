from __future__ import annotations

import html
import json
import math
import mimetypes
import re
import shutil
import subprocess
import textwrap
from collections import Counter
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import Image as PdfImage
    from reportlab.platypus import PageBreak, Paragraph, Preformatted, SimpleDocTemplate, Spacer, Table, TableStyle

    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False


REPO = Path(__file__).resolve().parents[1]
OUTPUT = REPO / "docs" / "OMB_Portfolio_Builder_Complete_Guide.docx"
DIAGRAM_DIR = REPO / "docs" / "guide-diagrams"
CODE_REFERENCE_DIR = REPO / "docs" / "code-reference"
CODE_REFERENCE_DOCX_DIR = REPO / "docs" / "code-reference-docx"
CODE_REFERENCE_PDF_DIR = REPO / "docs" / "code-reference-pdf"
CODE_REFERENCE_MASTER_DOCX = REPO / "docs" / "OMB_Portfolio_Builder_Code_Reference.docx"
CODE_REFERENCE_MASTER_PDF = REPO / "docs" / "OMB_Portfolio_Builder_Code_Reference.pdf"
FILE_REFERENCE_DOCX_DIR = REPO / "docs" / "important-code-reference-docx"
FILE_REFERENCE_PDF_DIR = REPO / "docs" / "important-code-reference-pdf"
FILE_REFERENCE_MASTER_DOCX = REPO / "docs" / "OMB_Portfolio_Builder_Important_Code_Reference.docx"
FILE_REFERENCE_MASTER_PDF = REPO / "docs" / "OMB_Portfolio_Builder_Important_Code_Reference.pdf"


FILE_DESCRIPTIONS = {
    ".github/workflows/build-windows-builder.yml": "GitHub Actions workflow that builds the Windows installer and portable application, checks the stable download link, enforces release version bumps, uploads artifacts, and publishes GitHub Releases.",
    ".github/workflows/main-branch-gate.yml": "GitHub Actions guard that rejects direct pushes to main and rejects pull requests into main unless the source branch is development.",
    ".gitignore": "Git ignore rules for build outputs, temporary files, and local-only data that should not be committed.",
    ".nojekyll": "Tells GitHub Pages not to process the site with Jekyll, so folders and static files publish exactly as generated.",
    ".well-known/security.txt": "Public security contact file used by browsers, researchers, and scanners to find responsible disclosure information.",
    "BRANCHING.md": "Human-readable branch model explaining development, main, feature branches, and release flow.",
    "Backgrounds/.gitkeep": "Keeps the Backgrounds folder in Git even when no background image files are present.",
    "README.md": "Main README for the standalone builder application: install, updates, workspaces, text editing, publishing, build commands, branch model, and uninstall.",
    "_headers": "Cloudflare Pages style HTTP headers for security and caching behavior on the public website.",
    "assets/apple-touch-icon.png": "Apple touch icon shown when the public site is saved on iOS or compatible launch surfaces.",
    "assets/favicon-16.png": "Small browser favicon used by tabs and browser UI.",
    "assets/favicon-192.png": "Large web app icon for Android and installable browser surfaces.",
    "assets/favicon-32.png": "Standard browser favicon size.",
    "assets/favicon-48.png": "Windows/browser icon size used by some desktop surfaces.",
    "assets/favicon.ico": "ICO favicon bundle used by browsers that request favicon.ico.",
    "assets/omb-app-icon-256.png": "Builder application icon source at 256 pixels.",
    "assets/omb-app-icon-512.png": "Builder application icon source at 512 pixels.",
    "assets/omb-app-icon.ico": "Windows ICO application icon used by Electron Builder and shortcuts.",
    "assets/omb-mark.png": "OMB snake/mark image used for branding in the site and builder.",
    "build/installer.nsh": "NSIS installer customization script for update behavior, shortcut placement, old install detection, and setup details.",
    "builder-rich-future-sections.js": "Builder-side helper code for richer future portfolio sections and section defaults.",
    "cloudflare/README.md": "Cloudflare setup notes for deploying the AI Worker endpoint and wiring the public site to the backend.",
    "cloudflare/portfolio-ai-worker.js": "Cloudflare Worker backend for the public portfolio AI assistant and fallback intelligence path.",
    "cloudflare/wrangler.toml": "Wrangler configuration used to deploy the Cloudflare Worker.",
    "docs/.gitkeep": "Keeps the docs folder in Git even before generated documents are added.",
    "electronics-search.js": "Electronics vocabulary and search support used by the website search behavior.",
    "index.html": "Public website HTML shell that recruiters and visitors load in the browser.",
    "installer-notes.txt": "Installer license/notes text shown during Windows setup.",
    "main.cjs": "Electron main process: creates the desktop window, starts the local backend, handles menus, update handoff, and application lifecycle.",
    "package.json": "Node/Electron package manifest with scripts, dependencies, version, installer settings, and extra resources.",
    "pnpm-lock.yaml": "Pinned dependency lockfile so installs and GitHub Actions use the same dependency versions.",
    "pnpm-workspace.yaml": "pnpm workspace declaration for the repository.",
    "portfolio-README.md": "README for the generated public portfolio website rather than the builder application.",
    "project-templates.json": "Appearance template definitions used by projects to control visual feel rather than content.",
    "projects.json": "Public project catalog consumed by the website after parsing builder data.",
    "robots.txt": "Crawler instruction file for search engines.",
    "script.js": "Public website JavaScript for navigation, portfolio rendering, search, AI UI, and interactive project windows.",
    "server.mjs": "Local backend server used by the builder for drafts, parsing, publishing, compiler execution, authentication checks, and file operations.",
    "styles.css": "Public website CSS for layout, contact area, projects, responsive behavior, AI panel, and mobile/desktop rendering.",
    "template-preview.css": "Builder application CSS for dialogs, windows, rich editors, compile code, dark mode, guide windows, and local previews.",
    "template-preview.html": "Builder HTML shell loaded inside the Electron app and local preview runtime.",
    "template-preview.js": "Builder frontend brain: state handling, rich editors, project builder, previews, parser triggers, publishing UI, compile workspace, and guide windows.",
}


ARCHITECTURE_PAGES = [
    ("High-Level System Map", "The system has two personalities: a private builder app and a public website. The builder is where content is created and verified. The website is what recruiters see after publishing.", "Owner -> OMB Builder App -> local backend -> drafts/files/parser -> Git target -> public website\nRecruiter -> browser -> GitHub Pages or Cloudflare -> index.html/styles.css/script.js -> projects.json/assets"),
    ("Why Electron Was Chosen", "Electron lets the builder look like a desktop app while still using HTML, CSS, and JavaScript. That means the same preview rendering logic can be reused between the builder and website.", "Windows OS -> Electron main.cjs -> BrowserWindow -> template-preview.html -> template-preview.js"),
    ("Private Builder Versus Public Website", "The builder includes editing controls, local draft storage, authentication checks, compiler tools, and publishing actions. The public website removes those controls and only displays approved parsed content.", "Local builder controls stay private\nParsed portfolio content becomes public only after Save Draft and Apply to site"),
    ("Local Backend Role", "The frontend cannot safely do everything by itself. The local backend handles file writes, repository operations, compiler execution, publishing checks, and generated site output.", "template-preview.js -> fetch('/api/...') -> server.mjs -> filesystem/git/compiler -> JSON response"),
    ("Portfolio Parser Role", "The parser converts builder data into clean website data. It decides what becomes an overview, what becomes a clickable section, which files are downloadable, and which images render inline.", "Draft fields -> normalized sections -> projects.json -> script.js renders recruiter-facing windows"),
    ("GitHub Pages Publishing", "The target repository receives static files. GitHub Pages then serves those files over HTTPS. The builder does not run on the public site.", "Local publish mirror -> git commit -> git push -> GitHub Pages build -> custom domain"),
    ("Cloudflare Role", "Cloudflare can sit in front of the domain and can also host Worker endpoints such as the portfolio AI backend. DNS points the custom domain to the website host.", "Visitor DNS lookup -> Cloudflare DNS -> Pages/GitHub host -> browser"),
    ("AI Backend Role", "The website AI should answer from portfolio context when the question is about Maurice's work and use general knowledge when the question is general. The Worker endpoint keeps API keys off the public page.", "Website chat -> /api/portfolio-ai -> Cloudflare Worker -> model/provider/context -> response"),
    ("Compiler Workspace Role", "Compile Code is a local project tool. It saves source into a compile workspace, runs compiler commands, captures terminal output, and lets successful code be appended into project sections.", "Project -> Compile Code -> source file -> compiler runner -> terminal output -> optional append to project"),
    ("Installer And Update Role", "The installer places the app in the per-user AppData Programs folder. Updates should replace that install instead of creating duplicate copies.", "GitHub Release -> installer download -> update handoff -> close app -> install over existing app -> reopen"),
    ("Branch Model", "Development collects work. Main is the release branch. Feature branches should merge into development first. Main should receive only development.", "feature branch -> development -> main -> release workflow -> installer"),
    ("Security Boundary", "Editing is local and authenticated. Public visitors should not get editing tools. Publishing requires GitHub write access to the target repository.", "Visitor can read public site\nOwner can edit builder\nGit credentials gate publishing"),
]

ARCHITECTURE_DIAGRAM_BY_TITLE = {
    "High-Level System Map": "system-overview",
    "Why Electron Was Chosen": "electron-runtime",
    "Portfolio Parser Role": "builder-to-site-files",
    "GitHub Pages Publishing": "builder-to-site-files",
    "Cloudflare Role": "cloudflare-ai-flow",
    "AI Backend Role": "cloudflare-ai-flow",
    "Compiler Workspace Role": "compile-code-flow",
    "Installer And Update Role": "release-pipeline",
    "Branch Model": "release-pipeline",
}


COMMAND_LESSONS = [
    ("git status", "Shows what branch you are on and whether files are changed, staged, or clean.", "Run it before edits, before commits, and before merges."),
    ("git switch development", "Moves your working copy to the development branch.", "Use development as the integration branch before anything goes to main."),
    ("git switch -c codex/example-change", "Creates a new feature branch from the branch you are currently on.", "Use this before changing the builder so development stays stable."),
    ("git add <file>", "Stages a changed file for the next commit.", "Only stage files that belong to the current change."),
    ("git commit -m \"message\"", "Records staged changes as a named checkpoint.", "Write messages that explain why the change exists."),
    ("git push origin <branch>", "Uploads your branch to GitHub.", "This lets GitHub run workflows and enables pull requests."),
    ("git pull --ff-only", "Updates the current branch only if Git can fast-forward safely.", "Avoids surprise merge commits."),
    ("git fetch --tags --force", "Downloads branch and tag metadata from GitHub.", "Used before checking whether a release tag already exists."),
    ("pnpm install --frozen-lockfile", "Installs Node dependencies exactly as pinned in pnpm-lock.yaml.", "Makes release builds repeatable."),
    ("pnpm run dev", "Starts the Electron builder in development mode.", "Useful while actively changing builder UI code."),
    ("pnpm run pack", "Builds a portable Windows application artifact.", "Good for quick packaging tests."),
    ("pnpm run installer", "Builds the NSIS Windows installer.", "Use this when testing installation and update behavior."),
    ("pnpm run dist", "Builds all configured Windows release artifacts.", "GitHub Actions uses this for the official release."),
    ("electron .", "Starts Electron using package.json main.cjs.", "Runs a desktop shell around the local web builder."),
    ("wrangler deploy", "Deploys the Cloudflare Worker backend.", "Use it when the AI endpoint code changes."),
    ("gcc file.c -o app.exe", "Compiles C source into a Windows executable.", "The builder uses similar logic when C tools are available."),
    ("g++ file.cpp -std=c++20 -o app.exe", "Compiles modern C++ into a Windows executable.", "C++ is compiled separately from C."),
    ("javac Main.java", "Compiles Java source into .class bytecode.", "Java has a compile step before running."),
    ("java Main", "Runs compiled Java bytecode.", "The builder shows this output in the terminal panel."),
    ("python script.py", "Runs Python source directly.", "Python is interpreted by default."),
    ("node script.js", "Runs JavaScript through Node.js.", "Useful for builder scripts and JavaScript examples."),
    ("iverilog -g2012 design.sv -o sim.vvp", "Compiles Verilog/SystemVerilog into an Icarus simulation file.", "Prepares HDL for simulation."),
    ("vvp sim.vvp", "Runs Icarus Verilog simulation output.", "This is how HDL terminal output appears."),
    ("git push --dry-run origin main", "Tests whether GitHub would accept a push without changing the remote.", "Used during authorization checks."),
    ("git remote -v", "Shows which GitHub repository the local folder talks to.", "Important for publishing targets."),
    ("git branch --show-current", "Prints the active branch name.", "The builder must know whether it is pushing the right branch."),
    ("git tag --list builder-v*", "Lists builder release tags.", "Release workflows use this to prevent duplicate releases."),
    ("Get-ChildItem", "PowerShell command that lists files and folders.", "Used to inspect workspaces and installer output."),
    ("Remove-Item", "PowerShell command that deletes files or folders.", "Use with care and verified paths."),
    ("Start-Process", "PowerShell command that starts another program.", "The updater uses detached process launching."),
]


FEATURE_TOPICS = [
    "Profile identity and contact details", "Front-page hero text", "Fun facts block", "Portfolio areas", "Project categories", "Project creation flow", "Project overview", "Design sections", "Simulation sections", "Files and downloadable evidence", "Images and captions", "Links and external URLs", "Resume viewing", "Project preview", "Full portfolio preview", "Save project", "Save all sections", "Save draft", "Apply to site", "Publishing target", "GitHub authorization", "Local draft recovery", "Mobile rendering", "Desktop rendering", "Search engine behavior", "Website search", "AI chat", "Cloudflare AI worker", "Code block insertion", "Compile Code workspace", "HDL testbench requirement", "HDL signal scope", "Terminal output", "Messages log", "Append code to project", "Dark mode", "Preferences menu", "Window controls", "Builder guide windows", "Installer workflow", "Updater workflow", "Uninstaller behavior", "Branch protection", "Release tags", "Security headers", "DNS and custom domain", "Offline editing", "Public/private boundary", "Asset synchronization", "Project parser", "Template appearance library", "Browser back and refresh support", "Error messages", "Troubleshooting logs",
]


TROUBLESHOOTING_TOPICS = [
    ("The app does not open", "Check whether the installed executable exists under AppData Local Programs. If it does, start it directly. If it fails, reinstall from the latest GitHub Release."),
    ("Update does not restart", "The update handoff must close Electron, run the installer, and then launch the new executable. Check the updater log and whether another builder process is still running."),
    ("GitHub authentication succeeds but builder cannot push", "Authentication and repository freshness are different. Pull or load from target first if the remote branch is ahead."),
    ("Dry-run push says non-fast-forward", "The target repository has commits not present locally. Load from target or pull before publishing."),
    ("Website changes show on desktop but not phone", "Check cache, confirm the published files changed, and verify the mobile renderer is using the same projects.json and CSS."),
    ("AI endpoint unavailable", "Verify the Worker route, environment secrets, Cloudflare deployment, and browser network errors."),
    ("Compiler takes too long", "Check whether the first run is installing or detecting tools. Cached runs should be faster after the first successful compile."),
    ("SystemVerilog output missing", "Check that Icarus Verilog and vvp are available and that the file is a simulation testbench or has a runnable top module."),
    ("Right-click formatting does not apply", "Confirm the active editor saved its selection before the menu changed font, size, or color."),
    ("Links are treated as formulas", "URL detection should win before formula detection. Pasted http, https, and www links must become normal links."),
    ("Images duplicate when moved", "Drag movement should move an existing block rather than copy it."),
    ("Installer reports an old installation", "The installer searches per-user and legacy locations. Remove stale install records only after confirming the active app path."),
    ("Workflow fails on main", "Main is protected. The source branch must be development, and package.json version must be new for release builds."),
    ("Release tag already exists", "Bump package.json to a new version before merging development into main."),
    ("Search cannot find a project", "Ensure the project was saved, parsed, and included in projects.json before publishing."),
]


CONCEPTS = [
    "What HTML does", "What CSS does", "What JavaScript does", "What JSON does", "What Markdown does", "What YAML does", "What TOML does", "What an executable is", "What an installer is", "What a portable app is", "What a local server is", "What a port is", "What localhost means", "What HTTPS means", "What DNS means", "What Cloudflare does", "What GitHub Pages does", "What a repository is", "What a commit is", "What a branch is", "What a pull request is", "What a merge conflict is", "What a release tag is", "What a workflow artifact is", "What a cache is", "What AppData is", "What Program Files is", "What UAC is", "What code signing is", "What a favicon is", "What an asset folder is", "What a parser is", "What a renderer is", "What a compiler is", "What an interpreter is", "What a simulator is", "What stdout is", "What stderr is", "What stdin is", "What a binary is", "What a class file is", "What a source file is", "What syntax highlighting is", "What a code beautifier is", "What a Cloudflare Worker is", "What an API endpoint is", "What a secret is", "What CORS means", "What CSP means", "What security headers do", "What robots.txt does", "What a resume link is", "What a downloadable file is", "What responsive design is", "What mobile breakpoint means", "What a modal window is", "What a route is", "What browser history is", "What an auth cache is", "What MFA means", "What IAM means", "What least privilege means", "What public and private data mean", "What local-first means", "What offline editing means", "What deployment means", "What rollback means", "What version number means", "What semantic versioning is", "What a dependency is", "What pnpm is", "What Node.js is", "What Electron Builder is", "What NSIS is", "What a desktop shortcut is", "What an uninstall entry is", "What a dry run is", "What non-fast-forward means", "What origin means", "What upstream means", "What a working tree is", "What staging means", "What binary asset means", "What a template means here", "What appearance means here", "What project evidence means", "What recruiter-facing means", "What a build pipeline is", "What release automation is", "What a branch gate is", "What status checks are", "What continuous integration is", "What continuous deployment is", "What a local draft is", "What a publish mirror is", "What a workspace root is", "What an endpoint fallback is", "What general AI vs portfolio AI means", "What context retrieval means", "What a prompt is", "What rate limits are", "What privacy boundaries are", "What audit logs are", "What download analytics are", "What IP address logging means", "What consent means", "What an environment variable is", "What a shell is", "What PowerShell is", "What a terminal is", "What a path is", "What a relative path is", "What an absolute path is", "What a file extension is", "What a MIME type is", "What a PDF is", "What a DOCX is", "What accessibility means", "What alt text is", "What line wrapping is", "What a hard line break is", "What rich text is", "What a contenteditable editor is", "What clipboard paste means", "What sanitization means", "What XSS means", "What static hosting cannot do", "What a backend can do", "What an API key is", "What Git Credential Manager is", "What an OAuth browser prompt is", "What branch protection can and cannot prove", "What a required status check is", "What a stale branch is", "What a fast-forward is", "What a generated file is", "What not to edit manually", "How to recover from a bad update", "How to recover from a broken draft", "How to verify the public site", "How to verify the mobile site", "How to verify installer output", "How to verify a compiler", "How to verify the AI backend", "How to verify DNS", "How to verify HTTPS", "How to verify search", "How to verify uploaded files", "How to verify resume viewing", "How to verify project parsing", "How to verify appearance templates", "How to verify branch rules", "How to prepare a main release", "How to explain this app in an interview", "How to onboard a new machine", "How to keep local data safe", "How to avoid duplicate installs", "How to think about future features", "How to read errors calmly",
]


TRADEOFFS = [
    ("Static site vs dynamic server", "A static site is cheaper, simpler, fast, and easy to host on GitHub Pages. A dynamic server could personalize more, but it needs uptime, security patching, and database operations."),
    ("Electron vs pure web builder", "Electron gives desktop file access and packaging. A pure web builder would be easier to access anywhere, but browser security would block many local file, compiler, and Git operations."),
    ("GitHub Releases vs custom update server", "GitHub Releases are reliable and transparent. A custom updater could be more controlled, but it would add server cost and security responsibility."),
    ("Per-user install vs Program Files", "Per-user installs update without admin prompts. Program Files feels traditional but causes UAC friction and update failures."),
    ("Local compiler execution vs cloud compilation", "Local compilation keeps source private and works offline. Cloud compilation could install fewer local tools but would require backend cost and security controls."),
    ("Rich text blocks vs one textarea", "Blocks make images, formulas, and code easier to store. A single textarea feels more like Word but is harder to parse into safe website output."),
    ("Git authentication cache vs always prompt", "Caching reduces annoyance. Always prompting is stricter but makes normal publishing painful."),
    ("Cloudflare Worker AI vs browser-only AI", "Worker AI hides secrets and can access server-side providers. Browser-only AI would expose keys or be limited to local browser models."),
    ("One repo vs split repos", "One repo simplifies early development. Split builder and portfolio repos can make release and content ownership cleaner later."),
    ("Template appearance vs content templates", "Appearance templates avoid injecting fake text. Content templates could speed setup, but they risk making portfolios look generic."),
]


DIAGRAM_SPECS = [
    {
        "name": "system-overview",
        "title": "Private Builder To Public Portfolio",
        "nodes": [
            ("owner", "Owner\ncreates content", 60, 160, 230, 115, "#DBEAFE"),
            ("builder", "OMB Builder App\nElectron desktop UI", 365, 140, 270, 135, "#CFFAFE"),
            ("backend", "Local backend\nserver.mjs", 710, 140, 240, 135, "#E0F2FE"),
            ("workspace", "Local workspaces\ndrafts, files, compile-code", 710, 345, 265, 130, "#DCFCE7"),
            ("git", "Git target repo\ncommit + push", 1040, 140, 250, 135, "#FEF3C7"),
            ("site", "Public website\nHTML/CSS/JS/assets", 1040, 345, 250, 130, "#FCE7F3"),
            ("visitor", "Recruiter browser\nread-only visitor", 1320, 345, 230, 130, "#EDE9FE"),
            ("ai", "Cloudflare Worker\nportfolio AI endpoint", 1040, 565, 250, 130, "#CCFBF1"),
        ],
        "edges": [
            ("owner", "builder", "clicks and types"),
            ("builder", "backend", "API requests"),
            ("backend", "workspace", "saves/parses"),
            ("backend", "git", "publishes"),
            ("git", "site", "hosts files"),
            ("site", "visitor", "serves pages"),
            ("site", "ai", "chat API"),
        ],
    },
    {
        "name": "electron-runtime",
        "title": "Electron Runtime Inside The Desktop App",
        "nodes": [
            ("windows", "Windows OS", 70, 150, 220, 110, "#E0E7FF"),
            ("main", "main.cjs\nElectron main process", 365, 135, 275, 140, "#DBEAFE"),
            ("server", "server.mjs\nlocal backend", 710, 135, 250, 140, "#CFFAFE"),
            ("window", "BrowserWindow\nbuilder screen", 365, 365, 275, 140, "#FEF9C3"),
            ("html", "template-preview.html\nCSS + JS", 710, 365, 250, 140, "#FDE68A"),
            ("system", "Filesystem, Git,\ncompilers, updater", 1045, 250, 300, 140, "#DCFCE7"),
        ],
        "edges": [
            ("windows", "main", "starts"),
            ("main", "server", "launches"),
            ("main", "window", "creates"),
            ("window", "html", "loads"),
            ("html", "server", "fetch /api"),
            ("server", "system", "executes"),
        ],
    },
    {
        "name": "builder-to-site-files",
        "title": "How Builder Edits Become Website Files",
        "nodes": [
            ("ui", "template-preview.js\nbuilder state", 70, 165, 260, 130, "#DBEAFE"),
            ("api", "server.mjs\nsave + parse APIs", 405, 165, 260, 130, "#CFFAFE"),
            ("draft", "local draft\nprojects.local.json\nuploads", 740, 120, 270, 130, "#DCFCE7"),
            ("parser", "portfolio parser\nclean website data", 740, 340, 270, 130, "#FEF3C7"),
            ("json", "projects.json\npublic catalog", 1085, 120, 250, 130, "#FCE7F3"),
            ("assets", "assets + files\nimages, PDFs, code", 1085, 340, 250, 130, "#EDE9FE"),
            ("site", "index.html\nscript.js\nstyles.css", 1085, 560, 250, 130, "#E0F2FE"),
        ],
        "edges": [
            ("ui", "api", "save buttons"),
            ("api", "draft", "writes raw edits"),
            ("draft", "parser", "normalizes"),
            ("parser", "json", "outputs catalog"),
            ("parser", "assets", "copies files"),
            ("json", "site", "rendered by"),
            ("assets", "site", "linked by"),
        ],
    },
    {
        "name": "release-pipeline",
        "title": "Builder Release Pipeline",
        "nodes": [
            ("feature", "feature/codex branch", 75, 165, 240, 120, "#DBEAFE"),
            ("dev", "development\nintegration", 385, 165, 240, 120, "#CFFAFE"),
            ("main", "main\nrelease branch", 695, 165, 240, 120, "#FDE68A"),
            ("gate", "Main Branch Gate\nmust come from development", 695, 370, 260, 130, "#FEE2E2"),
            ("workflow", "GitHub Actions\nbuild-windows-builder.yml", 1015, 165, 300, 130, "#DCFCE7"),
            ("release", "GitHub Release\ninstaller + portable", 1015, 370, 300, 130, "#EDE9FE"),
            ("update", "Installed app\nupdate dialog", 1325, 370, 220, 130, "#FCE7F3"),
        ],
        "edges": [
            ("feature", "dev", "merge"),
            ("dev", "main", "release PR"),
            ("main", "gate", "checked by"),
            ("main", "workflow", "push triggers"),
            ("workflow", "release", "publishes"),
            ("release", "update", "detected by"),
        ],
    },
    {
        "name": "compile-code-flow",
        "title": "Compile Code Workspace",
        "nodes": [
            ("source", "Add or import\nsource code", 75, 150, 250, 125, "#DBEAFE"),
            ("save", "Save source\ncompile-code folder", 390, 150, 250, 125, "#DCFCE7"),
            ("lang", "Language profile\nC, C++, Java,\nPython, JS, HDL", 705, 150, 260, 145, "#FEF3C7"),
            ("runner", "Compiler runner\nserver.mjs", 1030, 135, 250, 125, "#CFFAFE"),
            ("scope", "HDL testbench\n+ signal scope", 1030, 315, 250, 125, "#DCFCE7"),
            ("terminal", "Real terminal output\nstdout + stderr", 1030, 495, 250, 125, "#E0F2FE"),
            ("messages", "Messages log\nbuilder feedback", 705, 365, 260, 125, "#FCE7F3"),
            ("append", "Append code\nto project section", 390, 365, 250, 125, "#EDE9FE"),
        ],
        "edges": [
            ("source", "save", "save"),
            ("save", "lang", "detect"),
            ("lang", "runner", "compile/run"),
            ("runner", "scope", "VCD waveform"),
            ("runner", "terminal", "streams result"),
            ("runner", "messages", "status"),
            ("source", "append", "optional"),
        ],
    },
    {
        "name": "tool-build-map",
        "title": "What Tool Builds What",
        "nodes": [
            ("pnpm", "pnpm\ninstalls dependencies", 75, 140, 250, 115, "#DBEAFE"),
            ("electron", "Electron\nruns desktop app", 390, 140, 250, 115, "#CFFAFE"),
            ("builder", "electron-builder\npackages app", 705, 140, 260, 115, "#DCFCE7"),
            ("nsis", "NSIS\nWindows installer", 1030, 140, 250, 115, "#FEF3C7"),
            ("github", "GitHub Actions\ncloud build machine", 390, 365, 250, 125, "#FCE7F3"),
            ("wrangler", "Wrangler\nCloudflare Worker", 705, 365, 260, 125, "#EDE9FE"),
            ("git", "Git\npublishes website files", 1030, 365, 250, 125, "#E0F2FE"),
        ],
        "edges": [
            ("pnpm", "electron", "provides modules"),
            ("electron", "builder", "app source"),
            ("builder", "nsis", "installer artifact"),
            ("github", "pnpm", "runs"),
            ("github", "builder", "runs"),
            ("wrangler", "github", "separate deploy path"),
            ("git", "github", "pushes code"),
        ],
    },
    {
        "name": "file-communication-map",
        "title": "How Key Files Communicate",
        "nodes": [
            ("main", "main.cjs\napp shell", 65, 145, 230, 115, "#DBEAFE"),
            ("server", "server.mjs\nbackend APIs", 365, 145, 230, 115, "#CFFAFE"),
            ("html", "template-preview.html\nbuilder document", 665, 145, 250, 115, "#FEF3C7"),
            ("builderjs", "template-preview.js\nbuilder logic", 965, 145, 250, 115, "#DCFCE7"),
            ("buildercss", "template-preview.css\nbuilder style", 1265, 145, 250, 115, "#E0F2FE"),
            ("index", "index.html\npublic shell", 365, 390, 230, 115, "#FCE7F3"),
            ("script", "script.js\npublic logic", 665, 390, 250, 115, "#EDE9FE"),
            ("styles", "styles.css\npublic style", 965, 390, 250, 115, "#CCFBF1"),
            ("projects", "projects.json\npublic data", 1265, 390, 250, 115, "#FDE68A"),
        ],
        "edges": [
            ("main", "server", "starts"),
            ("main", "html", "loads"),
            ("html", "builderjs", "scripts"),
            ("html", "buildercss", "styles"),
            ("builderjs", "server", "fetch API"),
            ("server", "projects", "generates"),
            ("index", "script", "scripts"),
            ("index", "styles", "styles"),
            ("script", "projects", "fetches"),
        ],
    },
    {
        "name": "cloudflare-ai-flow",
        "title": "Website AI Backend Flow",
        "nodes": [
            ("chat", "Visitor types\nAI question", 75, 165, 240, 120, "#DBEAFE"),
            ("site", "script.js\nchat UI", 385, 165, 240, 120, "#CFFAFE"),
            ("worker", "Cloudflare Worker\nportfolio-ai-worker.js", 695, 165, 285, 130, "#DCFCE7"),
            ("context", "Portfolio context\nprojects, files, links", 1035, 120, 275, 130, "#FEF3C7"),
            ("model", "AI provider\nOpenAI or Worker AI", 1035, 340, 275, 130, "#EDE9FE"),
            ("answer", "Conversational answer\nplus relevant links", 695, 390, 285, 130, "#FCE7F3"),
        ],
        "edges": [
            ("chat", "site", "input"),
            ("site", "worker", "POST /api/portfolio-ai"),
            ("worker", "context", "retrieves"),
            ("worker", "model", "asks"),
            ("model", "answer", "returns"),
            ("answer", "site", "renders"),
        ],
    },
    {
        "name": "frontend-backend-cloudflare",
        "title": "Frontend, Backend, Cloudflare, And Secrets",
        "nodes": [
            ("builderfe", "Builder frontend\ntemplate-preview.js", 75, 150, 260, 125, "#DBEAFE"),
            ("localbe", "Local backend\nserver.mjs", 420, 150, 260, 125, "#CFFAFE"),
            ("parser", "Parser output\nprojects.json + assets", 760, 150, 275, 125, "#DCFCE7"),
            ("github", "GitHub Pages\nstatic hosting", 1110, 150, 260, 125, "#FEF3C7"),
            ("publicfe", "Public frontend\nscript.js + index.html", 1110, 390, 260, 125, "#EDE9FE"),
            ("worker", "Cloudflare Worker\n/api/portfolio-ai", 760, 390, 260, 125, "#CCFBF1"),
            ("secrets", "Cloudflare secrets\nOPENAI_API_KEY", 420, 390, 275, 125, "#FEE2E2"),
            ("model", "AI provider\nOpenAI or Worker AI", 75, 390, 260, 125, "#FCE7F3"),
        ],
        "edges": [
            ("builderfe", "localbe", "fetch /api"),
            ("localbe", "parser", "builds"),
            ("parser", "github", "publishes"),
            ("github", "publicfe", "serves"),
            ("publicfe", "worker", "chat request"),
            ("worker", "secrets", "reads server-side"),
            ("worker", "model", "calls"),
            ("model", "worker", "answer"),
        ],
    },
]


SOFTWARE_DECISIONS = [
    ("Separate builder and website", "The builder needs editing controls, private local files, authentication, compiler execution, and draft recovery. The website needs to be public, simple, fast, and read-only. Separating them protects visitors from builder-only controls and protects the owner from accidentally exposing private editing surfaces."),
    ("Use a local backend instead of only browser JavaScript", "Browser JavaScript is intentionally restricted. It cannot safely run compilers, write arbitrary files, manage Git repositories, or run installers. The local backend gives the builder controlled access to the filesystem, Git, compiler tools, and parser output."),
    ("Use a parser as the publishing boundary", "Raw builder state is convenient for editing but not ideal for public display. The parser cleans it, chooses visible sections, preserves rich text safely, copies allowed assets, and produces predictable website data."),
    ("Use AppData per-user install paths", "Per-user installation avoids writing into protected Program Files folders, which reduces update prompts and duplicate-install problems. It also matches the fact that this builder is a personal content tool."),
    ("Use GitHub Releases for installer distribution", "GitHub Releases provide versioned downloads, stable latest asset names, rollback access, and an update source the app can check without running a custom update server."),
    ("Use GitHub Actions for release automation", "Actions make releases repeatable. The same cloud build installs dependencies, runs electron-builder, creates artifacts, and publishes release assets, which is less error-prone than manual packaging."),
    ("Use Cloudflare Worker for AI", "The public website cannot safely contain AI provider secrets. A Worker can receive chat requests, add portfolio context, call an AI provider, and return an answer while keeping secrets server-side."),
    ("Use cached compile outputs", "Repeated C, C++, Java, and SystemVerilog runs should not recompile unchanged source every time. Caching reduces waiting, while Rebuild / run keeps a manual path for forcing a clean compile."),
]


TOOL_BUILD_ROWS = [
    ("pnpm", "Installs Node dependencies from pnpm-lock.yaml.", "Used locally and in GitHub Actions before packaging."),
    ("Electron", "Runs the desktop application shell.", "Loads main.cjs, starts the backend, and opens the builder window."),
    ("electron-builder", "Packages the Electron app.", "Produces portable and NSIS Windows artifacts."),
    ("NSIS", "Builds the Windows installer wizard.", "Handles install, update, shortcut, and uninstall behavior."),
    ("Git", "Moves source code and website files between local machine and GitHub.", "Used for development branches and Apply to site publishing."),
    ("GitHub Actions", "Runs release automation in the cloud.", "Builds installer and portable app when main or release tags are pushed."),
    ("Wrangler", "Deploys Cloudflare Worker code.", "Publishes cloudflare/portfolio-ai-worker.js to Cloudflare."),
    ("Compilers", "Build or run project source code.", "Used inside Compile Code for C, C++, Java, Python, JavaScript, Verilog, and SystemVerilog."),
]

ENTRY_POINTS = [
    ("Installed Windows app", "C:\\Users\\<you>\\AppData\\Local\\Programs\\OMB Portfolio Builder\\OMB Portfolio Builder.exe", "This is what a normal user double-clicks. It starts Electron, which starts the builder."),
    ("Electron package entry", "package.json -> main: main.cjs", "Electron reads package.json and uses main.cjs as the desktop application's first JavaScript file."),
    ("Electron main process", "main.cjs", "Creates the desktop window, starts the local backend, owns app menus, handles update handoff, and manages app lifecycle."),
    ("Local backend entry", "server.mjs", "Defines local API endpoints for saving drafts, parsing projects, publishing, compiler workspaces, authentication checks, and file operations."),
    ("Builder screen entry", "template-preview.html", "The HTML document loaded inside the Electron window. It includes the builder's CSS and JavaScript."),
    ("Builder logic entry", "template-preview.js", "The main frontend code for editing, rich text, projects, previews, compile code, guide windows, and publishing controls."),
    ("Builder style entry", "template-preview.css", "Controls the appearance of builder windows, dialogs, rich editors, dark mode, compile terminal, and guide panels."),
    ("Public website entry", "index.html", "The file a recruiter loads in the browser. It links to styles.css, script.js, projects.json, and assets."),
    ("Public website logic", "script.js", "Renders projects, search, website windows, AI chat UI, contact behavior, and mobile/desktop interactions."),
    ("Public website style", "styles.css", "Controls the public portfolio layout, mobile responsiveness, contact panel, AI panel, and project windows."),
    ("Public data entry", "projects.json", "The generated catalog the website reads to display projects and portfolio sections."),
    ("Release automation entry", ".github/workflows/build-windows-builder.yml", "GitHub reads this workflow when main or builder tags are pushed and builds installer artifacts."),
    ("Main branch guard entry", ".github/workflows/main-branch-gate.yml", "GitHub reads this workflow to reject direct or wrong-source changes into main."),
    ("Cloudflare AI entry", "cloudflare/portfolio-ai-worker.js", "The Worker receives public AI chat requests and talks to the configured AI provider."),
    ("Installer customization entry", "build/installer.nsh", "NSIS include file that customizes install/update behavior beyond electron-builder defaults."),
]


PROGRAMMING_SYNTAX_LESSONS = [
    (
        "JavaScript function",
        "function updateCompileTerminalPanel(project, file) {\n  const terminal = file?.lastResult?.terminal || \"No output yet\";\n  terminalElement.textContent = terminal;\n}",
        "The word function defines reusable behavior. The values inside parentheses are inputs. Curly braces contain the steps. const creates a value that should not be reassigned. The ?. operator means 'only read this if the left side exists.'",
    ),
    (
        "JavaScript async/await",
        "async function saveCompileFile(project, file) {\n  const response = await fetch('/api/code/save', {\n    method: 'POST',\n    headers: { 'Content-Type': 'application/json' },\n    body: JSON.stringify(file)\n  });\n  return response.json();\n}",
        "async means the function can wait for slow work. await pauses until fetch returns. fetch is how the browser side talks to server.mjs. JSON.stringify turns JavaScript data into a string for the request body.",
    ),
    (
        "JavaScript event handling",
        "sectionContent.addEventListener('click', async (event) => {\n  const button = event.target.closest('[data-compile-run]');\n  if (button) await compileActiveFile(project, file);\n});",
        "addEventListener says 'when this thing happens, run this code.' event.target is what was clicked. closest searches upward for a matching element. data-* attributes let buttons identify commands without hard-coding many IDs.",
    ),
    (
        "HTML structure",
        "<section class=\"compile-terminal-panel\" aria-label=\"Compiler terminal output\">\n  <div class=\"compile-terminal-heading\">OMB Local Terminal</div>\n  <pre class=\"compile-terminal\"></pre>\n</section>",
        "HTML is the skeleton. section groups a meaningful area. class names let CSS style it and JavaScript find it. aria-label helps accessibility tools. pre preserves terminal spacing and line breaks.",
    ),
    (
        "CSS selector",
        ".compile-terminal-panel.is-focused {\n  border-color: #0284c7;\n  box-shadow: 0 0 0 3px rgba(14, 165, 233, 0.18);\n}",
        "CSS selectors choose what to style. A dot means class. Two classes together mean the element must have both. Properties like border-color and box-shadow change appearance.",
    ),
    (
        "JSON data",
        "{\n  \"title\": \"Voltage Controlled Oscillator\",\n  \"category\": \"Analog and Mixed Signal\",\n  \"sections\": []\n}",
        "JSON stores structured data. Keys are in quotes. Strings are in quotes. Arrays use square brackets. Objects use curly braces. projects.json uses this kind of shape for public site data.",
    ),
    (
        "GitHub Actions YAML",
        "on:\n  push:\n    branches:\n      - main\njobs:\n  build-windows:\n    runs-on: windows-latest",
        "YAML uses indentation to show nesting. This says the workflow runs on pushes to main. jobs are units of work. runs-on chooses the temporary machine type.",
    ),
    (
        "Cloudflare Worker handler",
        "export default {\n  async fetch(request, env) {\n    const body = await request.json();\n    return Response.json({ answer: 'Hello' });\n  }\n};",
        "A Worker responds to web requests. fetch is the entry function Cloudflare calls. request contains what the browser sent. env contains configured secrets and bindings.",
    ),
    (
        "Node/Electron require",
        "const { app, BrowserWindow } = require('electron');\napp.whenReady().then(() => {\n  const win = new BrowserWindow({ width: 1200, height: 800 });\n});",
        "require imports Node modules. Electron provides app and BrowserWindow. app.whenReady waits until Electron is ready to create windows.",
    ),
    (
        "Regular expression",
        "const isUrl = /^(https?:\\/\\/|www\\.)/i.test(text);",
        "A regular expression is a pattern matcher. This one checks whether text starts with http://, https://, or www. The i flag means case-insensitive.",
    ),
]


SHELL_SYNTAX_LESSONS = [
    (
        "PowerShell variable",
        "$setup = Get-ChildItem dist -Filter \"*.exe\" | Select-Object -First 1",
        "PowerShell variables start with $. The pipeline | sends output from one command into the next command. This finds an installer file and stores it in $setup.",
    ),
    (
        "PowerShell condition",
        "if (-not $setup) {\n  throw \"Windows setup installer was not found.\"\n}",
        "if checks a condition. -not means false/not present. throw stops the script with an error message. GitHub Actions uses this style to fail clearly.",
    ),
    (
        "PowerShell copy",
        "Copy-Item $setup.FullName \"dist/OMB-Portfolio-Builder-Setup-latest-x64.exe\" -Force",
        "Copy-Item copies a file. $setup.FullName is the full path. -Force overwrites an existing file. The workflow uses this to create stable latest download names.",
    ),
    (
        "Bash environment variable",
        "if [ \"${{ github.head_ref }}\" != \"development\" ]; then\n  echo \"Wrong source branch\"\n  exit 1\nfi",
        "Bash condition syntax uses square brackets. != means not equal. exit 1 means fail the job. The main branch gate uses this pattern.",
    ),
    (
        "Git command shape",
        "git push origin development",
        "git is the program. push is the action. origin is the remote GitHub repository nickname. development is the branch being uploaded.",
    ),
    (
        "pnpm script shape",
        "pnpm run dist",
        "pnpm reads package.json, finds the dist script, and runs the command stored there. In this repo, dist runs electron-builder.",
    ),
    (
        "NSIS installer macro style",
        "!macro customInstall\n  DetailPrint \"Installing OMB Portfolio Builder\"\n!macroend",
        "NSIS uses macro blocks for installer customization. DetailPrint writes messages to the installer progress details.",
    ),
    (
        "Wrangler deploy",
        "cd cloudflare\nnpx wrangler deploy",
        "cd changes folders. npx runs a package command. wrangler deploy uploads the Worker code to Cloudflare.",
    ),
]

CACHING_METHODS = [
    ("HTTP no-store for local APIs", "template-preview.js and script.js call fetch with cache: 'no-store' or timestamp query strings for important local data.", "Prevents the browser from showing stale builder catalogs, templates, app update checks, security reports, and public project catalogs."),
    ("Update snooze and skip cache", "template-preview.js stores omb-snoozed-update-version, omb-snoozed-update-at, and omb-skipped-update-version in localStorage.", "Lets the app remember that a user temporarily dismissed a specific update version."),
    ("Builder preferences cache", "template-preview.js stores builder preferences in localStorage through preferenceStorageKey.", "Keeps local UI choices such as light/dark mode without publishing them to the website."),
    ("Publishing authorization cache", "server.mjs writes .omb-publish-session.json with repository, branch, checkedAt, expiresAt, success history, scope, and trust days.", "Avoids asking for GitHub authorization every time while still scoping trust to the same user, repo, branch, and machine context."),
    ("Extended trust cache", "server.mjs extends the authorization trust window when the same target has more than three successful authorizations in a week.", "Reduces repeated login friction for the owner while still limiting trust to a known target."),
    ("Compiler tool path cache", "server.mjs uses compileToolCache to remember where tools such as gcc, g++, javac, java, iverilog, and vvp were found.", "Speeds repeated compiler detection so every run does not search the full PATH again."),
    ("Compiler tool version cache", "server.mjs uses compileToolVersionCache keyed by tool path.", "Avoids repeatedly running version commands for tools that rarely change during a builder session."),
    ("Compile artifact cache", "server.mjs uses compileCacheKey and compileCacheDirectory under compile-code/.build-cache.", "Avoids recompiling unchanged C, C++, Java, Verilog, and SystemVerilog source unless Rebuild / run is requested."),
    ("Terminal result cache in builder state", "template-preview.js stores file.lastResult.terminal and workspace.terminal.", "Keeps the latest terminal output visible after re-rendering a builder section."),
    ("Saved source metadata", "template-preview.js and server.mjs store savedAt, savedPath, language, fileName, and dirty flags.", "Allows the UI to show whether a source is draft, saved, compiled, or needs fixes."),
]


INSTALLER_FUNCTIONS = [
    ("OMBToolsPageCreate", "Creates the installer tools/shortcuts page with checkboxes for desktop shortcut and publishing tools."),
    ("OMBToolsPageLeave", "Validates that updates stay in the existing install location instead of creating a second app copy."),
    ("OMBCheckGitAvailable", "Checks common Git installation paths and then tries git --version."),
    ("OMBCheckGitCredentialManagerAvailable", "Checks whether Git Credential Manager is available."),
    ("OMBInstallGitIfNeeded", "Installs or repairs Git for Windows through winget when publishing tools are missing."),
    ("OMBReadExistingAppVersion", "Reads the installed executable version from Windows file metadata."),
    ("OMBStopIfExistingInstallIsCurrent", "Stops setup if the installed version is already current or newer."),
    ("OMBCheckKnownBuilderWorkspace", "Rejects known legacy or development builder folders before a fresh install."),
    ("OMBWriteDuplicateScanScript", "Writes a PowerShell scanner that searches for duplicate builder copies across the system."),
    ("OMBScanForDuplicateBuilderCopies", "Runs the duplicate scanner and stops setup if another builder copy is found."),
    ("OMBStopBuilderProcessesForUpdate", "Stops running builder processes before an in-place update."),
    ("OMBDisableBuiltInOldUninstallerForInPlaceUpdate", "Removes old uninstall handoff registry values so updates can overwrite in place."),
    ("OMBUninstallExistingInstallIfPresent", "Finds registered or legacy installs, decides whether this is an update, and sets the installation directory."),
    ("customInit", "Installer entry macro that checks existing installs, duplicate workspaces, and default checkbox states."),
    ("customPageAfterChangeDir", "Adds the custom tools/shortcuts page after the directory chooser."),
    ("customInstall", "Creates/removes shortcuts, installs Git if selected, and cleans legacy app folders."),
    ("OMBUninstallPageCreate", "Creates the custom uninstall explanation page."),
    ("customUnInstall", "Removes desktop shortcuts during uninstall."),
]


FUNCTION_GROUPS = [
    ("main.cjs", "Electron app startup and desktop shell", [
        ("dispatchBuilderMenuAction", "Sends top-menu commands into the builder window."),
        ("quitForBuilderUpdate", "Quits cleanly when the updater has started."),
        ("setBuilderFullScreen", "Controls fullscreen state."),
        ("createAppMenu", "Builds the File/View/Help style application menu."),
        ("preparePackagedWorkspace", "Copies packaged site resources into a usable workspace."),
        ("resolveWorkspaceRoots", "Finds builder and portfolio workspace folders."),
        ("findFreePort", "Finds an available local port for the backend."),
        ("startBuilderServer", "Starts server.mjs in the background."),
        ("createWindow", "Creates the Electron BrowserWindow and loads the builder."),
        ("boot", "Main startup routine: prepare workspace, start backend, open window."),
    ]),
    ("server.mjs", "Local backend, publishing, AI, and compiler runner", [
        ("securityHeaders", "Adds no-store and security headers to local API responses."),
        ("sendJson", "Sends JSON API responses."),
        ("detectCodeLanguageFromSource", "Infers code language from source and file name."),
        ("findTool", "Locates compiler/runtime executables and caches the answer."),
        ("runProcess", "Runs a child process and captures stdout, stderr, exit code, timeout, and elapsed time."),
        ("saveCompileSource", "Writes compile-code source files into the local compile workspace."),
        ("compileCacheKey", "Builds a hash-like key from source, language, compiler, and runtime details."),
        ("compileAndRunCode", "Compiles/runs supported source code and returns terminal output."),
        ("readPublishAuthCache", "Reads the local GitHub publish authorization cache."),
        ("writePublishAuthCache", "Writes the local publish authorization cache with trust history."),
        ("publishAuthCacheIsFresh", "Checks whether cached authorization still matches the target and has not expired."),
        ("configurePublishTarget", "Saves and verifies a publishing target."),
        ("syncFromPublishTarget", "Loads compatible portfolio files from the GitHub target."),
        ("downloadAndLaunchAppUpdate", "Downloads the installer and starts the update handoff."),
        ("publishSiteChanges", "Commits and pushes generated public website files."),
        ("handlePortfolioAi", "Handles local/website AI assistant requests."),
        ("handleApi", "Routes /api endpoints to the correct backend function."),
    ]),
    ("template-preview.js", "Builder frontend and editor behavior", [
        ("renderCompileCodeSection", "Draws the Compile Code workspace UI."),
        ("showCompileOutput", "Scrolls directly to the terminal and focuses it."),
        ("copyCompileOutput", "Copies the current terminal output to the clipboard."),
        ("clearCompileOutput", "Clears stored terminal output for the active source."),
        ("compileActiveFile", "Saves the source, calls /api/code/compile, and updates terminal/messages."),
        ("appendCompileCodeToProject", "Inserts source code into a selected project rich-text destination."),
        ("renderSectionContent", "Draws the currently selected project builder section."),
        ("populateRichEditors", "Turns saved rich blocks into editable rich editor UI."),
        ("saveRichEditorToProject", "Serializes rich editor content back into project state."),
        ("scheduleAutosave", "Queues local draft saves after edits."),
        ("schedulePreviewRender", "Queues preview rebuilding after relevant changes."),
        ("configurePublishTarget", "Frontend path for target setup and authentication flow."),
    ]),
    ("script.js", "Public website runtime", [
        ("renderFunFacts", "Displays the fun facts area."),
        ("openProjectWindow", "Opens a project in the recruiter-facing website UI."),
        ("renderProjectDetail", "Renders a selected project from projects.json."),
        ("performSearch", "Searches visible website content and project data."),
        ("highlightSearchTerm", "Highlights matched search text."),
        ("handlePortfolioAiSubmit", "Sends visitor AI questions to the backend endpoint."),
        ("navigateSectionHistory", "Supports back/forward behavior inside website windows."),
    ]),
]

SYSTEM_LAYER_SECTIONS = [
    (
        "Frontend Layer: What The User Sees And Clicks",
        "The frontend is the part of the software that draws screens, buttons, text fields, windows, menus, project cards, previews, the AI chat interface, and the public website pages. This project has two frontend surfaces: the private builder frontend and the public portfolio frontend.",
        [
            ("Builder frontend files", "template-preview.html, template-preview.css, and template-preview.js."),
            ("Public frontend files", "index.html, styles.css, script.js, and electronics-search.js."),
            ("Main job", "Render UI, capture clicks and typing, call backend APIs, show previews, and display returned data."),
            ("What it should not do", "It should not store private API keys, directly run compilers, or bypass publishing authorization."),
            ("Beginner mental model", "The frontend is the dashboard and steering wheel. It lets the user drive, but the engine work happens elsewhere."),
        ],
        "const response = await fetch('/api/code/compile', {\n  method: 'POST',\n  headers: { 'Content-Type': 'application/json' },\n  body: JSON.stringify(payload)\n});",
    ),
    (
        "Backend Layer: What Does The Heavy Work",
        "The backend is the local service in server.mjs. It accepts requests from the builder frontend and does work the browser frontend should not do directly: file access, Git operations, compiler execution, publish target checks, update downloads, and parser output.",
        [
            ("Backend file", "server.mjs."),
            ("Main job", "Expose safe local /api endpoints for builder actions."),
            ("Important responsibilities", "Save drafts, parse projects, manage publish target, run Git, run compilers, fetch update info, handle AI requests locally when needed."),
            ("Security boundary", "The backend validates paths so file operations stay inside allowed roots."),
            ("Beginner mental model", "The backend is the workshop behind the counter. The frontend asks for work; the backend uses tools."),
        ],
        "if (url.pathname === '/api/code/compile') {\n  const payload = await readRequestJson(request);\n  const result = await compileAndRunCode(payload);\n  return sendJson(response, 200, { ok: result.ok, result });\n}",
    ),
    (
        "AI Layer: How The Portfolio Assistant Works",
        "The AI assistant is a visitor-facing helper. The public frontend collects a question, sends it to an endpoint, and receives a conversational answer. The AI should answer portfolio-specific questions using portfolio context and answer general electronics questions using general knowledge only when appropriate.",
        [
            ("Public UI", "The chat box lives in the portfolio contact area and is controlled by script.js."),
            ("Public endpoint", "The public website points to /api/portfolio-ai through the meta endpoint and fetch logic."),
            ("Backend endpoint", "cloudflare/portfolio-ai-worker.js is the preferred public backend for deployed AI."),
            ("Context source", "Portfolio text, project data, public links, and selected public sources can be added to the model request."),
            ("Answer rule", "Answer first in natural language, then provide related portfolio links only when they actually match the question."),
        ],
        "const endpoint = document.querySelector('meta[name=\"portfolio-ai-endpoint\"]')?.content || '/api/portfolio-ai';\nconst response = await fetch(endpoint, {\n  method: 'POST',\n  body: JSON.stringify({ question, context })\n});",
    ),
    (
        "Cloudflare Worker Layer: Public Backend For The AI",
        "The Cloudflare Worker is a small serverless backend that runs on Cloudflare instead of the visitor's browser. It receives AI chat requests from the website, reads server-side secrets, calls an AI provider or Worker AI, and returns the answer.",
        [
            ("Worker file", "cloudflare/portfolio-ai-worker.js."),
            ("Deployment config", "cloudflare/wrangler.toml."),
            ("Deploy tool", "Wrangler."),
            ("Why Worker exists", "The public browser cannot safely contain private API keys."),
            ("Beginner mental model", "The Worker is a secure receptionist between the public website and the AI provider."),
        ],
        "export default {\n  async fetch(request, env) {\n    const body = await request.json();\n    const key = env.OPENAI_API_KEY;\n    // Call the model server-side, never from the browser with a visible key.\n    return Response.json({ answer: '...' });\n  }\n};",
    ),
    (
        "Cloudflare API Keys And Secrets",
        "API keys are passwords for services. They must not be committed to GitHub, placed in index.html, placed in script.js, placed in projects.json, or pasted into public website code. The deployed Worker should read secrets from Cloudflare's encrypted secret storage.",
        [
            ("OpenAI secret name", "OPENAI_API_KEY, stored with wrangler secret put OPENAI_API_KEY."),
            ("Where it lives", "Cloudflare secret storage, exposed to the Worker as env.OPENAI_API_KEY at runtime."),
            ("Where it must not live", "GitHub repository files, browser JavaScript, public JSON, README examples with real values, screenshots, or issue comments."),
            ("Local testing", "Use local environment files only if they are ignored by Git, and never publish real keys."),
            ("Rotation rule", "If a key is exposed, revoke it and create a new one immediately."),
        ],
        "cd cloudflare\nnpx wrangler secret put OPENAI_API_KEY\nnpx wrangler deploy",
    ),
    (
        "Cloudflare Routing And DNS For The AI",
        "DNS points the custom domain to the public website. Cloudflare routing can also send /api/portfolio-ai to the Worker. That lets the same domain serve both the static portfolio and the AI backend endpoint.",
        [
            ("Static pages", "index.html, styles.css, script.js, projects.json, assets, and docs are served as static files."),
            ("AI route", "/api/portfolio-ai should be routed to the Cloudflare Worker."),
            ("Proxying", "Cloudflare proxy should be enabled for routes that need Worker handling."),
            ("Failure symptom", "If the AI falls back to canned answers, the Worker route or secret may be missing."),
            ("Debug start", "Check browser Network tab for /api/portfolio-ai status code and Worker logs in Cloudflare."),
        ],
        "Website browser request:\nGET https://mauriceotieno.com/index.html\nPOST https://mauriceotieno.com/api/portfolio-ai",
    ),
    (
        "API Request Flow: From Click To Result",
        "An API request is a structured message from one software layer to another. In this project, the builder frontend calls the local backend, while the public website calls the Cloudflare Worker for AI.",
        [
            ("Builder request", "template-preview.js -> server.mjs."),
            ("Website AI request", "script.js -> Cloudflare Worker."),
            ("Request body", "Usually JSON describing what the user wants."),
            ("Response body", "Usually JSON describing success, errors, or generated output."),
            ("No-store behavior", "Important requests avoid browser cache so the builder sees fresh state."),
        ],
        "Frontend sends JSON:\n{ \"question\": \"What projects use STM32?\" }\n\nBackend returns JSON:\n{ \"answer\": \"Maurice used STM32 in...\", \"links\": [...] }",
    ),
]

FLOW_WALKTHROUGHS = [
    (
        "Builder Startup Flow",
        "electron-runtime",
        [
            ("User action", "The user double-clicks OMB Portfolio Builder.exe or a desktop shortcut that points to it."),
            ("Windows loads Electron", "The packaged executable starts the Electron runtime and reads package.json to find main.cjs."),
            ("main.cjs boots", "boot waits for app readiness, resolves workspace folders, starts the backend, and creates the BrowserWindow."),
            ("server.mjs starts", "The local backend binds to a free localhost port and exposes local /api routes."),
            ("Builder HTML loads", "The BrowserWindow opens template-preview.html from the local backend origin."),
            ("Frontend initializes", "template-preview.js loads catalog/templates, restores local data, prepares rich editors, and draws the builder screen."),
            ("User sees app", "At this point the app behaves like a desktop app even though its UI is web technology."),
        ],
        "If startup fails, separate the problem into executable launch, Electron main process, backend server, port binding, and frontend load.",
    ),
    (
        "Edit, Save, Parse, Preview Flow",
        "builder-to-site-files",
        [
            ("User edits", "The owner types text, adds images, creates projects, attaches files, or changes appearance in the builder."),
            ("Frontend state changes", "template-preview.js updates in-memory builder/project state and marks affected content dirty."),
            ("Autosave or explicit save", "Save project, Save all sections, or Save draft sends data to server.mjs."),
            ("Backend writes local draft", "server.mjs writes local JSON and asset files into the managed builder workspace."),
            ("Parser normalizes", "The parser removes builder-only controls, decides what is public, and creates clean project/section data."),
            ("Preview renders", "The builder reads parsed output and renders a preview that should match the public website."),
            ("Publish mirror updates", "Generated public-safe files are copied into the portfolio publish workspace when ready."),
        ],
        "The key idea is that raw editing data and public parsed data are different. The parser is the boundary.",
    ),
    (
        "Apply To Site Flow",
        "release-pipeline",
        [
            ("Save requirement", "The builder should require Save draft before Apply to site so local data and parsed output are current."),
            ("Target check", "server.mjs reads the publishing target, branch, repository URL, custom domain, and cached auth state."),
            ("Authorization check", "If a fresh scoped authorization cache exists, it is reused; otherwise GitHub write access is verified."),
            ("Sync or stage", "The publish mirror is synchronized with generated website files."),
            ("Git status", "Git determines which public files changed."),
            ("Commit", "The builder creates a commit with the changed website files."),
            ("Push", "Git pushes to the target repository branch."),
            ("Host serves", "GitHub Pages or another host serves the updated static files."),
        ],
        "Publishing is intentionally stricter than saving. Saving is local. Apply to site changes the public website.",
    ),
    (
        "Compile Code Flow",
        "compile-code-flow",
        [
            ("Create/import source", "The user creates a compile source file or imports one from disk."),
            ("Choose language", "The builder stores a language profile such as C, C++, Java, JavaScript, Python, Verilog, or SystemVerilog."),
            ("Save source", "server.mjs writes source into the compile-code workspace."),
            ("Detect tools", "server.mjs finds compiler/runtime tools and caches their paths."),
            ("Build cache key", "The backend hashes source, language, compiler, runtime, and relevant flags."),
            ("Require HDL testbench", "For Verilog and SystemVerilog, the backend refuses to simulate a design-only file. At least one HDL file must be marked Testbench."),
            ("Require HDL waveform dump", "The HDL testbench must call $dumpfile and $dumpvars so the builder has signal-over-time data to draw in the scope."),
            ("Compile if needed", "If a matching artifact exists and Rebuild was not requested, the builder reuses it."),
            ("Run", "The runtime executes the binary/class/simulation/script and captures stdout/stderr."),
            ("Parse VCD scope", "For HDL simulation, server.mjs reads the generated VCD file and returns signal changes to template-preview.js."),
            ("Show output", "template-preview.js displays terminal text immediately and records messages in the log."),
        ],
        "Compile Code has two purposes: prove code runs locally and optionally append clean source evidence into a project section.",
    ),
    (
        "AI Chat Flow",
        "cloudflare-ai-flow",
        [
            ("Visitor asks", "A recruiter or visitor types a question in the portfolio AI panel."),
            ("Frontend classifies context", "script.js collects available site context and sends the question to the configured endpoint."),
            ("Worker receives", "cloudflare/portfolio-ai-worker.js receives the POST request."),
            ("Secrets stay hidden", "The Worker reads API keys from Cloudflare environment secrets, not from browser code."),
            ("Context is prepared", "Portfolio data, project snippets, public files, or public links can be included where useful."),
            ("Model answers", "The Worker calls OpenAI or Cloudflare Worker AI and receives a response."),
            ("Website renders answer", "script.js displays the answer in the chat, with related links only when relevant."),
        ],
        "The AI layer should not be a generic link sprinkler. It should answer the question first, then cite relevant portfolio locations when they match the conversation.",
    ),
    (
        "Installer Update Flow",
        "release-pipeline",
        [
            ("App checks release", "The builder checks GitHub Releases for a newer version."),
            ("User chooses Update", "The app starts the update handoff instead of acting like a fresh install."),
            ("Installer downloads", "server.mjs downloads the latest installer to a safe update folder."),
            ("App closes", "Electron emits the update-started event and quits so files can be replaced."),
            ("Updater script runs", "A detached PowerShell/cmd handoff runs the installer even after the app exits."),
            ("Installer detects current install", "NSIS finds the existing AppData install and updates it in place."),
            ("App reopens", "The updater waits for the executable and relaunches the updated app."),
        ],
        "Update problems usually mean one of three things: the old app did not close, the installer treated it as a new install, or the relaunch handoff failed.",
    ),
]


DATA_CONTRACTS = [
    (
        "Compile Code Request",
        "Sent by template-preview.js to server.mjs when saving, beautifying, compiling, or running code.",
        "{\n  \"projectId\": \"project-id\",\n  \"fileId\": \"source-id\",\n  \"title\": \"Counter design\",\n  \"fileName\": \"counter.sv\",\n  \"language\": \"systemverilog\",\n  \"role\": \"design\",\n  \"code\": \"module counter(...); ... endmodule\",\n  \"stdin\": \"optional input\",\n  \"workspaceFiles\": [\n    { \"fileName\": \"counter.sv\", \"language\": \"systemverilog\", \"role\": \"design\", \"code\": \"...\" },\n    { \"fileName\": \"tb_counter.sv\", \"language\": \"systemverilog\", \"role\": \"testbench\", \"code\": \"... $dumpfile ... $dumpvars ...\" }\n  ],\n  \"forceRebuild\": false\n}",
        "The backend must validate language, safe file name, safe project folder, rebuild intent, and for HDL simulation whether at least one workspace file is a testbench with waveform dump calls.",
    ),
    (
        "Compile Code Response",
        "Returned by server.mjs after compile/run.",
        "{\n  \"ok\": true,\n  \"result\": {\n    \"language\": \"systemverilog\",\n    \"terminal\": \"simulation output...\",\n    \"saved\": { \"sourcePath\": \"...\", \"savedAt\": \"...\" },\n    \"waveform\": {\n      \"source\": \"waveform.vcd\",\n      \"timeScale\": \"1 ns\",\n      \"maxTime\": 100,\n      \"signals\": [{ \"name\": \"tb_counter.clk\", \"width\": 1, \"changes\": [{ \"time\": 0, \"value\": \"0\" }] }]\n    }\n  }\n}",
        "template-preview.js stores result.terminal in file.lastResult and workspace.terminal, stores result.waveform for HDL, repaints the terminal panel, and draws the signal scope.",
    ),
    (
        "Publish Target",
        "Stored locally so the builder knows which repository can receive generated website files.",
        "{\n  \"remote\": \"https://github.com/user/user.github.io.git\",\n  \"branch\": \"main\",\n  \"customDomain\": \"example.com\",\n  \"authorizationChecked\": true,\n  \"expiresAt\": \"2026-07-05T...Z\"\n}",
        "This data is local control data, not public portfolio content. Authentication must be verified before writing to the website target.",
    ),
    (
        "Publishing Authorization Cache",
        "Written to .omb-publish-session.json after successful verification.",
        "{\n  \"remote\": \"https://github.com/...\",\n  \"branch\": \"main\",\n  \"repository\": \"owner/repo\",\n  \"scope\": \"machine/user scope\",\n  \"checkedAt\": \"...\",\n  \"expiresAt\": \"...\",\n  \"successHistory\": [\"...\"]\n}",
        "Freshness requires same remote, branch, repository, scope, and valid expiration time.",
    ),
    (
        "Project Catalog Item",
        "Public website data consumed by script.js from projects.json.",
        "{\n  \"id\": \"vco-project\",\n  \"title\": \"Voltage Controlled Oscillator\",\n  \"category\": \"Analog and Mixed Signal\",\n  \"overview\": { \"blocks\": [...] },\n  \"sections\": [...],\n  \"assets\": [...]\n}",
        "The public catalog should contain only content intended for visitors. Builder-only controls and private fields should not appear.",
    ),
    (
        "Rich Text Block",
        "Internal content unit used by builder editors and parser rendering.",
        "{\n  \"type\": \"paragraph | image | formula | code | link\",\n  \"text\": \"optional text\",\n  \"html\": \"optional formatted html\",\n  \"src\": \"optional asset path\",\n  \"language\": \"optional code language\"\n}",
        "Blocks let the parser preserve text formatting, hard line breaks, images, formulas, and code without treating everything as plain text.",
    ),
    (
        "AI Chat Request",
        "Sent by the public website to the AI endpoint.",
        "{\n  \"question\": \"What embedded projects does Maurice have?\",\n  \"history\": [...],\n  \"portfolioContext\": {...}\n}",
        "The frontend should send enough context for a good answer, but secrets and private files must stay out of the browser.",
    ),
    (
        "AI Chat Response",
        "Returned by Cloudflare Worker or local AI handler.",
        "{\n  \"answer\": \"Maurice has embedded work involving...\",\n  \"links\": [\n    { \"label\": \"PWM VCO\", \"href\": \"#projects\" }\n  ]\n}",
        "The answer should be conversational first. Links should be specific to the prompt, not generic portfolio shortcuts.",
    ),
]


DEEP_DETAIL_TOPICS = [
    (
        "How To Read The Repository Like A Map",
        "Start with package.json to learn what the app is and what scripts exist. Then open main.cjs to understand desktop startup. Then open server.mjs to understand local backend capabilities. Then open template-preview.html/css/js to understand the builder UI. Finally open index.html/styles.css/script.js/projects.json to understand the public website.",
    ),
    (
        "How Files Communicate Without Directly Importing Each Other",
        "Some files communicate by imports or requires, but many communicate through files and HTTP requests. template-preview.js does not directly edit server internals; it calls /api endpoints. server.mjs writes projects.json. script.js later reads projects.json. This loose coupling keeps the public site simpler.",
    ),
    (
        "How To Debug A Feature From Button To File",
        "Find the button text in template-preview.html or generated template-preview.js markup. Find the data-* attribute on the button. Search template-preview.js for that dataset key. Follow the handler to the fetch call or state update. If it calls /api, find that route in server.mjs. Then inspect what file is read or written.",
    ),
    (
        "Why Generated Files Need Discipline",
        "A generated file can be overwritten. For example, projects.json should come from the parser, not from hand edits. If you hand edit generated output, the builder may replace your changes during the next Save draft or Apply to site.",
    ),
    (
        "How To Think About Errors",
        "Errors usually name the layer that failed. A JavaScript console error points toward frontend code. A terminal compile error points toward compiler/tool/source code. A Git error points toward branch, remote, credentials, or repository state. A 404/500 from /api points toward route or backend logic.",
    ),
    (
        "How To Decide Where A New Feature Belongs",
        "If the feature changes what the user sees while editing, start in template-preview.js/css/html. If it writes files, runs commands, authenticates, or publishes, start in server.mjs. If it changes recruiter-facing behavior, start in script.js/styles.css/index.html. If it changes install/update, start in package.json and build/installer.nsh.",
    ),
]


def tracked_files() -> list[str]:
    result = subprocess.run(["git", "ls-files"], cwd=REPO, text=True, capture_output=True, check=True)
    return [line.strip() for line in result.stdout.splitlines() if line.strip()]


GENERATED_REFERENCE_PREFIXES = (
    "docs/code-reference/",
    "docs/code-reference-docx/",
    "docs/code-reference-pdf/",
    "docs/file-reference-docx/",
    "docs/file-reference-pdf/",
    "docs/important-code-reference-docx/",
    "docs/important-code-reference-pdf/",
    "docs/guide-diagrams/",
)

GENERATED_REFERENCE_FILES = {
    "docs/OMB_Portfolio_Builder_Complete_Guide.docx",
    "docs/OMB_Portfolio_Builder_Complete_Guide.pdf",
    "docs/OMB_Portfolio_Builder_Code_Reference.docx",
    "docs/OMB_Portfolio_Builder_Code_Reference.pdf",
    "docs/OMB_Portfolio_Builder_File_Reference.docx",
    "docs/OMB_Portfolio_Builder_File_Reference.pdf",
    "docs/OMB_Portfolio_Builder_Important_Code_Reference.docx",
    "docs/OMB_Portfolio_Builder_Important_Code_Reference.pdf",
}


def is_generated_reference_file(repo_file: str) -> bool:
    return repo_file in GENERATED_REFERENCE_FILES or repo_file.startswith(GENERATED_REFERENCE_PREFIXES)


def primary_tracked_files(files: list[str]) -> list[str]:
    return [file for file in files if not is_generated_reference_file(file)]


IMPORTANT_CODE_FILES = [
    "main.cjs",
    "server.mjs",
    "template-preview.js",
    "template-preview.html",
    "template-preview.css",
    "script.js",
    "index.html",
    "styles.css",
    "cloudflare/portfolio-ai-worker.js",
    "cloudflare/wrangler.toml",
    "build/installer.nsh",
    "docs/build_complete_guide.py",
    "builder-rich-future-sections.js",
    "electronics-search.js",
    "package.json",
    "project-templates.json",
    "projects.json",
    ".github/workflows/build-windows-builder.yml",
    ".github/workflows/main-branch-gate.yml",
]


def important_code_files(files: list[str]) -> list[str]:
    available = set(files)
    return [file for file in IMPORTANT_CODE_FILES if file in available and not is_generated_reference_file(file)]


def load_package() -> dict:
    try:
        return json.loads((REPO / "package.json").read_text(encoding="utf-8"))
    except Exception:
        return {}


def load_font(size: int, bold: bool = False) -> ImageFont.ImageFont:
    font_name = "arialbd.ttf" if bold else "arial.ttf"
    font_path = Path(r"C:\Windows\Fonts") / font_name
    if font_path.exists():
        return ImageFont.truetype(str(font_path), size=size)
    return ImageFont.load_default()


def wrap_for_box(label: str, font: ImageFont.ImageFont, max_width: int) -> list[str]:
    lines: list[str] = []
    for part in str(label).splitlines():
        words = part.split()
        if not words:
            lines.append("")
            continue
        current = words[0]
        for word in words[1:]:
            candidate = f"{current} {word}"
            bbox = font.getbbox(candidate)
            if bbox[2] - bbox[0] <= max_width:
                current = candidate
            else:
                lines.append(current)
                current = word
        lines.append(current)
    return lines


def draw_arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], color: str, width: int = 4) -> None:
    draw.line([start, end], fill=color, width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    length = 18
    spread = 0.55
    points = [
        end,
        (int(end[0] - length * math.cos(angle - spread)), int(end[1] - length * math.sin(angle - spread))),
        (int(end[0] - length * math.cos(angle + spread)), int(end[1] - length * math.sin(angle + spread))),
    ]
    draw.polygon(points, fill=color)


def edge_points(source: tuple[int, int, int, int], target: tuple[int, int, int, int]) -> tuple[tuple[int, int], tuple[int, int]]:
    sx1, sy1, sx2, sy2 = source
    tx1, ty1, tx2, ty2 = target
    sc = ((sx1 + sx2) // 2, (sy1 + sy2) // 2)
    tc = ((tx1 + tx2) // 2, (ty1 + ty2) // 2)
    dx = tc[0] - sc[0]
    dy = tc[1] - sc[1]
    if abs(dx) >= abs(dy):
        start = (sx2 if dx >= 0 else sx1, sc[1])
        end = (tx1 if dx >= 0 else tx2, tc[1])
    else:
        start = (sc[0], sy2 if dy >= 0 else sy1)
        end = (tc[0], ty1 if dy >= 0 else ty2)
    return start, end


def make_diagrams() -> dict[str, Path]:
    DIAGRAM_DIR.mkdir(parents=True, exist_ok=True)
    title_font = load_font(38, bold=True)
    box_font = load_font(24, bold=True)
    label_font = load_font(18, bold=False)
    diagram_paths: dict[str, Path] = {}

    for spec in DIAGRAM_SPECS:
        image = Image.new("RGB", (1600, 900), "#F8FBFD")
        draw = ImageDraw.Draw(image)
        draw.rounded_rectangle((24, 24, 1576, 876), radius=26, outline="#B7D7EA", width=3, fill="#FFFFFF")
        draw.text((60, 48), spec["title"], fill="#0B2545", font=title_font)
        draw.line((60, 105, 1540, 105), fill="#B7D7EA", width=3)

        boxes: dict[str, tuple[int, int, int, int]] = {}
        for node_id, label, x, y, w, h, fill in spec["nodes"]:
            boxes[node_id] = (x, y, x + w, y + h)

        for source_id, target_id, *maybe_label in spec["edges"]:
            source = boxes[source_id]
            target = boxes[target_id]
            start, end = edge_points(source, target)
            draw_arrow(draw, start, end, "#2563EB", width=4)
            if maybe_label:
                label = maybe_label[0]
                mid = ((start[0] + end[0]) // 2, (start[1] + end[1]) // 2)
                text_box = label_font.getbbox(label)
                tw = text_box[2] - text_box[0] + 16
                th = text_box[3] - text_box[1] + 10
                if math.hypot(end[0] - start[0], end[1] - start[1]) > max(tw + 28, 95):
                    draw.rounded_rectangle((mid[0] - tw // 2, mid[1] - th // 2, mid[0] + tw // 2, mid[1] + th // 2), radius=9, fill="#EFF6FF", outline="#BFDBFE", width=1)
                    draw.text((mid[0] - tw // 2 + 8, mid[1] - th // 2 + 4), label, fill="#1E3A8A", font=label_font)

        for node_id, label, x, y, w, h, fill in spec["nodes"]:
            draw.rounded_rectangle((x, y, x + w, y + h), radius=20, fill=fill, outline="#668EA8", width=3)
            lines = wrap_for_box(label, box_font, w - 28)
            line_height = 30
            total_h = len(lines) * line_height
            start_y = y + (h - total_h) // 2
            for index, line in enumerate(lines):
                bbox = box_font.getbbox(line)
                tw = bbox[2] - bbox[0]
                draw.text((x + (w - tw) // 2, start_y + index * line_height), line, fill="#102033", font=box_font)

        path = DIAGRAM_DIR / f"{spec['name']}.png"
        image.save(path)
        diagram_paths[spec["name"]] = path
    return diagram_paths


def add_diagram(doc: Document, path: Path, caption: str) -> None:
    if not path.exists():
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(6.35))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(85, 85, 85)


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    run = cell.paragraphs[0].add_run(str(text))
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(1.55)
    table.columns[1].width = Inches(4.95)
    head = table.rows[0].cells
    cell_text(head[0], "Item", True)
    cell_text(head[1], "Explanation", True)
    shade(head[0], "E8EEF5")
    shade(head[1], "E8EEF5")
    for label, value in rows:
        cells = table.add_row().cells
        cell_text(cells[0], label, True)
        cell_text(cells[1], value, False)


def add_code(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["CodeBlock"]
    run = p.add_run((text or "").rstrip())
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")


def add_note(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.35)
    cell = table.cell(0, 0)
    shade(cell, "F4F8FC")
    title_run = cell.paragraphs[0].add_run(title)
    title_run.bold = True
    title_run.font.color.rgb = RGBColor(31, 77, 120)
    title_run.font.size = Pt(10)
    paragraph = cell.add_paragraph(body)
    paragraph.style = "Body Text"


def bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def numbers(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def setup_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)
    for section in doc.sections:
        section.top_margin = Inches(0.65)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.65)
        section.right_margin = Inches(0.65)
        section.header_distance = Inches(0.28)
        section.footer_distance = Inches(0.28)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, color, before, after in [
        ("Heading 1", 14, "2E74B5", 9, 4),
        ("Heading 2", 12, "2E74B5", 7, 3),
        ("Heading 3", 10.5, "1F4D78", 5, 2),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    code = doc.styles.add_style("CodeBlock", 1)
    code.font.name = "Consolas"
    code._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    code.font.size = Pt(7.2)
    code.paragraph_format.space_before = Pt(2)
    code.paragraph_format.space_after = Pt(3)
    code.paragraph_format.line_spacing = 1.0

    header = doc.sections[0].header.paragraphs[0]
    header.text = "OMB Portfolio Builder and Website Complete Guide"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(85, 85, 85)
    return doc


def add_cover(doc: Document, package: dict) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("OMB Portfolio Builder and Website\nComplete Creation Guide")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(11, 37, 69)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("A beginner-friendly manual explaining the desktop app, public website, GitHub workflows, installer, parser, compiler tools, AI backend, and publishing system.")
    r.font.size = Pt(12)
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Repository: otienomaurice/omb-portfolio-builder\nVersion documented: {package.get('version', 'unknown')}\nGenerated: July 4, 2026")
    add_note(doc, "How this guide is written", "This guide starts from the highest-level idea and then drills down. It intentionally explains terms that engineers usually skip: what Electron is, what a workflow is, what a local backend is, why Git branches exist, and why the builder and website are separated.")
    doc.add_page_break()


def add_reading_map(doc: Document) -> None:
    doc.add_heading("Reading Map", level=1)
    doc.add_paragraph("Use this manual in layers. First read the high-level design pages. Then read the app flow pages. After that, use the file-by-file and command pages as a reference when you need to understand a specific file or command.")
    add_table(doc, [
        ("Part 1", "High-level architecture and big-picture mental model."),
        ("Part 2", "Builder app internals: Electron, local backend, parser, rich editors, previews, compile code, and publishing."),
        ("Part 3", "Public website internals: HTML, CSS, JavaScript, projects.json, assets, search, AI, and mobile behavior."),
        ("Part 4", "GitHub, branches, workflows, releases, installers, and updates."),
        ("Part 5", "Every tracked file explained in plain English."),
        ("Part 6", "Commands, tradeoffs, troubleshooting, and safe operating procedures."),
    ])
    doc.add_page_break()


def add_entry_points(doc: Document) -> None:
    doc.add_heading("Where Someone Starts: Entry Points", level=1)
    doc.add_paragraph("An entry point is the first file, executable, or function that starts a flow. This project has several entry points because normal users, developers, GitHub Actions, Cloudflare, and installers all start in different places.")
    add_table(doc, [(name, f"{entry}. {meaning}") for name, entry, meaning in ENTRY_POINTS])
    doc.add_heading("Fastest way to orient yourself", level=2)
    numbers(doc, [
        "If you are using the installed app, start with OMB Portfolio Builder.exe.",
        "If you are debugging desktop startup, start with package.json and main.cjs.",
        "If you are debugging local APIs, start with server.mjs.",
        "If you are debugging builder UI, start with template-preview.html, template-preview.js, and template-preview.css.",
        "If you are debugging the public website, start with index.html, script.js, styles.css, and projects.json.",
        "If you are debugging installer behavior, start with package.json build.nsis and build/installer.nsh.",
        "If you are debugging releases, start with .github/workflows/build-windows-builder.yml.",
    ])
    doc.add_page_break()


def add_flow_walkthroughs(doc: Document, diagrams: dict[str, Path]) -> None:
    for title, diagram_name, steps, debug_note in FLOW_WALKTHROUGHS:
        doc.add_heading(f"Detailed Walkthrough: {title}", level=1)
        add_diagram(doc, diagrams[diagram_name], f"Context diagram for {title}.")
        doc.add_paragraph("This walkthrough follows the real direction of work through the system. Read it slowly: each line is one handoff from one layer to another.")
        add_table(doc, steps)
        doc.add_heading("Debug note", level=2)
        doc.add_paragraph(debug_note)
        doc.add_page_break()


def add_data_contracts(doc: Document) -> None:
    for title, purpose, example, note in DATA_CONTRACTS:
        doc.add_heading(f"Data Contract: {title}", level=1)
        doc.add_paragraph(purpose)
        doc.add_heading("Example shape", level=2)
        add_code(doc, example)
        doc.add_heading("How to read this", level=2)
        doc.add_paragraph(note)
        doc.add_heading("Why contracts matter", level=2)
        doc.add_paragraph("A data contract is an agreement between two pieces of software. If the frontend sends one shape and the backend expects another shape, the feature breaks even if both files are individually valid code.")
        doc.add_page_break()


def endpoint_purpose(endpoint: str) -> str:
    if "code/compile" in endpoint:
        return "Compiles or runs a source file and returns terminal output."
    if "code/save" in endpoint:
        return "Saves source code into the local compile workspace."
    if "code/beautify" in endpoint:
        return "Formats source code for cleaner display and editing."
    if "code/tools" in endpoint:
        return "Checks compiler/runtime availability."
    if "code/install" in endpoint:
        return "Attempts to install missing compiler tools."
    if "publish-target" in endpoint:
        return "Reads or writes the Git publishing target and authorization state."
    if "publish" in endpoint:
        return "Publishes generated website files to the target repository."
    if "catalog" in endpoint:
        return "Loads project/catalog data for the builder."
    if "templates" in endpoint:
        return "Loads appearance template definitions."
    if "app-update" in endpoint:
        return "Checks or starts builder app update behavior."
    if "security-report" in endpoint:
        return "Returns security/download/auth reporting for the builder."
    if "portfolio-ai" in endpoint:
        return "Handles AI assistant questions."
    if "system-check" in endpoint:
        return "Checks local tool/system readiness."
    return "Project-specific local API endpoint. Inspect server.mjs and the fetch caller for exact behavior."


def extract_api_endpoints() -> list[tuple[str, str, str]]:
    endpoints: dict[tuple[str, str], int] = {}
    for file_name in ["template-preview.js", "script.js", "server.mjs", "index.html"]:
        path = REPO / file_name
        if not path.exists():
            continue
        for line_no, line in enumerate(path.read_text(encoding="utf-8", errors="replace").splitlines(), start=1):
            for match in re.findall(r"['\"](/api/[A-Za-z0-9_./-]+)", line):
                clean = match.rstrip("/")
                endpoints[(clean, file_name)] = min(line_no, endpoints.get((clean, file_name), line_no))
    return [(endpoint, file_name, str(line_no)) for (endpoint, file_name), line_no in sorted(endpoints.items())]


def add_api_endpoint_catalog(doc: Document) -> None:
    endpoints = extract_api_endpoints()
    doc.add_heading("API Endpoint Catalog", level=1)
    doc.add_paragraph("This catalog is generated from strings found in the source files. It tells you which /api paths are used and where to start reading when a request fails.")
    if not endpoints:
        doc.add_paragraph("No /api endpoints were found by the scanner.")
        doc.add_page_break()
        return
    rows = [(endpoint, f"Seen in {file_name} near line {line_no}. Purpose: {endpoint_purpose(endpoint)}") for endpoint, file_name, line_no in endpoints]
    add_table(doc, rows)
    doc.add_page_break()
    for endpoint, file_name, line_no in endpoints:
        doc.add_heading(f"API Detail: {endpoint}", level=1)
        add_table(doc, [
            ("Where seen", f"{file_name} near line {line_no}"),
            ("Purpose", endpoint_purpose(endpoint)),
            ("Frontend question", "Which button, form, or page action sends this request?"),
            ("Backend question", "Which handler in server.mjs receives this path and what files/tools does it touch?"),
            ("Debugging", "Check browser network status, response JSON, server logs, and whether stale cache was avoided with no-store or a timestamp."),
        ])
        doc.add_page_break()


def function_purpose(file_name: str, function_name: str) -> str:
    name = function_name.lower()
    if name.startswith("render"):
        return "Renders UI, markup, or display output."
    if name.startswith("handle"):
        return "Handles an event, request, command, or user action."
    if name.startswith("update"):
        return "Updates UI, state, cache, or stored data."
    if name.startswith("save") or name.startswith("write") or name.startswith("store"):
        return "Persists data to local state, disk, credentials, or browser storage."
    if name.startswith("read") or name.startswith("load") or name.startswith("fetch"):
        return "Loads data from disk, network, browser storage, or a service."
    if "compile" in name or "compiler" in name:
        return "Part of the Compile Code language/tool/build/run flow."
    if "publish" in name or "git" in name or "remote" in name:
        return "Part of publishing, Git, target repository, or authorization behavior."
    if "auth" in name or "credential" in name or "access" in name:
        return "Part of authentication, authorization, or credential checking."
    if "cache" in name:
        return "Reads, writes, or validates cached state."
    if name.startswith("validate") or name.startswith("normalize") or name.startswith("safe") or name.startswith("sanitize"):
        return "Validates or normalizes input so later code receives safe predictable values."
    if name.startswith("create") or name.startswith("new") or name.startswith("build"):
        return "Creates an object, UI structure, process, payload, or generated artifact."
    if name.startswith("open") or name.startswith("close") or name.startswith("show") or name.startswith("hide"):
        return "Controls a window, dialog, panel, menu, or visible state."
    if name.startswith("sync"):
        return "Keeps two places aligned, such as local data and target files."
    if file_name == "main.cjs":
        return "Part of Electron desktop startup, window control, workspace setup, menu behavior, or update handoff."
    if file_name == "server.mjs":
        return "Part of the local backend API, filesystem, compiler, Git, AI, update, or publish workflow."
    if file_name == "template-preview.js":
        return "Part of the private builder frontend and editing experience."
    if file_name == "script.js":
        return "Part of the public website runtime."
    return "Named function in the repository. Inspect the surrounding lines to learn its exact inputs and outputs."


FUNCTION_PATTERNS = [
    re.compile(r"^\s*(?:async\s+)?function\s+([A-Za-z_$][\w$]*)\s*\("),
    re.compile(r"^\s*const\s+([A-Za-z_$][\w$]*)\s*=\s*(?:async\s*)?\([^=]*\)\s*=>"),
    re.compile(r"^\s*const\s+([A-Za-z_$][\w$]*)\s*=\s*(?:async\s+)?function\b"),
    re.compile(r"^\s*export\s+(?:async\s+)?function\s+([A-Za-z_$][\w$]*)\s*\("),
    re.compile(r"^\s*def\s+([A-Za-z_][\w]*)\s*\("),
    re.compile(r"^\s*Function\s+([A-Za-z0-9_.-]+)\s*$"),
]


SOURCE_SIGNAL_RULES = [
    ("Git command", re.compile(r"\bgit\b|simpleGit|GIT_", re.IGNORECASE), "Repository state, publishing, branch checks, or release automation."),
    ("File read", re.compile(r"\breadFile(?:Sync)?\b|Get-Content|createReadStream|fs\.read", re.IGNORECASE), "Reads local builder state, project files, website assets, or configuration."),
    ("File write", re.compile(r"\bwriteFile(?:Sync)?\b|appendFile(?:Sync)?\b|copyFile(?:Sync)?\b|mkdir(?:Sync)?\b|fs\.write", re.IGNORECASE), "Writes drafts, generated portfolio output, compiler artifacts, uploaded files, or installer data."),
    ("Process execution", re.compile(r"\bspawn\b|\bexecFile\b|\bexec\b|\bStart-Process\b", re.IGNORECASE), "Runs Git, compilers, installers, or helper tools outside the JavaScript runtime."),
    ("Compiler or simulator", re.compile(r"\bgcc\b|\bg\+\+\b|\bjavac\b|\bjava\b|\bpython\b|\bnode\b|\biverilog\b|\bvvp\b|\bverilog\b|\bsystemverilog\b", re.IGNORECASE), "Part of the Compile Code workspace or language/tool detection."),
    ("Cloudflare or AI", re.compile(r"cloudflare|wrangler|OPENAI|Workers AI|portfolio-ai|chat|model", re.IGNORECASE), "Public AI assistant, Worker deployment, or model-provider fallback behavior."),
    ("Security or auth", re.compile(r"auth|token|credential|secret|oauth|permission|authorize|scope|csrf|csp|headers", re.IGNORECASE), "Publishing protection, API key safety, HTTP security, or user authorization."),
    ("Browser DOM", re.compile(r"document\.|querySelector|getElementById|classList|dataset|addEventListener", re.IGNORECASE), "Builder or website interface behavior in the browser/Electron renderer."),
    ("Local storage", re.compile(r"localStorage|sessionStorage|IndexedDB|cache", re.IGNORECASE), "Local persistence, draft state, auth cache, update snooze, or browser-side caching."),
    ("Network request", re.compile(r"\bfetch\s*\(|XMLHttpRequest|https?://|/api/", re.IGNORECASE), "Frontend-backend communication or public web/API calls."),
    ("Rich text", re.compile(r"contenteditable|execCommand|Selection|Range|clipboard|paste|font|color|sanitize", re.IGNORECASE), "Rich editor behavior, formatting, paste handling, or content sanitization."),
    ("Installer", re.compile(r"NSIS|installer|uninstall|shortcut|AppData|Program Files|UAC|electron-builder", re.IGNORECASE), "Windows setup, update, shortcut, or installation behavior."),
]


COMMON_CALL_WORDS = {
    "if", "for", "while", "switch", "catch", "return", "typeof", "await", "function", "class", "new", "throw",
    "console", "Math", "JSON", "String", "Number", "Boolean", "Array", "Object", "Promise", "Date", "RegExp",
    "setTimeout", "clearTimeout", "setInterval", "parseInt", "parseFloat", "encodeURIComponent",
    "decodeURIComponent", "require", "import", "describe", "it", "test",
}


def line_excerpt(lines: list[str], start: int, end: int, max_lines: int = 120) -> str:
    selected = lines[max(start - 1, 0): min(end, len(lines))]
    truncated = len(selected) > max_lines
    selected = selected[:max_lines]
    rendered = [f"{start + index:>5}: {line}" for index, line in enumerate(selected)]
    if truncated:
        rendered.append(f"      ... excerpt truncated after {max_lines} lines. Open the source file for the rest of this function.")
    return "\n".join(rendered)


def find_block_end(lines: list[str], start_index: int) -> int:
    brace_balance = 0
    saw_brace = False
    for index in range(start_index, min(start_index + 500, len(lines))):
        line = re.sub(r"//.*$", "", lines[index])
        brace_balance += line.count("{")
        brace_balance -= line.count("}")
        if "{" in line:
            saw_brace = True
        if saw_brace and brace_balance <= 0 and index > start_index:
            return index + 1
        if not saw_brace and index > start_index and not lines[index].startswith((" ", "\t")):
            return index
    return min(start_index + 80, len(lines))


def extract_function_blocks(repo_file: str, lines: list[str]) -> list[dict]:
    functions = []
    for index, line in enumerate(lines):
        stripped = line.strip()
        for pattern in FUNCTION_PATTERNS:
            match = pattern.search(line)
            if not match:
                continue
            name = match.group(1)
            if line.lstrip().startswith("Function "):
                end = next((i + 1 for i in range(index + 1, len(lines)) if lines[i].lstrip().startswith("FunctionEnd")), min(index + 80, len(lines)))
            elif line.lstrip().startswith("def "):
                end = next((i for i in range(index + 1, len(lines)) if lines[i] and not lines[i].startswith((" ", "\t", "@"))), min(index + 120, len(lines)))
            else:
                end = find_block_end(lines, index)
            body = "\n".join(lines[index:end])
            call_candidates = re.findall(r"\b([A-Za-z_$][\w$]*)\s*\(", body)
            calls = []
            for candidate in call_candidates:
                if candidate not in COMMON_CALL_WORDS and candidate != name and candidate not in calls:
                    calls.append(candidate)
            endpoint_hits = sorted(set(match.rstrip("/") for match in re.findall(r"['\"](/api/[A-Za-z0-9_./-]+)", body)))
            local_storage_keys = sorted(set(match for match in re.findall(r"(?:localStorage|sessionStorage)\.(?:getItem|setItem|removeItem)\(['\"]([^'\"]+)", body)))
            functions.append({
                "line": index + 1,
                "end_line": end,
                "name": name,
                "purpose": function_purpose(Path(repo_file).name, name),
                "signature": stripped[:240],
                "excerpt": line_excerpt(lines, index + 1, end, max_lines=120),
                "line_count": max(end - index, 1),
                "calls": calls[:30],
                "endpoints": endpoint_hits[:30],
                "storage_keys": local_storage_keys[:30],
                "returns": len(re.findall(r"\breturn\b", body)),
                "awaits": len(re.findall(r"\bawait\b", body)),
                "touches_dom": bool(re.search(r"document\.|querySelector|getElementById|classList|dataset", body)),
                "touches_files": bool(re.search(r"readFile|writeFile|appendFile|copyFile|mkdir|rm|unlink|fs\.", body, re.IGNORECASE)),
                "runs_process": bool(re.search(r"\bspawn\b|\bexecFile\b|\bexec\b", body, re.IGNORECASE)),
                "uses_network": bool(re.search(r"\bfetch\s*\(|https?://|/api/", body, re.IGNORECASE)),
                "uses_security": bool(re.search(r"auth|token|secret|credential|permission|authorize|scope", body, re.IGNORECASE)),
            })
            break
    return functions


def collect_signal_hits(lines: list[str]) -> list[dict]:
    hits = []
    for line_no, line in enumerate(lines, start=1):
        stripped = line.strip()
        if not stripped:
            continue
        for name, pattern, explanation in SOURCE_SIGNAL_RULES:
            if pattern.search(line):
                hits.append({
                    "line": line_no,
                    "type": name,
                    "evidence": stripped[:220],
                    "meaning": explanation,
                })
                break
    return hits


def function_detail_rows(item: dict) -> list[tuple[str, str]]:
    touches = []
    if item["touches_dom"]:
        touches.append("browser/Electron DOM")
    if item["touches_files"]:
        touches.append("filesystem")
    if item["runs_process"]:
        touches.append("external processes")
    if item["uses_network"]:
        touches.append("network/API requests")
    if item["uses_security"]:
        touches.append("authentication/security data")
    return [
        ("Source location", f"Lines {item['line']}-{item['end_line']} ({item['line_count']} source lines)."),
        ("Purpose", item["purpose"]),
        ("Declaration", item["signature"]),
        ("Touches", ", ".join(touches) if touches else "Mostly local calculations or local UI state."),
        ("Calls detected", ", ".join(item["calls"][:12]) if item["calls"] else "No obvious named calls detected by the generator."),
        ("API endpoints", ", ".join(item["endpoints"]) if item["endpoints"] else "None detected inside this function body."),
        ("Storage keys", ", ".join(item["storage_keys"]) if item["storage_keys"] else "No local/session storage keys detected."),
        ("Async/return clues", f"{item['awaits']} await expression(s), {item['returns']} return statement(s)."),
        ("Debug approach", "Trigger the user action that reaches this function, inspect inputs/state before the function, then inspect the returned value, DOM update, file write, process output, or API response after it runs."),
    ]


def extract_functions() -> list[tuple[str, int, str, str]]:
    function_rows: list[tuple[str, int, str, str]] = []
    for file_name in ["main.cjs", "server.mjs", "template-preview.js", "script.js", "cloudflare/portfolio-ai-worker.js", "docs/build_complete_guide.py", "build/installer.nsh"]:
        path = REPO / file_name
        if not path.exists():
            continue
        for line_no, line in enumerate(path.read_text(encoding="utf-8", errors="replace").splitlines(), start=1):
            for pattern in FUNCTION_PATTERNS:
                match = pattern.search(line)
                if match:
                    name = match.group(1)
                    function_rows.append((file_name, line_no, name, function_purpose(file_name, name)))
                    break
    return function_rows


TEXT_CODE_EXTENSIONS = {
    ".cjs", ".css", ".html", ".js", ".json", ".md", ".mjs", ".nsh", ".py", ".toml", ".txt", ".yaml", ".yml"
}


def is_text_code_file(repo_file: str) -> bool:
    path = REPO / repo_file
    if is_generated_reference_file(repo_file):
        return False
    if repo_file.endswith((".docx", ".pdf", ".png", ".ico")):
        return False
    return path.is_file() and path.suffix.lower() in TEXT_CODE_EXTENSIONS


def source_lines(repo_file: str) -> list[str]:
    path = REPO / repo_file
    if not path.exists() or not path.is_file():
        return []
    return path.read_text(encoding="utf-8", errors="replace").splitlines()


def extract_source_facts(repo_file: str) -> dict:
    lines = source_lines(repo_file)
    text = "\n".join(lines)
    functions = extract_function_blocks(repo_file, lines)
    variables = []
    imports = []
    endpoints = []
    selectors = []
    events = []
    routes = []
    json_keys = []
    for line_no, line in enumerate(lines, start=1):
        stripped = line.strip()
        var_match = re.match(r"^\s*(?:const|let|var)\s+([A-Za-z_$][\w$]*)\s*=", line)
        if var_match and len(variables) < 140:
            variables.append({
                "line": line_no,
                "name": var_match.group(1),
                "statement": stripped[:180],
            })
        if re.match(r"^\s*import\s+", line) or re.match(r"^\s*(?:const|let|var)\s+\S+\s*=\s*require\(", line):
            imports.append({"line": line_no, "statement": stripped[:220]})
        for match in re.findall(r"['\"](/api/[A-Za-z0-9_./-]+)", line):
            endpoints.append({"line": line_no, "endpoint": match.rstrip("/")})
        route_match = re.search(r"\bapp\.(get|post|put|delete|patch)\(['\"]([^'\"]+)", line)
        if route_match:
            routes.append({"line": line_no, "method": route_match.group(1).upper(), "path": route_match.group(2)})
        if ".addEventListener(" in line:
            events.append({"line": line_no, "statement": stripped[:180]})
        json_match = re.match(r'^\s*"([^"]+)"\s*:', line)
        if json_match and len(json_keys) < 120:
            json_keys.append({"line": line_no, "key": json_match.group(1), "statement": stripped[:180]})
        if repo_file.endswith(".css"):
            selector_match = re.match(r"^\s*([.#A-Za-z0-9_\-:\[\]\(\)\s>,+.=\"']+)\s*\{", line)
            if selector_match and len(selectors) < 140:
                selectors.append({"line": line_no, "selector": selector_match.group(1).strip()[:160]})
    return {
        "repo_file": repo_file,
        "line_count": len(lines),
        "size_bytes": (REPO / repo_file).stat().st_size if (REPO / repo_file).exists() else 0,
        "description": FILE_DESCRIPTIONS.get(repo_file, "Tracked source/configuration file in the OMB Portfolio Builder repository."),
        "functions": functions,
        "variables": variables,
        "imports": imports,
        "endpoints": endpoints,
        "routes": routes,
        "selectors": selectors,
        "events": events,
        "json_keys": json_keys,
        "signals": collect_signal_hits(lines),
        "mentions_server": "server.mjs" in text,
        "mentions_template_preview": "template-preview" in text,
        "mentions_script": "script.js" in text,
        "mentions_projects_json": "projects.json" in text,
        "mentions_github": "github" in text.lower(),
        "mentions_cloudflare": "cloudflare" in text.lower() or "worker" in text.lower(),
        "snippet": line_excerpt(lines, 1, min(72, len(lines)), max_lines=72) if lines else "",
    }


def fact_relationship_summary(facts: dict) -> str:
    related = []
    if facts.get("mentions_template_preview"):
        related.append("builder frontend")
    if facts.get("mentions_server"):
        related.append("local backend")
    if facts.get("mentions_script"):
        related.append("public website runtime")
    if facts.get("mentions_projects_json"):
        related.append("portfolio catalog")
    if facts.get("mentions_cloudflare"):
        related.append("Cloudflare/AI layer")
    if facts.get("mentions_github"):
        related.append("GitHub/release layer")
    return ", ".join(related) if related else "Mostly local to its own feature area."


def folder_role(repo_file: str) -> str:
    if repo_file.startswith(".github/workflows/"):
        return "GitHub automation. GitHub Actions reads this during CI/release/branch-gate runs."
    if repo_file.startswith(".well-known/"):
        return "Public web metadata. Browsers, security tools, and crawlers may request this path."
    if repo_file.startswith("assets/"):
        return "Public and desktop visual asset. Used for favicons, app icons, branding, or install surfaces."
    if repo_file.startswith("build/"):
        return "Packaging and installer customization. Used while creating the Windows app installer."
    if repo_file.startswith("cloudflare/"):
        return "Cloudflare Worker/deployment area. Used for public AI backend and Cloudflare deployment notes."
    if repo_file.startswith("docs/"):
        return "Documentation and generated reference area. Used to explain the app and source structure."
    if repo_file.startswith("Backgrounds/"):
        return "Portfolio background asset folder placeholder."
    return "Repository root. Usually an entry file, app config, public website file, package manifest, or top-level documentation."


def file_category(repo_file: str, path: Path) -> str:
    suffix = path.suffix.lower()
    if repo_file.endswith(".gitkeep"):
        return "Folder marker"
    if repo_file.startswith(".github/workflows/"):
        return "GitHub workflow"
    if suffix in {".png", ".ico", ".jpg", ".jpeg", ".webp", ".svg"}:
        return "Image/icon asset"
    if suffix in {".docx", ".pdf"}:
        return "Generated document"
    if suffix in {".js", ".mjs", ".cjs"}:
        return "JavaScript source"
    if suffix in {".html", ".css"}:
        return "Website/builder UI source"
    if suffix in {".json", ".toml", ".yaml", ".yml"}:
        return "Configuration or structured data"
    if suffix in {".md", ".txt", ""}:
        return "Documentation or text control file"
    if suffix == ".nsh":
        return "NSIS installer script"
    if suffix == ".py":
        return "Python tooling"
    return "Repository file"


def image_metadata(path: Path) -> dict:
    if not path.exists() or path.suffix.lower() not in {".png", ".jpg", ".jpeg", ".webp", ".ico"}:
        return {}
    try:
        with Image.open(path) as image:
            return {
                "format": image.format or path.suffix.lower().lstrip(".").upper(),
                "width": image.size[0],
                "height": image.size[1],
                "mode": image.mode,
                "frames": getattr(image, "n_frames", 1),
            }
    except Exception as exc:
        return {"error": str(exc)}


def file_specific_notes(repo_file: str, category: str) -> list[str]:
    notes = []
    if repo_file in FILE_DESCRIPTIONS:
        notes.append(FILE_DESCRIPTIONS[repo_file])
    if repo_file.startswith("assets/favicon"):
        notes.append("This favicon participates in browser tab identity, mobile install surfaces, and public website branding.")
    if repo_file.startswith("assets/omb-app-icon"):
        notes.append("This app icon is used by the Windows desktop application, installer metadata, and shortcuts.")
    if repo_file == "assets/omb-mark.png":
        notes.append("This is the OMB visual mark used by the website and builder branding.")
    if repo_file == "_headers":
        notes.append("This file describes security and caching headers for hosts that support static `_headers` files, such as Cloudflare Pages style deployments.")
    if repo_file == ".nojekyll":
        notes.append("This makes GitHub Pages serve static files directly without Jekyll filtering.")
    if repo_file == ".gitattributes":
        notes.append("This protects binary artifacts such as DOCX, PDF, PNG, ICO, and installers from line-ending conversion.")
    if repo_file == ".gitignore":
        notes.append("This keeps temporary build output, local-only data, and generated noise out of version control.")
    if repo_file.endswith(".gitkeep"):
        notes.append("This empty marker exists so Git keeps an otherwise empty folder.")
    if category == "Image/icon asset":
        notes.append("Do not edit binary assets with a text editor. Replace them with properly exported image/icon files.")
    if category == "GitHub workflow":
        notes.append("Workflow changes should be tested on a feature branch because mistakes can block merges or releases.")
    if not notes:
        notes.append("Tracked repository file used by the builder, website, installer, documentation, or release process.")
    return notes


def extract_file_facts(repo_file: str) -> dict:
    path = REPO / repo_file
    suffix = path.suffix.lower()
    category = file_category(repo_file, path)
    text_file = is_text_code_file(repo_file) or suffix in {".gitignore", ".gitattributes", ".nojekyll", "", ".txt", ".md"}
    lines = source_lines(repo_file) if text_file and path.exists() and path.is_file() else []
    mime_type, _encoding = mimetypes.guess_type(str(path))
    facts = {
        "repo_file": repo_file,
        "path": path,
        "exists": path.exists(),
        "size_bytes": path.stat().st_size if path.exists() and path.is_file() else 0,
        "suffix": suffix or "(none)",
        "category": category,
        "mime_type": mime_type or "unknown",
        "folder_role": folder_role(repo_file),
        "description": FILE_DESCRIPTIONS.get(repo_file, "Tracked repository file used by the builder or public website."),
        "notes": file_specific_notes(repo_file, category),
        "is_text": bool(lines or text_file),
        "line_count": len(lines),
        "snippet": line_excerpt(lines, 1, min(72, len(lines)), max_lines=72) if lines else "",
        "image": image_metadata(path),
    }
    if is_text_code_file(repo_file):
        facts["source"] = extract_source_facts(repo_file)
    else:
        facts["source"] = None
    return facts


def add_file_fact_summary(doc: Document, facts: dict) -> None:
    rows = [
        ("File", facts["repo_file"]),
        ("Category", facts["category"]),
        ("Folder role", facts["folder_role"]),
        ("Size", f"{facts['size_bytes']:,} bytes"),
        ("Extension", facts["suffix"]),
        ("MIME type", facts["mime_type"]),
        ("Text lines", f"{facts['line_count']:,}" if facts["is_text"] else "Not a readable text file"),
    ]
    if facts["image"]:
        image = facts["image"]
        if "error" in image:
            rows.append(("Image metadata", f"Could not inspect image: {image['error']}"))
        else:
            rows.append(("Image metadata", f"{image['format']} {image['width']}x{image['height']}, mode {image['mode']}, frames {image['frames']}"))
    add_table(doc, rows)


def add_file_fact_sections(doc: Document, facts: dict) -> None:
    doc.add_heading("Why this file exists", level=2)
    for note in facts["notes"]:
        doc.add_paragraph(note, style="List Bullet")
    doc.add_heading("How it is used", level=2)
    doc.add_paragraph(facts["folder_role"])
    if facts["source"]:
        add_source_fact_sections(doc, facts["source"], include_excerpts=True)
    elif facts["snippet"]:
        doc.add_heading("Readable content snippet", level=2)
        add_code(doc, facts["snippet"])
    elif facts["image"]:
        doc.add_heading("Asset handling note", level=2)
        doc.add_paragraph("This file is documented by metadata instead of raw bytes. Binary image/icon bytes are not useful to print in a Word/PDF reference. Use an image editor or icon tool when changing it.")
    else:
        doc.add_heading("Binary or marker note", level=2)
        doc.add_paragraph("This file is either binary, empty, or a marker/config file whose value comes from its presence and path rather than readable content.")
    doc.add_heading("Safe change checklist", level=2)
    numbers(doc, [
        "Confirm which app, website, installer, workflow, or document consumes this file.",
        "Change it on a feature branch from development.",
        "Regenerate docs if the file purpose, API surface, or behavior changes.",
        "Run the smallest relevant smoke test before merging into development.",
    ])


def code_reference_name(repo_file: str, suffix: str) -> str:
    return re.sub(r"[^A-Za-z0-9_.-]+", "__", repo_file).strip("_") + suffix


def remove_directory(path: Path) -> None:
    if not path.exists():
        return

    def remove_readonly(function, item_path, excinfo):
        try:
            Path(item_path).chmod(0o700)
            function(item_path)
        except Exception:
            raise excinfo

    shutil.rmtree(path, onexc=remove_readonly)


def add_source_fact_summary(doc: Document, facts: dict) -> None:
    add_table(doc, [
        ("File", facts["repo_file"]),
        ("Purpose", facts["description"]),
        ("Lines", f"{facts['line_count']:,}"),
        ("Size", f"{facts['size_bytes']:,} bytes"),
        ("Talks to", fact_relationship_summary(facts)),
        ("Functions", str(len(facts["functions"]))),
        ("Variables/constants", str(len(facts["variables"]))),
        ("API mentions", str(len(facts["endpoints"]))),
        ("Signals", str(len(facts["signals"]))),
    ])


def add_source_fact_sections(doc: Document, facts: dict, include_excerpts: bool = True) -> None:
    doc.add_heading("How this file participates in the system", level=2)
    doc.add_paragraph("This generated section reads the actual file and points out how it communicates with other pieces of the app. It is intentionally concrete: line numbers, declarations, endpoints, events, source excerpts, and debugging cues.")
    if facts["imports"]:
        doc.add_heading("Imports, requires, and external dependencies", level=2)
        add_table(doc, [(f"Line {item['line']}", item["statement"]) for item in facts["imports"][:60]])
    if facts["routes"]:
        doc.add_heading("Backend routes implemented here", level=2)
        add_table(doc, [(f"{item['method']} {item['path']}", f"Line {item['line']}. {endpoint_purpose(item['path'])}") for item in facts["routes"][:80]])
    if facts["endpoints"]:
        doc.add_heading("API endpoints mentioned or called", level=2)
        add_table(doc, [(item["endpoint"], f"Line {item['line']}. {endpoint_purpose(item['endpoint'])}") for item in facts["endpoints"][:80]])
    if facts["signals"]:
        doc.add_heading("Important implementation signals", level=2)
        add_table(doc, [(f"{item['type']} at line {item['line']}", f"{item['meaning']} Evidence: {item['evidence']}") for item in facts["signals"][:80]])
    if facts["variables"]:
        doc.add_heading("Important constants and variables", level=2)
        add_table(doc, [(item["name"], f"Line {item['line']}: {item['statement']}") for item in facts["variables"][:80]])
    if facts["json_keys"]:
        doc.add_heading("JSON/YAML-style keys detected", level=2)
        add_table(doc, [(item["key"], f"Line {item['line']}: {item['statement']}") for item in facts["json_keys"][:80]])
    if facts["events"]:
        doc.add_heading("Event handlers and user interactions", level=2)
        add_table(doc, [(f"Line {item['line']}", item["statement"]) for item in facts["events"][:80]])
    if facts["selectors"]:
        doc.add_heading("CSS selectors and visual hooks", level=2)
        add_table(doc, [(f"Line {item['line']}", item["selector"]) for item in facts["selectors"][:120]])
    if facts["functions"]:
        doc.add_heading("Function-by-function detail", level=2)
        for item in facts["functions"]:
            doc.add_heading(f"{item['name']} - lines {item['line']}-{item['end_line']}", level=3)
            add_table(doc, function_detail_rows(item))
            if include_excerpts:
                doc.add_paragraph("Source excerpt:")
                add_code(doc, item["excerpt"])
    doc.add_heading("Representative opening snippet", level=2)
    add_code(doc, facts["snippet"] or "(empty file)")
    doc.add_heading("Beginner debugging checklist for this file", level=2)
    numbers(doc, [
        "Name the user action, command, or workflow that reaches this file.",
        "Find the exact function or route that owns the behavior.",
        "Check the inputs: DOM state, request body, file path, local storage value, compiler source, or Git branch.",
        "Check the outputs: returned JSON, updated DOM, written file, terminal output, generated project catalog, or published website asset.",
        "If the issue crosses a boundary, test the boundary directly: browser console for frontend, endpoint call for backend, command line for Git/compiler, Cloudflare logs for Worker behavior.",
    ])


def setup_code_reference_document(title: str, subtitle: str) -> Document:
    doc = setup_document()
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title_p.add_run(title)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(11, 37, 69)
    doc.add_paragraph(subtitle)
    doc.add_paragraph(f"Generated from repository state at {REPO}.")
    doc.add_paragraph("")
    return doc


def compact_generated_docx(doc: Document, keep_first_page_break: bool = True) -> int:
    """Remove generator-inserted page breaks so pages fill naturally."""
    kept = 0
    removed = 0
    for paragraph in list(doc.paragraphs):
        has_page_break = any(
            element.tag == qn("w:br") and element.get(qn("w:type")) == "page"
            for element in paragraph._element.iter()
        )
        if not has_page_break:
            continue
        if keep_first_page_break and kept == 0:
            kept += 1
            continue
        parent = paragraph._element.getparent()
        if parent is not None:
            parent.remove(paragraph._element)
            removed += 1
    return removed


def write_single_code_reference_docx(facts: dict, output_path: Path) -> None:
    doc = setup_code_reference_document(
        f"Code Reference: {facts['repo_file']}",
        "A source-level explanation with function details, file communication, variables, endpoints, and debugging notes.",
    )
    add_source_fact_summary(doc, facts)
    add_source_fact_sections(doc, facts, include_excerpts=True)
    compact_generated_docx(doc, keep_first_page_break=False)
    doc.save(output_path)


def pdf_styles():
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(
        name="CodeTiny",
        parent=styles["Code"],
        fontName="Courier",
        fontSize=6.4,
        leading=7.6,
        backColor=colors.HexColor("#F5F8FB"),
        borderColor=colors.HexColor("#D8E6F0"),
        borderWidth=0.25,
        borderPadding=4,
        spaceAfter=7,
    ))
    styles.add(ParagraphStyle(
        name="SmallBody",
        parent=styles["BodyText"],
        fontSize=8.5,
        leading=11.0,
        spaceAfter=5,
    ))
    styles.add(ParagraphStyle(
        name="TinyTable",
        parent=styles["BodyText"],
        fontSize=7.0,
        leading=8.5,
    ))
    return styles


def pdf_paragraph(text: str, style) -> Paragraph:
    return Paragraph(html.escape(str(text)).replace("\n", "<br/>"), style)


def pdf_code(text: str, styles) -> Preformatted:
    wrapped_lines = []
    for raw_line in (text or "").splitlines():
        expanded = raw_line.replace("\t", "    ")
        chunks = textwrap.wrap(expanded, width=112, replace_whitespace=False, drop_whitespace=False) or [""]
        wrapped_lines.extend(chunks)
    return Preformatted("\n".join(wrapped_lines), styles["CodeTiny"])


def pdf_table(rows: list[tuple[str, str]], styles):
    data = [[pdf_paragraph("Item", styles["TinyTable"]), pdf_paragraph("Explanation", styles["TinyTable"])]]
    data.extend([[pdf_paragraph(label, styles["TinyTable"]), pdf_paragraph(value, styles["TinyTable"])] for label, value in rows])
    table = Table(data, colWidths=[1.45 * inch, 5.25 * inch], repeatRows=1)
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E8EEF5")),
        ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#B7CFE0")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 5),
        ("RIGHTPADDING", (0, 0), (-1, -1), 5),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    return table


def add_pdf_source_fact_sections(story: list, facts: dict, styles, include_excerpts: bool = True) -> None:
    story.append(Paragraph("How this file participates in the system", styles["Heading2"]))
    story.append(pdf_paragraph("This generated section reads the actual file and points out line numbers, declarations, endpoints, events, source excerpts, and debugging cues.", styles["SmallBody"]))
    if facts["imports"]:
        story.append(Paragraph("Imports, requires, and external dependencies", styles["Heading2"]))
        story.append(pdf_table([(f"Line {item['line']}", item["statement"]) for item in facts["imports"][:60]], styles))
    if facts["routes"]:
        story.append(Paragraph("Backend routes implemented here", styles["Heading2"]))
        story.append(pdf_table([(f"{item['method']} {item['path']}", f"Line {item['line']}. {endpoint_purpose(item['path'])}") for item in facts["routes"][:80]], styles))
    if facts["endpoints"]:
        story.append(Paragraph("API endpoints mentioned or called", styles["Heading2"]))
        story.append(pdf_table([(item["endpoint"], f"Line {item['line']}. {endpoint_purpose(item['endpoint'])}") for item in facts["endpoints"][:80]], styles))
    if facts["signals"]:
        story.append(Paragraph("Important implementation signals", styles["Heading2"]))
        story.append(pdf_table([(f"{item['type']} at line {item['line']}", f"{item['meaning']} Evidence: {item['evidence']}") for item in facts["signals"][:80]], styles))
    if facts["variables"]:
        story.append(Paragraph("Important constants and variables", styles["Heading2"]))
        story.append(pdf_table([(item["name"], f"Line {item['line']}: {item['statement']}") for item in facts["variables"][:80]], styles))
    if facts["json_keys"]:
        story.append(Paragraph("JSON/YAML-style keys detected", styles["Heading2"]))
        story.append(pdf_table([(item["key"], f"Line {item['line']}: {item['statement']}") for item in facts["json_keys"][:80]], styles))
    if facts["events"]:
        story.append(Paragraph("Event handlers and user interactions", styles["Heading2"]))
        story.append(pdf_table([(f"Line {item['line']}", item["statement"]) for item in facts["events"][:80]], styles))
    if facts["selectors"]:
        story.append(Paragraph("CSS selectors and visual hooks", styles["Heading2"]))
        story.append(pdf_table([(f"Line {item['line']}", item["selector"]) for item in facts["selectors"][:120]], styles))
    if facts["functions"]:
        story.append(Paragraph("Function-by-function detail", styles["Heading2"]))
        for item in facts["functions"]:
            story.append(Paragraph(f"{item['name']} - lines {item['line']}-{item['end_line']}", styles["Heading3"]))
            story.append(pdf_table(function_detail_rows(item), styles))
            if include_excerpts:
                story.append(pdf_paragraph("Source excerpt:", styles["SmallBody"]))
                story.append(pdf_code(item["excerpt"], styles))
    story.append(Paragraph("Representative opening snippet", styles["Heading2"]))
    story.append(pdf_code(facts["snippet"] or "(empty file)", styles))


def write_single_code_reference_pdf(facts: dict, output_path: Path) -> None:
    if not REPORTLAB_AVAILABLE:
        return
    styles = pdf_styles()
    story = [
        Paragraph(f"Code Reference: {facts['repo_file']}", styles["Title"]),
        pdf_paragraph("A source-level explanation with function details, file communication, variables, endpoints, and debugging notes.", styles["SmallBody"]),
        pdf_table([
            ("File", facts["repo_file"]),
            ("Purpose", facts["description"]),
            ("Lines", f"{facts['line_count']:,}"),
            ("Size", f"{facts['size_bytes']:,} bytes"),
            ("Talks to", fact_relationship_summary(facts)),
            ("Functions", str(len(facts["functions"]))),
            ("Variables/constants", str(len(facts["variables"]))),
            ("API mentions", str(len(facts["endpoints"]))),
            ("Signals", str(len(facts["signals"]))),
        ], styles),
        Spacer(1, 0.12 * inch),
    ]
    add_pdf_source_fact_sections(story, facts, styles, include_excerpts=True)
    SimpleDocTemplate(str(output_path), pagesize=LETTER, rightMargin=0.55 * inch, leftMargin=0.55 * inch, topMargin=0.55 * inch, bottomMargin=0.55 * inch).build(story)


def write_master_code_reference_docx(code_files: list[str], facts_by_file: dict[str, dict], diagrams: dict[str, Path]) -> None:
    doc = setup_code_reference_document(
        "OMB Portfolio Builder Code Reference",
        "A generated Word reference that explains the important source files, all discovered functions, communication paths, variables, endpoints, installers, AI paths, and compiler paths.",
    )
    add_diagram(doc, diagrams["file-communication-map"], "Main source file communication map.")
    add_diagram(doc, diagrams["frontend-backend-cloudflare"], "Frontend, backend, Cloudflare, and AI boundary.")
    add_diagram(doc, diagrams["compile-code-flow"], "Compile Code workspace and simulator flow.")
    doc.add_heading("How to use this document", level=1)
    numbers(doc, [
        "Start with the file you want to understand.",
        "Read the summary table to understand its role.",
        "Read implementation signals to see whether the file touches Git, files, compilers, Cloudflare, authentication, DOM, storage, or network calls.",
        "Read function details to understand line ranges, side effects, calls, endpoints, and source excerpts.",
        "Use the debugging checklist to trace a behavior from click to function to file/API/output.",
    ])
    doc.add_heading("Generated file index", level=1)
    add_table(doc, [(repo_file, f"{facts_by_file[repo_file]['line_count']:,} lines, {len(facts_by_file[repo_file]['functions'])} function(s), {len(facts_by_file[repo_file]['signals'])} signal(s).") for repo_file in code_files])
    doc.add_page_break()
    for repo_file in code_files:
        facts = facts_by_file[repo_file]
        doc.add_heading(f"Source File: {repo_file}", level=1)
        add_source_fact_summary(doc, facts)
        add_source_fact_sections(doc, facts, include_excerpts=True)
        doc.add_page_break()
    compact_generated_docx(doc, keep_first_page_break=True)
    doc.save(CODE_REFERENCE_MASTER_DOCX)


def write_master_code_reference_pdf(code_files: list[str], facts_by_file: dict[str, dict], diagrams: dict[str, Path]) -> None:
    if not REPORTLAB_AVAILABLE:
        return
    styles = pdf_styles()
    story = [
        Paragraph("OMB Portfolio Builder Code Reference", styles["Title"]),
        pdf_paragraph("A generated PDF reference that explains the important source files, discovered functions, communication paths, variables, endpoints, installers, AI paths, and compiler paths.", styles["SmallBody"]),
    ]
    for diagram_name, caption in [
        ("file-communication-map", "Main source file communication map."),
        ("frontend-backend-cloudflare", "Frontend, backend, Cloudflare, and AI boundary."),
        ("compile-code-flow", "Compile Code workspace and simulator flow."),
    ]:
        path = diagrams.get(diagram_name)
        if path and path.exists():
            story.append(PdfImage(str(path), width=6.6 * inch, height=3.72 * inch))
            story.append(pdf_paragraph(caption, styles["SmallBody"]))
    story.append(Paragraph("Generated file index", styles["Heading1"]))
    story.append(pdf_table([(repo_file, f"{facts_by_file[repo_file]['line_count']:,} lines, {len(facts_by_file[repo_file]['functions'])} function(s), {len(facts_by_file[repo_file]['signals'])} signal(s).") for repo_file in code_files], styles))
    story.append(PageBreak())
    for repo_file in code_files:
        facts = facts_by_file[repo_file]
        story.append(Paragraph(f"Source File: {repo_file}", styles["Heading1"]))
        story.append(pdf_table([
            ("File", facts["repo_file"]),
            ("Purpose", facts["description"]),
            ("Lines", f"{facts['line_count']:,}"),
            ("Size", f"{facts['size_bytes']:,} bytes"),
            ("Talks to", fact_relationship_summary(facts)),
            ("Functions", str(len(facts["functions"]))),
            ("Variables/constants", str(len(facts["variables"]))),
            ("API mentions", str(len(facts["endpoints"]))),
            ("Signals", str(len(facts["signals"]))),
        ], styles))
        add_pdf_source_fact_sections(story, facts, styles, include_excerpts=True)
        story.append(PageBreak())
    SimpleDocTemplate(str(CODE_REFERENCE_MASTER_PDF), pagesize=LETTER, rightMargin=0.55 * inch, leftMargin=0.55 * inch, topMargin=0.55 * inch, bottomMargin=0.55 * inch).build(story)


def write_single_file_reference_docx(facts: dict, output_path: Path) -> None:
    doc = setup_code_reference_document(
        f"Important Code Reference: {facts['repo_file']}",
        "A focused code/control-file explanation covering purpose, owner layer, APIs, variables, implementation signals, source excerpts, and debugging rules.",
    )
    add_file_fact_summary(doc, facts)
    add_file_fact_sections(doc, facts)
    compact_generated_docx(doc, keep_first_page_break=False)
    doc.save(output_path)


def add_pdf_file_fact_sections(story: list, facts: dict, styles) -> None:
    story.append(Paragraph("Why this file exists", styles["Heading2"]))
    for note in facts["notes"]:
        story.append(pdf_paragraph(f"- {note}", styles["SmallBody"]))
    story.append(Paragraph("How it is used", styles["Heading2"]))
    story.append(pdf_paragraph(facts["folder_role"], styles["SmallBody"]))
    if facts["source"]:
        add_pdf_source_fact_sections(story, facts["source"], styles, include_excerpts=True)
    elif facts["snippet"]:
        story.append(Paragraph("Readable content snippet", styles["Heading2"]))
        story.append(pdf_code(facts["snippet"], styles))
    elif facts["image"]:
        story.append(Paragraph("Asset handling note", styles["Heading2"]))
        story.append(pdf_paragraph("This file is documented by metadata instead of raw bytes. Binary image/icon bytes are not useful to print in a reference document.", styles["SmallBody"]))
    else:
        story.append(Paragraph("Binary or marker note", styles["Heading2"]))
        story.append(pdf_paragraph("This file is either binary, empty, or a marker/config file whose value comes from its presence and path rather than readable content.", styles["SmallBody"]))
    story.append(Paragraph("Safe change checklist", styles["Heading2"]))
    for item in [
        "Confirm which app, website, installer, workflow, or document consumes this file.",
        "Change it on a feature branch from development.",
        "Regenerate docs if the file purpose, API surface, or behavior changes.",
        "Run the smallest relevant smoke test before merging into development.",
    ]:
        story.append(pdf_paragraph(f"- {item}", styles["SmallBody"]))


def write_single_file_reference_pdf(facts: dict, output_path: Path) -> None:
    if not REPORTLAB_AVAILABLE:
        return
    styles = pdf_styles()
    rows = [
        ("File", facts["repo_file"]),
        ("Category", facts["category"]),
        ("Folder role", facts["folder_role"]),
        ("Size", f"{facts['size_bytes']:,} bytes"),
        ("Extension", facts["suffix"]),
        ("MIME type", facts["mime_type"]),
        ("Text lines", f"{facts['line_count']:,}" if facts["is_text"] else "Not a readable text file"),
    ]
    if facts["image"]:
        image = facts["image"]
        if "error" in image:
            rows.append(("Image metadata", f"Could not inspect image: {image['error']}"))
        else:
            rows.append(("Image metadata", f"{image['format']} {image['width']}x{image['height']}, mode {image['mode']}, frames {image['frames']}"))
    story = [
        Paragraph(f"Important Code Reference: {facts['repo_file']}", styles["Title"]),
        pdf_paragraph("A focused code/control-file explanation covering purpose, owner layer, APIs, variables, implementation signals, source excerpts, and debugging rules.", styles["SmallBody"]),
        pdf_table(rows, styles),
        Spacer(1, 0.12 * inch),
    ]
    add_pdf_file_fact_sections(story, facts, styles)
    SimpleDocTemplate(str(output_path), pagesize=LETTER, rightMargin=0.55 * inch, leftMargin=0.55 * inch, topMargin=0.55 * inch, bottomMargin=0.55 * inch).build(story)


def write_master_file_reference_docx(file_facts: list[dict], diagrams: dict[str, Path]) -> None:
    doc = setup_code_reference_document(
        "OMB Portfolio Builder Important Code Reference",
        "A generated Word reference for the code and control files that drive the builder, website, installer, Cloudflare AI worker, workflows, parser, compile workspace, and generated documentation.",
    )
    add_diagram(doc, diagrams["file-communication-map"], "Main file communication map for the builder and public website.")
    category_counts = Counter(facts["category"] for facts in file_facts)
    doc.add_heading("Coverage summary", level=1)
    add_table(doc, [
        ("Important files documented", str(len(file_facts))),
        ("Selection rule", "Only behavior-driving code/control files are included. Favicons, binary assets, and passive marker files are intentionally skipped here."),
        ("Categories", ", ".join(f"{name}: {count}" for name, count in sorted(category_counts.items()))),
    ])
    doc.add_heading("File index", level=1)
    add_table(doc, [(facts["repo_file"], f"{facts['category']}. {facts['folder_role']}") for facts in file_facts])
    doc.add_page_break()
    for facts in file_facts:
        doc.add_heading(f"Important Code Reference: {facts['repo_file']}", level=1)
        add_file_fact_summary(doc, facts)
        add_file_fact_sections(doc, facts)
        doc.add_page_break()
    compact_generated_docx(doc, keep_first_page_break=True)
    doc.save(FILE_REFERENCE_MASTER_DOCX)


def write_master_file_reference_pdf(file_facts: list[dict], diagrams: dict[str, Path]) -> None:
    if not REPORTLAB_AVAILABLE:
        return
    styles = pdf_styles()
    category_counts = Counter(facts["category"] for facts in file_facts)
    story = [
        Paragraph("OMB Portfolio Builder Important Code Reference", styles["Title"]),
        pdf_paragraph("A generated PDF reference for the code and control files that drive the builder, website, installer, Cloudflare AI worker, workflows, parser, compile workspace, and generated documentation.", styles["SmallBody"]),
    ]
    file_map = diagrams.get("file-communication-map")
    if file_map and file_map.exists():
        story.append(PdfImage(str(file_map), width=6.6 * inch, height=3.72 * inch))
    story.append(Paragraph("Coverage summary", styles["Heading1"]))
    story.append(pdf_table([
        ("Important files documented", str(len(file_facts))),
        ("Selection rule", "Only behavior-driving code/control files are included. Favicons, binary assets, and passive marker files are intentionally skipped here."),
        ("Categories", ", ".join(f"{name}: {count}" for name, count in sorted(category_counts.items()))),
    ], styles))
    story.append(Paragraph("File index", styles["Heading1"]))
    story.append(pdf_table([(facts["repo_file"], f"{facts['category']}. {facts['folder_role']}") for facts in file_facts], styles))
    story.append(PageBreak())
    for facts in file_facts:
        story.append(Paragraph(f"Important Code Reference: {facts['repo_file']}", styles["Heading1"]))
        rows = [
            ("File", facts["repo_file"]),
            ("Category", facts["category"]),
            ("Folder role", facts["folder_role"]),
            ("Size", f"{facts['size_bytes']:,} bytes"),
            ("Extension", facts["suffix"]),
            ("MIME type", facts["mime_type"]),
            ("Text lines", f"{facts['line_count']:,}" if facts["is_text"] else "Not a readable text file"),
        ]
        if facts["image"]:
            image = facts["image"]
            rows.append(("Image metadata", f"Could not inspect image: {image['error']}" if "error" in image else f"{image['format']} {image['width']}x{image['height']}, mode {image['mode']}, frames {image['frames']}"))
        story.append(pdf_table(rows, styles))
        add_pdf_file_fact_sections(story, facts, styles)
        story.append(PageBreak())
    SimpleDocTemplate(str(FILE_REFERENCE_MASTER_PDF), pagesize=LETTER, rightMargin=0.55 * inch, leftMargin=0.55 * inch, topMargin=0.55 * inch, bottomMargin=0.55 * inch).build(story)


def write_file_reference_docs(files: list[str], diagrams: dict[str, Path]) -> list[Path]:
    remove_directory(FILE_REFERENCE_DOCX_DIR)
    remove_directory(FILE_REFERENCE_PDF_DIR)
    FILE_REFERENCE_DOCX_DIR.mkdir(parents=True, exist_ok=True)
    FILE_REFERENCE_PDF_DIR.mkdir(parents=True, exist_ok=True)
    file_facts = [extract_file_facts(file) for file in important_code_files(files)]
    written: list[Path] = []
    for facts in file_facts:
        docx_path = FILE_REFERENCE_DOCX_DIR / code_reference_name(facts["repo_file"], ".docx")
        write_single_file_reference_docx(facts, docx_path)
        written.append(docx_path)
        pdf_path = FILE_REFERENCE_PDF_DIR / code_reference_name(facts["repo_file"], ".pdf")
        write_single_file_reference_pdf(facts, pdf_path)
        if pdf_path.exists():
            written.append(pdf_path)
    write_master_file_reference_docx(file_facts, diagrams)
    written.append(FILE_REFERENCE_MASTER_DOCX)
    write_master_file_reference_pdf(file_facts, diagrams)
    if FILE_REFERENCE_MASTER_PDF.exists():
        written.append(FILE_REFERENCE_MASTER_PDF)
    return written


def write_code_reference_docs(files: list[str], diagrams: dict[str, Path]) -> list[Path]:
    remove_directory(CODE_REFERENCE_DIR)
    remove_directory(CODE_REFERENCE_DOCX_DIR)
    remove_directory(CODE_REFERENCE_PDF_DIR)
    CODE_REFERENCE_DOCX_DIR.mkdir(parents=True, exist_ok=True)
    CODE_REFERENCE_PDF_DIR.mkdir(parents=True, exist_ok=True)
    code_files = [file for file in files if is_text_code_file(file)]
    facts_by_file = {repo_file: extract_source_facts(repo_file) for repo_file in code_files}
    written: list[Path] = []
    for repo_file in code_files:
        facts = facts_by_file[repo_file]
        docx_path = CODE_REFERENCE_DOCX_DIR / code_reference_name(repo_file, ".docx")
        write_single_code_reference_docx(facts, docx_path)
        written.append(docx_path)
        pdf_path = CODE_REFERENCE_PDF_DIR / code_reference_name(repo_file, ".pdf")
        write_single_code_reference_pdf(facts, pdf_path)
        if pdf_path.exists():
            written.append(pdf_path)
    write_master_code_reference_docx(code_files, facts_by_file, diagrams)
    written.append(CODE_REFERENCE_MASTER_DOCX)
    write_master_code_reference_pdf(code_files, facts_by_file, diagrams)
    if CODE_REFERENCE_MASTER_PDF.exists():
        written.append(CODE_REFERENCE_MASTER_PDF)
    return written


def add_generated_code_reference(doc: Document, files: list[str], diagrams: dict[str, Path]) -> None:
    doc.add_heading("Generated Word And PDF Code References", level=1)
    written = write_code_reference_docs(files, diagrams)
    doc.add_paragraph("The repository now includes generated Word and PDF code-reference documents rather than only Markdown notes. These documents are meant for source-level reading beside the actual files. They explain functions, variables, API endpoints, imports, event handlers, file relationships, security/authentication cues, compiler calls, installer behavior, and debugging paths.")
    add_table(doc, [
        ("Word folder", str(CODE_REFERENCE_DOCX_DIR.relative_to(REPO))),
        ("PDF folder", str(CODE_REFERENCE_PDF_DIR.relative_to(REPO))),
        ("Master Word reference", str(CODE_REFERENCE_MASTER_DOCX.relative_to(REPO))),
        ("Master PDF reference", str(CODE_REFERENCE_MASTER_PDF.relative_to(REPO)) if CODE_REFERENCE_MASTER_PDF.exists() else "PDF generation skipped because ReportLab is unavailable."),
        ("Artifacts generated", str(len(written))),
        ("Regeneration command", "python docs/build_complete_guide.py"),
    ])
    doc.add_page_break()

    for repo_file in [file for file in files if is_text_code_file(file)]:
        facts = extract_source_facts(repo_file)
        doc.add_heading(f"Code Reference: {repo_file}", level=1)
        doc.add_paragraph(facts["description"])
        add_table(doc, [
            ("Lines", f"{facts['line_count']:,}"),
            ("Size", f"{facts['size_bytes']:,} bytes"),
            ("Talks to", fact_relationship_summary(facts)),
            ("Functions discovered", str(len(facts["functions"]))),
            ("Variables discovered", str(len(facts["variables"]))),
            ("API endpoints mentioned", str(len(facts["endpoints"]))),
            ("Implementation signals", str(len(facts["signals"]))),
        ])
        if facts["signals"]:
            doc.add_heading("Signals detected", level=2)
            add_table(doc, [(f"{item['type']} at line {item['line']}", f"{item['meaning']} Evidence: {item['evidence']}") for item in facts["signals"][:20]])
        if facts["endpoints"]:
            doc.add_heading("Endpoints in this file", level=2)
            add_table(doc, [(item["endpoint"], f"Line {item['line']}. {endpoint_purpose(item['endpoint'])}") for item in facts["endpoints"][:20]])
        if facts["functions"]:
            doc.add_heading("Function details", level=2)
            for item in facts["functions"][:60]:
                doc.add_heading(f"{item['name']} (lines {item['line']}-{item['end_line']})", level=3)
                add_table(doc, function_detail_rows(item))
                add_code(doc, item["excerpt"])
        if facts["variables"]:
            doc.add_heading("Variables and constants", level=2)
            add_table(doc, [(item["name"], f"Line {item['line']}: {item['statement']}") for item in facts["variables"][:20]])
        if facts["imports"]:
            doc.add_heading("Imports and dependencies", level=2)
            add_table(doc, [(f"Line {item['line']}", item["statement"]) for item in facts["imports"][:20]])
        doc.add_heading("Opening snippet", level=2)
        add_code(doc, facts["snippet"] or "(empty file)")
        doc.add_page_break()


def add_generated_file_reference(doc: Document, files: list[str], diagrams: dict[str, Path]) -> None:
    doc.add_heading("Generated Word And PDF Important Code References", level=1)
    written = write_file_reference_docs(files, diagrams)
    selected_files = important_code_files(files)
    doc.add_paragraph("The repository now includes generated file-specific Word and PDF documents for the important code and control files that explain how the system actually works. This focused set avoids passive binary assets and marker files so the documentation stays useful for learning the builder internals.")
    add_table(doc, [
        ("Word folder", str(FILE_REFERENCE_DOCX_DIR.relative_to(REPO))),
        ("PDF folder", str(FILE_REFERENCE_PDF_DIR.relative_to(REPO))),
        ("Master Word reference", str(FILE_REFERENCE_MASTER_DOCX.relative_to(REPO))),
        ("Master PDF reference", str(FILE_REFERENCE_MASTER_PDF.relative_to(REPO)) if FILE_REFERENCE_MASTER_PDF.exists() else "PDF generation skipped because ReportLab is unavailable."),
        ("Important files documented", str(len(selected_files))),
        ("Artifacts generated", str(len(written))),
        ("Skipped here", "Favicons, app icons, binary assets, passive marker files, and generated reference documents."),
        ("Regeneration command", "python docs/build_complete_guide.py"),
    ])
    doc.add_heading("What this adds beyond the code reference", level=2)
    bullets(doc, [
        "The selected files are the places a developer actually starts when debugging app behavior.",
        "Each selected file gets source excerpts, implementation signals, API/DOM/storage clues, and safe-change notes.",
        "The master reference acts like a guided map of the app internals rather than a dump of every asset.",
    ])
    doc.add_page_break()


def add_complete_function_inventory(doc: Document) -> None:
    rows = extract_functions()
    doc.add_heading("Generated Function Inventory", level=1)
    doc.add_paragraph("This section is generated from the real source files. It does not replace reading the code, but it gives you a map of where functions live and what their names suggest they own.")
    add_table(doc, [
        ("Files scanned", "main.cjs, server.mjs, template-preview.js, script.js"),
        ("Functions found", str(len(rows))),
        ("How descriptions are made", "Descriptions are heuristic based on file name and function naming patterns. For exact behavior, inspect the source around the listed line."),
    ])
    doc.add_page_break()

    by_file: dict[str, list[tuple[str, int, str, str]]] = {}
    for row in rows:
        by_file.setdefault(row[0], []).append(row)
    for file_name, file_rows in by_file.items():
        doc.add_heading(f"Function Inventory: {file_name}", level=1)
        doc.add_paragraph(f"{file_name} contains {len(file_rows)} named functions discovered by the guide generator.")
        chunk_size = 45
        for index in range(0, len(file_rows), chunk_size):
            chunk = file_rows[index:index + chunk_size]
            add_table(doc, [(f"{name} (line {line_no})", purpose) for _, line_no, name, purpose in chunk])
            if index + chunk_size < len(file_rows):
                doc.add_paragraph("Continued on the next table.")
        doc.add_page_break()


def add_deep_detail_topics(doc: Document) -> None:
    for title, body in DEEP_DETAIL_TOPICS:
        doc.add_heading(f"Deep Detail: {title}", level=1)
        doc.add_paragraph(body)
        doc.add_heading("How to use this detail", level=2)
        bullets(doc, [
            "Start with the user-visible behavior.",
            "Find the file that owns that behavior.",
            "Trace the data handoff to the next file, endpoint, cache, or generated artifact.",
            "Change the smallest responsible piece and test the flow end to end.",
        ])
        doc.add_page_break()


def add_programming_syntax(doc: Document) -> None:
    for title, snippet, explanation in PROGRAMMING_SYNTAX_LESSONS:
        doc.add_heading(f"Programming Syntax: {title}", level=1)
        add_code(doc, snippet)
        doc.add_paragraph(explanation)
        doc.add_heading("How this appears in the project", level=2)
        doc.add_paragraph("You will see this pattern in main.cjs, server.mjs, template-preview.js, script.js, Cloudflare Worker code, HTML files, CSS files, JSON catalogs, and GitHub workflow YAML.")
        doc.add_page_break()


def add_shell_syntax(doc: Document) -> None:
    for title, snippet, explanation in SHELL_SYNTAX_LESSONS:
        doc.add_heading(f"Shell Or Script Syntax: {title}", level=1)
        add_code(doc, snippet)
        doc.add_paragraph(explanation)
        doc.add_heading("Why this matters", level=2)
        doc.add_paragraph("Shell scripts glue tools together. They are not the app UI, but they install tools, build releases, deploy workers, scan directories, and enforce rules.")
        doc.add_page_break()


def add_caching_methods(doc: Document, diagrams: dict[str, Path]) -> None:
    doc.add_heading("Caching Methods Used In The Builder And Website", level=1)
    add_diagram(doc, diagrams["compile-code-flow"], "Compile artifact caching happens inside the Compile Code workspace.")
    doc.add_paragraph("Caching means remembering a result so the app does not repeat slow or annoying work. This project uses several different caches, each with a different safety boundary.")
    add_table(doc, [(name, f"{where} Purpose: {purpose}") for name, where, purpose in CACHING_METHODS])
    doc.add_heading("Important safety rule", level=2)
    doc.add_paragraph("A cache must be scoped. Publishing authorization is scoped to the same repository, branch, user/machine scope, and expiration time. Compiler caches are scoped to source code, language, compiler path, runtime path, and file name. UI caches like update snooze are local convenience data and do not grant publishing access.")
    doc.add_page_break()
    for name, where, purpose in CACHING_METHODS:
        doc.add_heading(f"Cache Detail: {name}", level=1)
        if "Compiler" in name or "Compile" in name or "Terminal" in name or "source" in name:
            add_diagram(doc, diagrams["compile-code-flow"], f"Context diagram for {name}.")
        elif "Publishing" in name or "Extended trust" in name:
            add_diagram(doc, diagrams["release-pipeline"], f"Publishing authorization cache supports this release/publish workflow.")
        elif "HTTP" in name:
            add_diagram(doc, diagrams["frontend-backend-cloudflare"], f"HTTP no-store protects request freshness across frontend/backend paths.")
        add_table(doc, [
            ("Where implemented", where),
            ("Why it exists", purpose),
            ("What can go stale", "Cached data can become outdated when files, tools, versions, repositories, or user choices change."),
            ("How to refresh", "Use the explicit refresh/rebuild/authenticate action for that feature, clear the relevant local data, or restart the builder when appropriate."),
        ])
        doc.add_page_break()


def add_installer_details(doc: Document, diagrams: dict[str, Path]) -> None:
    doc.add_heading("How The Installer Works", level=1)
    add_diagram(doc, diagrams["release-pipeline"], "Installer and update behavior sits at the end of the release pipeline.")
    doc.add_paragraph("The installer is produced by electron-builder and customized by NSIS script code in build/installer.nsh. electron-builder gathers app files and creates the installer shell. NSIS controls Windows setup pages, duplicate install checks, shortcut behavior, Git installation checks, update behavior, and uninstall pages.")
    add_table(doc, [
        ("electron-builder", "Reads package.json build settings and packages the Electron app."),
        ("NSIS", "Runs installer wizard logic and custom script functions."),
        ("installer.nsh", "Adds OMB-specific setup pages, duplicate detection, update handling, shortcuts, and Git checks."),
        ("AppData install", "Default per-user install location avoids Program Files admin prompts."),
        ("Update install", "Existing installation is updated in place instead of creating a duplicate copy."),
    ])
    doc.add_heading("Installer flow", level=2)
    numbers(doc, [
        "customInit runs first.",
        "The installer checks for an existing registered install.",
        "If a current or newer install exists, setup stops.",
        "If an older install exists, setup treats the run as an update and uses the existing location.",
        "If no install exists, setup scans for duplicate builder copies.",
        "The custom tools page asks about desktop shortcut and publishing tools.",
        "Files are installed into the chosen per-user location.",
        "customInstall creates/removes shortcuts and optionally installs Git for Windows.",
        "The uninstaller has its own page explaining what will and will not be removed.",
    ])
    doc.add_page_break()
    for name, description in INSTALLER_FUNCTIONS:
        doc.add_heading(f"Installer Function: {name}", level=1)
        doc.add_paragraph(description)
        add_table(doc, [
            ("File", "build/installer.nsh"),
            ("Syntax family", "NSIS installer scripting"),
            ("Function job", description),
            ("Beginner note", "Installer functions run during setup/uninstall, not while the normal builder UI is being used."),
        ])
        doc.add_page_break()


def add_function_maps(doc: Document, diagrams: dict[str, Path]) -> None:
    for file_name, group_purpose, functions in FUNCTION_GROUPS:
        doc.add_heading(f"Function Map: {file_name}", level=1)
        if file_name == "main.cjs":
            add_diagram(doc, diagrams["electron-runtime"], "Electron startup functions sit inside this runtime flow.")
        elif file_name == "server.mjs":
            add_diagram(doc, diagrams["frontend-backend-cloudflare"], "server.mjs is the local backend layer in this map.")
        elif file_name == "template-preview.js":
            add_diagram(doc, diagrams["builder-to-site-files"], "template-preview.js is the builder frontend that sends edits toward parser output.")
        elif file_name == "script.js":
            add_diagram(doc, diagrams["cloudflare-ai-flow"], "script.js owns the public website behavior, including the AI chat request path.")
        doc.add_paragraph(group_purpose)
        add_table(doc, [(name, description) for name, description in functions])
        doc.add_heading("How to use this map", level=2)
        doc.add_paragraph("When debugging, find the user action first, then find the function group that owns it. For example, a compile terminal issue starts in template-preview.js, then calls server.mjs, then returns output back to template-preview.js.")
        doc.add_page_break()
        for name, description in functions:
            doc.add_heading(f"Function Detail: {file_name} -> {name}", level=1)
            doc.add_paragraph(description)
            add_table(doc, [
                ("File", file_name),
                ("Function", name),
                ("Responsibility", description),
                ("Debug question", "What input does this function receive, what output or state change does it produce, and what function calls it next?"),
            ])
            doc.add_page_break()


def add_architecture(doc: Document, diagrams: dict[str, Path]) -> None:
    for title, body, diagram in ARCHITECTURE_PAGES:
        doc.add_heading(title, level=1)
        doc.add_paragraph(body)
        diagram_name = ARCHITECTURE_DIAGRAM_BY_TITLE.get(title)
        if diagram_name:
            add_diagram(doc, diagrams[diagram_name], f"Generated block diagram: {title}.")
        doc.add_heading("Text version of the block diagram", level=2)
        add_code(doc, diagram)
        doc.add_heading("Beginner explanation", level=2)
        doc.add_paragraph("Read each arrow as 'this thing talks to the next thing.' The important rule is that editing happens locally in the builder, while viewing happens publicly on the website after generated static files are published.")
        doc.add_heading("What can break", level=2)
        bullets(doc, [
            "If local builder state is not saved, the parser may not see the latest edits.",
            "If Git authentication or branch state is wrong, Apply to site cannot safely push.",
            "If public hosting cache is stale, a phone may show an older copy until refresh or cache expiry.",
        ])
        doc.add_page_break()


def add_file_communication(doc: Document, diagrams: dict[str, Path]) -> None:
    doc.add_heading("How The Important Files Communicate", level=1)
    add_diagram(doc, diagrams["file-communication-map"], "Generated map of the main builder files and public website files.")
    doc.add_paragraph("This is the file-level mental model. The desktop app starts in main.cjs, the backend lives in server.mjs, the builder screen is template-preview.html with template-preview.css and template-preview.js, and the public website is index.html with styles.css, script.js, projects.json, and assets.")
    add_table(doc, [
        ("main.cjs -> server.mjs", "Electron starts the backend so the builder UI can call local APIs."),
        ("main.cjs -> template-preview.html", "Electron opens the builder window and loads the builder HTML file."),
        ("template-preview.html -> template-preview.js", "The builder HTML loads the JavaScript that handles editing, saving, previewing, compile code, and publishing UI."),
        ("template-preview.js -> server.mjs", "The builder uses fetch requests to ask the backend to save files, parse projects, run Git, run compilers, or publish."),
        ("server.mjs -> projects.json", "The parser writes public project data used by the website."),
        ("index.html -> script.js/styles.css", "The public page loads its behavior and visual design."),
        ("script.js -> projects.json", "The public website reads the generated catalog to render project windows."),
    ])
    doc.add_heading("Why this communication shape is useful", level=2)
    bullets(doc, [
        "The public website can stay static and fast.",
        "The builder can safely use local-only features without exposing them online.",
        "The parser creates a clean boundary between editable state and visitor-facing output.",
        "The same project data can be previewed locally before it is published.",
    ])
    doc.add_page_break()


def add_tool_build_map(doc: Document, diagrams: dict[str, Path]) -> None:
    doc.add_heading("What Tool Builds What", level=1)
    add_diagram(doc, diagrams["tool-build-map"], "Generated map showing which tool owns each part of the build and release process.")
    doc.add_paragraph("A common source of confusion is thinking one tool builds everything. In this project, each tool has a narrow job. pnpm installs dependencies, Electron runs the app, electron-builder packages the app, NSIS makes the installer, Git moves files, GitHub Actions automates release builds, Wrangler deploys Cloudflare Worker code, and compilers run project source code.")
    add_table(doc, [(tool, f"{builds} Used for: {where}") for tool, builds, where in TOOL_BUILD_ROWS])
    doc.add_heading("Beginner shortcut", level=2)
    doc.add_paragraph("When something fails, ask which tool owns that step. Installer failure points toward NSIS or electron-builder. AI backend failure points toward Cloudflare Worker or Wrangler. Website publish failure points toward Git, GitHub, branch state, or authentication.")
    doc.add_page_break()


def add_software_decisions(doc: Document) -> None:
    for title, body in SOFTWARE_DECISIONS:
        doc.add_heading(f"Software Engineering Decision: {title}", level=1)
        doc.add_paragraph(body)
        add_table(doc, [
            ("Decision", title),
            ("Reason", body),
            ("What it protects", "Predictable publishing, safer public/private separation, easier updates, clearer debugging, or better user experience."),
            ("Tradeoff", "The design may add more files or moving parts, but each moving part has a defined job."),
        ])
        doc.add_heading("How to evaluate this later", level=2)
        bullets(doc, [
            "If the feature becomes hard to maintain, check whether responsibilities are mixed across too many files.",
            "If users are confused, improve the builder UI or guide rather than exposing implementation details.",
            "If publishing becomes risky, tighten the parser and authentication boundary first.",
        ])
        doc.add_page_break()


def add_system_layers(doc: Document, diagrams: dict[str, Path]) -> None:
    doc.add_heading("AI, Frontend, Backend, Cloudflare, And Secrets", level=1)
    add_diagram(doc, diagrams["frontend-backend-cloudflare"], "Generated map of the frontend, backend, Cloudflare Worker, AI provider, and secret boundary.")
    doc.add_paragraph("This section separates the major layers so you know where to look when something breaks. The frontend draws and captures input. The backend performs protected work. Cloudflare hosts the public AI endpoint. Secrets stay server-side.")
    doc.add_page_break()
    for title, body, rows, snippet in SYSTEM_LAYER_SECTIONS:
        doc.add_heading(title, level=1)
        doc.add_paragraph(body)
        add_table(doc, rows)
        doc.add_heading("Representative code or command", level=2)
        add_code(doc, snippet)
        doc.add_heading("Debug question", level=2)
        doc.add_paragraph("Ask which layer owns the problem: did the UI send the request, did the backend receive it, did the service have the needed secret or tool, and did the response come back in the format the UI expects?")
        doc.add_page_break()


def add_workflows(doc: Document, files: list[str], diagrams: dict[str, Path]) -> None:
    workflows = [
        ("Content Creation Workflow", ["Open the builder.", "Edit profile, front-page text, portfolio areas, or projects.", "Save project or Save all sections.", "Open preview and inspect the exact public layout.", "Save draft.", "Apply to site after the target is authenticated."]),
        ("Project Creation Workflow", ["Click Add project.", "Choose category.", "Choose appearance.", "Enter title.", "Edit overview and sections.", "Add files, images, links, code blocks, and compile-code evidence.", "Save project so the parser can build the project preview."]),
        ("Publishing Workflow", ["Enter repository target.", "Authenticate GitHub write access.", "Load from target if the target already has content.", "Save draft.", "Apply to site.", "Git commits and pushes generated files.", "GitHub Pages or Cloudflare serves the updated website."]),
        ("Update Workflow", ["The installed app checks GitHub Releases.", "If a newer version exists, the app shows an update dialog.", "The updater downloads the installer.", "The app closes itself.", "The installer updates the existing AppData install.", "The app relaunches after the installer finishes."]),
        ("Compile Code Workflow", ["Create or import a source file.", "Pick the language.", "For Verilog/SystemVerilog, mark design files as Design and create or import a Testbench file.", "Keep $dumpfile and $dumpvars in the HDL testbench so the signal scope can draw waveforms.", "Save source.", "Run beautifier if desired.", "Compile or run.", "Read terminal output, messages log, and HDL scope when available.", "Append the source code to a project section when ready."]),
    ]
    for title, steps in workflows:
        doc.add_heading(title, level=1)
        if title in {"Content Creation Workflow", "Project Creation Workflow", "Publishing Workflow"}:
            add_diagram(doc, diagrams["builder-to-site-files"], f"Generated flow diagram supporting {title}.")
        if title == "Update Workflow":
            add_diagram(doc, diagrams["release-pipeline"], "Generated release/update pipeline diagram.")
        if title == "Compile Code Workflow":
            add_diagram(doc, diagrams["compile-code-flow"], "Generated Compile Code workflow diagram.")
        numbers(doc, steps)
        doc.add_heading("Plain-English mental model", level=2)
        doc.add_paragraph("A workflow is just a repeatable recipe. The builder turns complicated actions into visible buttons so you do not need to remember shell commands every time.")
        doc.add_page_break()

    for workflow in [f for f in files if f.startswith(".github/workflows/")]:
        text = (REPO / workflow).read_text(encoding="utf-8", errors="replace")
        doc.add_heading(f"GitHub Workflow File: {workflow}", level=1)
        doc.add_paragraph(FILE_DESCRIPTIONS.get(workflow, "GitHub Actions workflow file."))
        doc.add_heading("What a GitHub workflow is", level=2)
        doc.add_paragraph("A workflow is an automated script GitHub runs in the cloud. Instead of you manually building installers or enforcing branch rules, GitHub reads this YAML file and executes each job on a temporary machine.")
        doc.add_heading("Important YAML excerpt", level=2)
        add_code(doc, "\n".join(text.splitlines()[:80]))
        doc.add_heading("Tradeoff", level=2)
        doc.add_paragraph("Automated workflows reduce mistakes, but a bad workflow can block good merges. That is why the main branch gate should be strict but understandable.")
        doc.add_page_break()


def file_type(path: Path, repo_path: str) -> str:
    if repo_path.endswith(".gitkeep"):
        return "folder marker"
    if path.suffix:
        return path.suffix.lower().lstrip(".")
    return "no extension"


def add_files(doc: Document, files: list[str]) -> None:
    text_exts = {".js", ".mjs", ".cjs", ".html", ".css", ".json", ".md", ".txt", ".yml", ".yaml", ".toml", ".nsh", ""}
    for repo_file in files:
        path = REPO / repo_file
        doc.add_heading(f"File Explained: {repo_file}", level=1)
        doc.add_paragraph(FILE_DESCRIPTIONS.get(repo_file, "Tracked repository file used by the builder or public website."))
        add_table(doc, [
            ("Type", file_type(path, repo_file)),
            ("Size", f"{path.stat().st_size:,} bytes" if path.exists() else "missing locally"),
            ("Used by", "Builder app, public site, installer, Cloudflare, or GitHub workflow depending on the path."),
            ("Beginner idea", "A repository is a folder Git tracks. This file is one object Git can version, review, and publish."),
        ])
        if path.exists() and path.is_file() and path.suffix.lower() in text_exts:
            lines = path.read_text(encoding="utf-8", errors="replace").splitlines()
            doc.add_heading("Representative snippet", level=2)
            add_code(doc, "\n".join(lines[:24]) if lines else "(empty file)")
            if len(lines) > 24:
                doc.add_paragraph(f"Only the first 24 lines are shown here. The full file has {len(lines):,} lines and should be opened in the repository when editing.")
        else:
            doc.add_heading("Binary or asset note", level=2)
            doc.add_paragraph("This is treated as an asset. The guide explains its purpose, but the raw binary content is not printed because raw image/icon bytes would not be useful to read in Word.")
        doc.add_heading("Safe editing rule", level=2)
        bullets(doc, [
            "Make changes on a feature branch from development.",
            "Preview or test the behavior before merging.",
            "Do not edit generated/public files by hand if the builder parser is supposed to generate them.",
        ])
        doc.add_page_break()


def add_commands(doc: Document) -> None:
    for command, explanation, why in COMMAND_LESSONS:
        doc.add_heading(f"Command Explained: {command}", level=1)
        add_code(doc, command)
        add_table(doc, [
            ("What it does", explanation),
            ("Why it matters", why),
            ("Where used", "Local development, GitHub Actions, builder publishing, installer testing, or compiler execution."),
            ("Beginner warning", "Run commands from the correct folder. In this project, most development commands belong in the builder repository root."),
        ])
        doc.add_heading("How to read command output", level=2)
        doc.add_paragraph("Command output is the program talking back to you. Success output usually names what happened. Error output usually names a missing file, missing tool, denied permission, wrong branch, or failed compile step.")
        doc.add_page_break()


def add_features(doc: Document, diagrams: dict[str, Path]) -> None:
    for topic in FEATURE_TOPICS:
        doc.add_heading(f"Builder Feature: {topic}", level=1)
        if topic in {"Compile Code workspace", "HDL testbench requirement", "HDL signal scope", "Terminal output", "Messages log", "Append code to project"}:
            add_diagram(doc, diagrams["compile-code-flow"], f"Context diagram for {topic}.")
        elif topic in {"AI chat", "Cloudflare AI worker"}:
            add_diagram(doc, diagrams["cloudflare-ai-flow"], f"Context diagram for {topic}.")
        elif topic in {"Installer workflow", "Updater workflow", "Release tags", "Branch protection"}:
            add_diagram(doc, diagrams["release-pipeline"], f"Context diagram for {topic}.")
        elif topic in {"Project preview", "Full portfolio preview", "Project parser", "Asset synchronization"}:
            add_diagram(doc, diagrams["builder-to-site-files"], f"Context diagram for {topic}.")
        doc.add_paragraph(f"This page explains the {topic.lower()} area of the OMB builder and website system.")
        add_table(doc, [
            ("Owner view", "The builder exposes controls for creating, editing, saving, previewing, or publishing this area."),
            ("Visitor view", "The public website shows only parsed, approved output. Editing controls stay private."),
            ("Parser role", "The parser converts local builder state into website-safe data and markup."),
            ("Risk", "If this area is not saved before preview or publishing, the public site may show an older version."),
        ])
        doc.add_heading("Plain-English example", level=2)
        doc.add_paragraph("Think of the builder as the kitchen and the public website as the plate served to visitors. The parser is the person plating the content neatly. Visitors should never see the prep controls.")
        doc.add_page_break()


def add_tradeoffs(doc: Document) -> None:
    for title, body in TRADEOFFS:
        doc.add_heading(f"Tradeoff: {title}", level=1)
        doc.add_paragraph(body)
        add_table(doc, [
            ("Benefit", "The chosen direction fits the current builder goal: owner-controlled, local-first, preview-before-publish portfolio creation."),
            ("Cost", "Every architecture choice gives up something. Usually the cost is complexity, setup work, or less flexibility in another direction."),
            ("Decision rule", "Prefer the option that keeps publishing safe, previews accurate, and editing understandable."),
        ])
        doc.add_page_break()


def add_troubleshooting(doc: Document) -> None:
    for title, body in TROUBLESHOOTING_TOPICS:
        doc.add_heading(f"Troubleshooting: {title}", level=1)
        doc.add_paragraph(body)
        numbers(doc, [
            "Read the exact error message before changing anything.",
            "Confirm whether the issue is local builder state, Git state, installer state, public website cache, or backend service availability.",
            "Fix the smallest proven cause first.",
            "Save a clean checkpoint after the fix works.",
        ])
        doc.add_page_break()


def add_concepts(doc: Document) -> None:
    for concept in CONCEPTS:
        doc.add_heading(f"Beginner Concept: {concept}", level=1)
        doc.add_paragraph(f"{concept} is one of the building blocks behind the OMB builder and portfolio website. You do not need to memorize every term at once. Learn the role it plays, then come back when you see it in an error message or file name.")
        add_table(doc, [
            ("Simple meaning", "This concept helps explain how the app, website, GitHub workflow, installer, or backend behaves."),
            ("Where it appears", "Look for it in README instructions, source files, GitHub Actions logs, installer behavior, browser behavior, or terminal output."),
            ("How to debug it", "Ask: what is supposed to happen, what actually happened, what file or service owns that behavior, and what changed most recently?"),
        ])
        doc.add_heading("Tiny mental model", level=2)
        doc.add_paragraph("When confused, draw boxes. Put the user on the left, the thing they click in the middle, and the file or service that responds on the right. Most software problems become easier once the boxes are visible.")
        doc.add_page_break()


def add_closing(doc: Document) -> None:
    doc.add_heading("Final Operating Checklist", level=1)
    numbers(doc, [
        "Make feature changes on a feature or codex branch from development.",
        "Update README and the in-app guide when behavior changes.",
        "Run local syntax checks and useful smoke tests.",
        "Merge feature work into development.",
        "Only merge development into main for release-ready builds.",
        "Bump package.json before a new main release.",
        "Confirm GitHub Actions publish a fresh installer and stable latest asset.",
        "Open the builder, check update behavior, and verify the public website link downloads the latest installer.",
    ])
    add_note(doc, "The core rule", "The builder is the source of truth for creating content. The preview is the truth test before publishing. GitHub is the delivery path. The public website is the final read-only result.")


def main() -> None:
    files = tracked_files()
    package = load_package()
    diagrams = make_diagrams()
    doc = setup_document()
    add_cover(doc, package)
    add_reading_map(doc)
    add_entry_points(doc)
    add_architecture(doc, diagrams)
    add_flow_walkthroughs(doc, diagrams)
    add_file_communication(doc, diagrams)
    add_tool_build_map(doc, diagrams)
    add_software_decisions(doc)
    add_system_layers(doc, diagrams)
    add_data_contracts(doc)
    add_api_endpoint_catalog(doc)
    add_programming_syntax(doc)
    add_shell_syntax(doc)
    add_caching_methods(doc, diagrams)
    add_installer_details(doc, diagrams)
    add_function_maps(doc, diagrams)
    add_complete_function_inventory(doc)
    add_generated_code_reference(doc, files, diagrams)
    add_generated_file_reference(doc, files, diagrams)
    add_deep_detail_topics(doc)
    add_workflows(doc, files, diagrams)
    add_files(doc, files)
    add_commands(doc)
    add_features(doc, diagrams)
    add_tradeoffs(doc)
    add_troubleshooting(doc)
    add_concepts(doc)
    add_closing(doc)
    removed_breaks = compact_generated_docx(doc, keep_first_page_break=True)
    print(f"Compacted generated DOCX layout by removing {removed_breaks} page breaks.")
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
