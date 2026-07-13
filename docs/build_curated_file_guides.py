from __future__ import annotations

import shutil
import subprocess
import tempfile
import re
from collections import Counter
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


REPO = Path(__file__).resolve().parents[1]
DOCX_DIR = REPO / "docs" / "curated-file-guides-docx"
PDF_DIR = REPO / "docs" / "curated-file-guides-pdf"
MASTER_DOCX = REPO / "docs" / "OMB_Portfolio_Builder_Curated_File_Guides.docx"
MASTER_PDF = REPO / "docs" / "OMB_Portfolio_Builder_Curated_File_Guides.pdf"


def paragraph(text: str) -> str:
    return " ".join(str(text).strip().split())


JS_KEYWORDS = {
    "async", "await", "break", "case", "catch", "class", "const", "continue", "default", "delete",
    "do", "else", "export", "extends", "false", "finally", "for", "from", "function", "if",
    "import", "in", "instanceof", "let", "new", "null", "of", "return", "switch", "throw",
    "true", "try", "typeof", "undefined", "var", "void", "while", "yield",
}


def set_paragraph_shading(paragraph_obj, fill: str) -> None:
    p_pr = paragraph_obj._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_run_font(run, color: str = "0F172A", bold: bool = False, italic: bool = False, size: float = 7.5) -> None:
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def code_token_runs(line: str, language: str = "javascript") -> list[tuple[str, str, bool, bool]]:
    if language == "html":
        pattern = re.compile(r"(<!--.*?-->|</?[\w:-]+|[A-Za-z_:][-A-Za-z0-9_:.]*(?=\=)|\"[^\"]*\"|'[^']*'|=|>|/?>|\s+|.)")
        tokens = []
        for match in pattern.finditer(line):
            token = match.group(0)
            if token.startswith("<!--"):
                tokens.append((token, "64748B", False, True))
            elif token.startswith("<"):
                tokens.append((token, "0E7490", True, False))
            elif token.startswith(("\"", "'")):
                tokens.append((token, "B45309", False, False))
            elif token in {"=", ">", "/>"}:
                tokens.append((token, "475569", False, False))
            elif token.strip() and re.match(r"[A-Za-z_:]", token):
                tokens.append((token, "7C3AED", False, False))
            else:
                tokens.append((token, "0F172A", False, False))
        return tokens

    pattern = re.compile(
        r"(//.*|/\*.*?\*/|`[^`]*`|\"(?:\\.|[^\"])*\"|'(?:\\.|[^'])*'|\b\d+(?:\.\d+)?\b|\b[A-Za-z_$][\w$]*\b|\s+|.)"
    )
    tokens = []
    previous_non_space = ""
    for match in pattern.finditer(line):
        token = match.group(0)
        if token.startswith("//") or token.startswith("/*"):
            tokens.append((token, "64748B", False, True))
        elif token.startswith(("\"", "'", "`")):
            tokens.append((token, "B45309", False, False))
        elif token in JS_KEYWORDS:
            tokens.append((token, "1D4ED8", True, False))
        elif re.fullmatch(r"\d+(?:\.\d+)?", token):
            tokens.append((token, "7C3AED", False, False))
        elif re.fullmatch(r"[A-Za-z_$][\w$]*", token) and previous_non_space == "function":
            tokens.append((token, "0E7490", True, False))
        elif re.fullmatch(r"[A-Za-z_$][\w$]*", token):
            tokens.append((token, "0F172A", False, False))
        elif token in "{}[]().,;:=>+-*/%!?&|":
            tokens.append((token, "475569", False, False))
        else:
            tokens.append((token, "0F172A", False, False))
        if token.strip():
            previous_non_space = token
    return tokens


def add_code_block(doc: Document, code: str, language: str = "javascript", title: str = "Source code") -> None:
    heading = doc.add_paragraph()
    heading.paragraph_format.first_line_indent = Inches(0)
    heading.paragraph_format.space_before = Pt(4)
    heading.paragraph_format.space_after = Pt(2)
    heading_run = heading.add_run(title)
    heading_run.bold = True
    heading_run.font.color.rgb = RGBColor(15, 82, 110)
    heading_run.font.size = Pt(9)

    for raw_line in code.splitlines() or [""]:
        line = raw_line if raw_line else " "
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.first_line_indent = Inches(0)
        p.paragraph_format.right_indent = Inches(0.08)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        set_paragraph_shading(p, "F8FAFC")
        for token, color, bold, italic in code_token_runs(line, language):
            run = p.add_run(token)
            set_run_font(run, color=color, bold=bold, italic=italic)


@dataclass
class FunctionDoc:
    name: str
    signature: str
    params: str
    body: str
    source: str
    is_async: bool
    occurrence: int = 1


@dataclass
class VariableDoc:
    kind: str
    name: str
    initializer: str
    source: str
    scope: str = "top-level"


def split_top_level_commas(value: str) -> list[str]:
    parts: list[str] = []
    current: list[str] = []
    depth = 0
    quote = ""
    for char in value:
        if quote:
            current.append(char)
            if char == quote:
                quote = ""
            elif char == "\\":
                current.append(char)
            continue
        if char in {"'", '"', "`"}:
            quote = char
            current.append(char)
            continue
        if char in "([{":
            depth += 1
        elif char in ")]}" and depth:
            depth -= 1
        if char == "," and depth == 0:
            part = "".join(current).strip()
            if part:
                parts.append(part)
            current = []
        else:
            current.append(char)
    part = "".join(current).strip()
    if part:
        parts.append(part)
    return parts


def function_ranges(source: str) -> list[FunctionDoc]:
    """Extract top-level function blocks without relying on line numbers."""
    pattern = re.compile(r"(?m)^(async\s+)?function\s+([A-Za-z_$][\w$]*)\s*\(")
    results: list[FunctionDoc] = []
    counts: Counter[str] = Counter()
    matches = list(pattern.finditer(source))
    for match_index, match in enumerate(matches):
        is_async = bool(match.group(1))
        name = match.group(2)
        params_start = match.end()
        depth = 1
        index = params_start
        quote = ""
        while index < len(source) and depth:
            char = source[index]
            if quote:
                if char == "\\":
                    index += 2
                    continue
                if char == quote:
                    quote = ""
            else:
                if char in {"'", '"', "`"}:
                    quote = char
                elif char == "(":
                    depth += 1
                elif char == ")":
                    depth -= 1
            index += 1
        params = source[params_start : index - 1].strip()
        brace_index = source.find("{", index)
        if brace_index == -1:
            continue
        next_start = matches[match_index + 1].start() if match_index + 1 < len(matches) else len(source)
        segment = source[match.start() : next_start].rstrip()
        close_index = segment.rfind("\n}")
        if close_index != -1:
            full_source = segment[: close_index + 2]
        else:
            full_source = segment
        body_start = full_source.find("{")
        body_end = full_source.rfind("}")
        body = full_source[body_start + 1 : body_end] if body_start != -1 and body_end != -1 else source[brace_index + 1 : next_start]
        counts[name] += 1
        signature = f"{'async ' if is_async else ''}function {name}({params})"
        results.append(FunctionDoc(name=name, signature=signature, params=params, body=body, source=full_source, is_async=is_async, occurrence=counts[name]))
    return results


def brace_depth_before(source: str, position: int) -> int:
    depth = 0
    quote = ""
    index = 0
    while index < position:
        char = source[index]
        if quote:
            if char == "\\":
                index += 2
                continue
            if char == quote:
                quote = ""
        else:
            if char in {"'", '"', "`"}:
                quote = char
            elif char == "{":
                depth += 1
            elif char == "}":
                depth = max(0, depth - 1)
        index += 1
    return depth


def top_level_variables(source: str) -> list[VariableDoc]:
    results: list[VariableDoc] = []
    pattern = re.compile(r"(?m)^(const|let|var)\s+([A-Za-z_$][\w$]*|\{[^=]+?\}|\[[^=]+?\])\s*=\s*")
    for match in pattern.finditer(source):
        if brace_depth_before(source, match.start()) != 0:
            continue
        kind = match.group(1)
        raw_name = match.group(2).strip()
        name = raw_name
        if raw_name.startswith("{") or raw_name.startswith("["):
            name = "destructured " + raw_name[:54].replace("\n", " ")
        start = match.end()
        index = start
        depth = 0
        quote = ""
        while index < len(source):
            char = source[index]
            if quote:
                if char == "\\":
                    index += 2
                    continue
                if char == quote:
                    quote = ""
            else:
                if char in {"'", '"', "`"}:
                    quote = char
                elif char in "{[(":
                    depth += 1
                elif char in "}])" and depth:
                    depth -= 1
                elif char == ";" and depth == 0:
                    break
                elif char == "\n" and depth == 0:
                    next_text = source[index + 1 : index + 24].lstrip()
                    if next_text.startswith(("const ", "let ", "var ", "function ", "async function ")):
                        break
            index += 1
        initializer = source[start:index].strip()
        source_code = source[match.start() : index + 1].strip()
        results.append(VariableDoc(kind=kind, name=name, initializer=initializer, source=source_code))
    return results


def local_variables(func: FunctionDoc) -> list[VariableDoc]:
    results: list[VariableDoc] = []
    pattern = re.compile(r"(?m)\b(const|let|var)\s+([A-Za-z_$][\w$]*|\{[^=]+?\}|\[[^=]+?\])\s*(?:=\s*([^;\n]+))?")
    for match in pattern.finditer(func.body):
        raw_name = match.group(2).strip()
        name = raw_name
        if raw_name.startswith("{") or raw_name.startswith("["):
            name = "destructured " + raw_name[:60].replace("\n", " ")
        initializer = (match.group(3) or "").strip()
        results.append(VariableDoc(kind=match.group(1), name=name, initializer=initializer, source=match.group(0).strip(), scope=func.name))
    return results


def initializer_shape(initializer: str) -> str:
    value = initializer.strip()
    if not value:
        return "no explicit initializer"
    if value.startswith("{"):
        return "object literal"
    if value.startswith("["):
        return "array literal"
    if value.startswith("new "):
        return "constructed object"
    if value.startswith("async ") or "=>" in value or value.startswith("("):
        return "function or callback value"
    if value.startswith(("'", '"', "`")):
        return "string value"
    if re.match(r"^-?\d", value):
        return "numeric value"
    if value in {"true", "false"}:
        return "boolean value"
    if "process.env" in value:
        return "environment-derived value"
    if "Path" in value or "path." in value or ".join" in value:
        return "path-derived value"
    if "document." in value or "querySelector" in value or "getElementById" in value:
        return "DOM reference"
    if "fetch(" in value:
        return "network result"
    if "await " in value:
        return "awaited asynchronous result"
    return "computed value"


def object_keys(initializer: str, limit: int = 18) -> list[str]:
    text = initializer.strip()
    if not text.startswith("{"):
        return []
    keys: list[str] = []
    for match in re.finditer(r"(?m)^\s*([A-Za-z_$][\w$-]*|['\"][^'\"]+['\"])\s*:", text):
        key = match.group(1).strip("'\"")
        if key not in keys:
            keys.append(key)
        if len(keys) >= limit:
            break
    return keys


TOP_LEVEL_VARIABLE_DETAILS: dict[str, dict[str, str]] = {
    "server.mjs": {
        "root": "root is the absolute folder that contains server.mjs. Every local static-file lookup starts from this folder unless the app is deliberately publishing or compiling somewhere else. It is calculated from import.meta.url instead of the current shell directory so the backend behaves the same whether Electron, PowerShell, or another launcher starts it.",
        "portfolioRoot": "portfolioRoot is the folder treated as the publishable website mirror. It normally equals root, but OMB_PORTFOLIO_WORKSPACE can redirect it to a separate portfolio folder. Publishing functions use this boundary when copying and staging public website files.",
        "compileRoot": "compileRoot is the separate workspace for Compile Code projects. It is deliberately outside the public site folder so temporary source files, binaries, VCD waveforms, and compiler output do not get mixed into the portfolio website.",
        "port": "port is the TCP port used by the local HTTP server. PORT can override it for development or packaging, while the default supports the local builder preview without requiring a user to choose a port.",
        "host": "host is the network interface binding for the server. The default allows the local runtime to listen broadly, but write operations still rely on isLocalRequest so privilege is not granted merely because a browser can reach the server.",
        "execFileAsync": "execFileAsync is the Promise-based wrapper around Node's execFile. It lets Git, version checks, and other short commands use async/await instead of nested callbacks.",
        "packageJson": "packageJson holds metadata from package.json, especially the current builder version used by update checks, release comparisons, and status messages.",
        "types": "types maps file extensions to HTTP Content-Type headers. Static serving uses it so browsers interpret CSS as CSS, JavaScript as JavaScript, PDFs as PDFs, icons as icons, and uploaded assets as the right media type.",
        "draftPath": "draftPath is the private local draft catalog path. It stores builder-side edits that have not necessarily been applied to the public projects.json file.",
        "catalogPath": "catalogPath is the public portfolio catalog path. The website runtime reads this file to render the public project list, profile, sections, fun facts, links, and rich content.",
        "publishAuthCachePath": "publishAuthCachePath is the local file that remembers recent publishing authorization. It is intentionally local to the machine and target so GitHub verification does not have to repeat on every Apply to site click.",
        "publishPaths": "publishPaths is the allowlist of files and folders that can be copied and staged during publishing. This is a critical safety object: it prevents builder-only files, temp compiler output, credentials, and internal docs from being pushed as website content.",
        "publishAuthCacheTtlMs": "publishAuthCacheTtlMs is the normal one-day authorization window. It implements the rule that GitHub publishing should not ask again every time the app opens, but also should not trust an old session forever.",
        "publishAuthExtendedTtlMs": "publishAuthExtendedTtlMs is the thirty-day extended authorization window. It is used after repeated successful authorizations prove the same device and target are consistently owned by the same publisher.",
        "publishAuthHistoryWindowMs": "publishAuthHistoryWindowMs is the seven-day window used to count successful authorizations. The cache logic uses it to decide whether the user qualifies for extended trust.",
        "publishAuthExtendedThreshold": "publishAuthExtendedThreshold is the number of successful authorizations required before the builder extends trust. It turns the user's requested security rule into a concrete number.",
        "defaultSiteRepository": "defaultSiteRepository stores the optional repository configured through the environment. It gives packaged or developer-launched builds a default target without hardcoding one into the code.",
        "blockedAppUpdateVersions": "blockedAppUpdateVersions is a Set of release versions the updater must skip. It exists so a known bad release can be blocked without deleting every release artifact from GitHub.",
        "publishAuthorizationHelp": "publishAuthorizationHelp is the human-readable troubleshooting text shown when publishing is blocked. It explains the sequence the user should follow: configure target, authenticate, load if needed, then apply.",
        "gitCandidates": "gitCandidates lists command names and likely Windows Git paths. Git operations use this list to find Git even when PATH is incomplete, which is common on Windows installs.",
        "compileLanguageProfiles": "compileLanguageProfiles is the language rulebook for the IDE feature. Each language entry defines the default filename, valid extensions, display label, required tools, and installer hints. Compile, build, run, simulate, and beautify operations all depend on this map.",
        "compileToolCandidates": "compileToolCandidates maps compiler tool names to possible executable locations. It lets the app find gcc, g++, javac, java, iverilog, vvp, python, node, and LTspice even when a normal terminal cannot.",
        "compileToolCache": "compileToolCache remembers the resolved path for compiler tools. It prevents every compile click from scanning the filesystem again.",
        "compileToolVersionCache": "compileToolVersionCache remembers version strings for tools after they have been checked. This makes status panels faster and keeps repeated diagnostics from slowing the UI.",
        "portfolioAiInstructions": "portfolioAiInstructions is the system prompt for the portfolio assistant. It defines how the model should distinguish greetings, general electronics questions, portfolio-specific questions, public GitHub code requests, and uncertain answers.",
        "githubTextFilePattern": "githubTextFilePattern tells the AI source reader which GitHub files are worth fetching as text. It includes code, README, HDL, schematic-like text formats, scripts, and engineering project files.",
        "githubSkipFilePattern": "githubSkipFilePattern tells the AI source reader which GitHub files to skip because they are binary, huge, media-oriented, or not useful as prompt text.",
    },
    "script.js": {
        "grid": "grid points at the project-grid container. renderProjects writes category sections and project cards into this element, making it the public project's visible mounting point.",
        "searchInput": "searchInput points at the project search box. Search handlers read its value, build dropdown results, apply highlights, and decide which projects remain visible.",
        "projectFilters": "projectFilters points at the category filter container. renderCategoryFilters fills it with filter buttons derived from the actual catalog categories.",
        "filterButtons": "filterButtons stores the current filter button elements. It is updated after dynamic filters are rendered because the original NodeList would not include newly created buttons.",
        "projectCount": "projectCount displays the number of visible or total projects. It is updated after catalog load and render operations so the summary band matches the data.",
        "projectTrackCount": "projectTrackCount displays how many project categories/tracks exist. It is derived from the saved categories rather than hardcoded.",
        "projectTrackLabels": "projectTrackLabels displays the category names. It helps the top summary reflect Analog, Digital, Embedded, or any custom categories the builder creates.",
        "year": "year points at the footer year span. Startup code can set it from the current date so the footer remains current without manual HTML edits.",
        "funFactsCallout": "funFactsCallout is the near-top personality callout. renderFunFacts only reveals it when saved fun facts exist.",
        "heroEyebrow": "heroEyebrow is the small label above the hero title. renderSiteContent can replace it from the builder's front page settings.",
        "heroTitle": "heroTitle is the main public headline. It receives the owner-controlled front-page title instead of forcing a fixed sentence into every portfolio.",
        "heroCopy": "heroCopy is the supporting paragraph in the hero section. It is one of the fields affected by rich text and plain text styling choices.",
        "brandName": "brandName is the bold brand text in the header. renderProfile updates it to match the portfolio owner or site title.",
        "brandSubtitle": "brandSubtitle is the small subtitle beside the brand. It usually communicates local builder/public portfolio context or a tagline.",
        "brandIcon": "brandIcon is the image element for the OMB visual mark. Profile settings can change it when the builder is generalized for other users.",
        "brandText": "brandText is the OMB engraving text in the header logo area. It carries the small trademark-like label.",
        "headerAvatar": "headerAvatar is the clickable top-right profile image container. renderProfile reveals it only when a profile image exists.",
        "headerAvatarImage": "headerAvatarImage is the actual img element inside the avatar link. It receives the saved profile picture URL.",
        "heroImage": "heroImage is the main hero/background image element. Site content settings decide whether it is displayed.",
        "resumeSection": "resumeSection is the public resume panel. renderProfile hides or shows it depending on whether a resume URL exists.",
        "contactBand": "contactBand is the final dark contact area. It also contains the AI assistant panel, so layout code must keep chat growth from distorting contact information.",
        "profilePhoto": "profilePhoto is the large contact-section profile image. It is separate from the header avatar so desktop/mobile layouts can size them differently.",
        "contactTitle": "contactTitle is the heading for the contact panel. renderProfile updates it from the owner identity.",
        "contactIntro": "contactIntro is the contact-section introduction text. It can be empty, plain text, or rich content depending on saved profile data.",
        "contactDetails": "contactDetails is the list of direct contact methods such as email and phone. renderProfile builds clickable mailto and tel links inside it.",
        "contactLinks": "contactLinks is the container for external profile links such as GitHub, LinkedIn, resume, or custom websites.",
        "footerOwner": "footerOwner is the footer ownership text. It lets the footer follow the current profile owner rather than remaining generic.",
        "searchWrap": "searchWrap is the parent around the search box. The search dropdown attaches near it so results appear where the user is typing.",
        "aiAssistantForm": "aiAssistantForm is the chat form. Its submit handler starts answerAssistantQuestion.",
        "aiAssistantPanel": "aiAssistantPanel is the whole Ask My Portfolio panel. Growth and layout code use it to keep the assistant right-justified and scrollable.",
        "aiAssistantInput": "aiAssistantInput is the question field. It is cleared/focused by chat handlers and intentionally does not keep long instructional text inside the typed value.",
        "aiAssistantLog": "aiAssistantLog is the scrollable message history. appendAssistantMessage and replaceAssistantMessage write user and assistant messages into it.",
        "aiAssistantStatus": "aiAssistantStatus is the small status line used for ready/thinking/error states.",
        "aiClearChatButton": "aiClearChatButton is the explicit Clear chat control. It resets conversation history, source maps, and the visible log.",
        "categories": "categories is the runtime copy of project categories loaded from the catalog. It can include more than the original Analog, Digital, and Embedded groups.",
        "projects": "projects is the runtime project array. Rendering, search, assistant context, and project windows all read from it.",
        "siteSections": "siteSections stores custom main-page sections such as personal profile, hobbies, sports, social media, or professional material.",
        "funFacts": "funFacts stores simple fun fact strings. renderFunFacts turns at most a few lines into a near-top callout.",
        "funFactsRich": "funFactsRich stores rich-format fun facts when the builder saves styled content instead of plain strings.",
        "fieldStyles": "fieldStyles stores plain-field font, size, and color choices. It lets front-page, profile, and contact fields render consistently with builder formatting.",
        "siteContent": "siteContent stores front-page text and site-level settings such as hero copy and titles.",
        "siteContentRich": "siteContentRich stores rich versions of site-level text fields.",
        "profile": "profile stores owner identity: name, email, phone, images, resume, GitHub, LinkedIn, and other public contact details.",
        "profileRich": "profileRich stores rich versions of profile/contact fields so styled text can survive from builder to website.",
        "activeFilter": "activeFilter stores the currently selected project category filter.",
        "activeSectionDialogDrag": "activeSectionDialogDrag stores the drag state while a project window is being moved.",
        "activeSectionDialogResize": "activeSectionDialogResize stores the resize state while a project window edge is being dragged.",
        "sectionDialogDragEnabled": "sectionDialogDragEnabled prevents drag handlers from being installed more than once.",
        "isApplyingSectionRoute": "isApplyingSectionRoute prevents browser history updates from recursively triggering themselves while a saved route is being restored.",
        "searchPanel": "searchPanel stores the dropdown result panel created for search.",
        "searchStatus": "searchStatus stores the status element inside the search dropdown.",
        "searchLimitSelect": "searchLimitSelect stores the dropdown control that lets the visitor choose 5, 10, 15, or 20 results.",
        "currentSearchResults": "currentSearchResults is the current ranked search-result array. It lets click handlers navigate results without recalculating the search.",
        "assistantSourceCounter": "assistantSourceCounter creates stable ids for assistant source buttons.",
        "assistantSourceMap": "assistantSourceMap connects source button ids to source records so clicking a source opens the right place.",
        "assistantChatHistory": "assistantChatHistory stores recent user and assistant messages. It gives the chatbot conversation memory.",
        "searchableEntries": "searchableEntries is the full in-browser search index. It contains project titles, sections, files, captions, page sections, uploaded text, and electronics keyword expansions.",
        "searchResultLimit": "searchResultLimit stores how many dropdown results to display.",
        "searchHighlightTimer": "searchHighlightTimer debounces highlight updates so typing does not repaint the page too aggressively.",
        "searchHighlightName": "searchHighlightName is the CSS class applied to highlighted search matches.",
        "sectionRouteKey": "sectionRouteKey identifies portfolio section entries in browser history state.",
        "sectionRouteParams": "sectionRouteParams defines the URL parameter names for project id, section index, and nested subsection path.",
        "supportedCodeLanguages": "supportedCodeLanguages is the public language list for code rendering and detection. It names C, C++, Verilog/SystemVerilog, Java, JavaScript, Python, HTML, and related aliases.",
        "legacyTemplateSkins": "legacyTemplateSkins maps older template ids to the newer appearance-only template ids. It keeps older saved projects rendering after template behavior changed.",
        "defaultSiteContent": "defaultSiteContent is the fallback front-page copy for a blank install. It keeps the site readable before the builder user customizes hero text and site messaging.",
        "defaultProfile": "defaultProfile is the empty-owner fallback. It defines all expected profile fields so renderProfile can run even before the user enters contact details.",
    },
}


def function_references_for_variable(name: str, source: str) -> list[str]:
    if name.startswith("destructured"):
        return []
    refs: list[str] = []
    for func in function_ranges(source):
        if re.search(rf"\b{re.escape(name)}\b", func.body):
            refs.append(func.name)
        if len(refs) >= 10:
            break
    return refs


def local_variable_usage(name: str, body: str) -> list[str]:
    if name.startswith("destructured") or not re.match(r"^[A-Za-z_$][\w$]*$", name):
        return []
    uses: list[str] = []
    if re.search(rf"\breturn\s+{re.escape(name)}\b", body):
        uses.append("returned to the caller")
    if re.search(rf"\b{re.escape(name)}\.(map|filter|reduce|forEach|sort)\s*\(", body):
        uses.append("iterated or transformed as a collection")
    if re.search(rf"\b{re.escape(name)}\.(push|set|add|append|appendChild)\s*\(", body):
        uses.append("mutated as a collection or DOM container")
    if re.search(rf"\b{re.escape(name)}\.[A-Za-z_$][\w$]*", body):
        props = sorted(set(re.findall(rf"\b{re.escape(name)}\.([A-Za-z_$][\w$]*)", body)))[:8]
        if props:
            uses.append("accesses properties/methods " + ", ".join(props))
    call_names = sorted(set(re.findall(rf"\b([A-Za-z_$][\w$]*)\s*\([^)]*\b{re.escape(name)}\b", body)))[:8]
    call_names = [item for item in call_names if item not in {"if", "for", "while", "switch", "return"}]
    if call_names:
        uses.append("passed into " + ", ".join(call_names))
    if re.search(rf"\b{re.escape(name)}\s*=", body):
        uses.append("assigned or updated later in the function")
    return uses[:5]


def variable_category(name: str, initializer: str, file_name: str) -> str:
    lower = name.lower()
    init = initializer.lower()
    if name.startswith("destructured"):
        return "imported or unpacked API handles"
    if "root" in lower or "path" in lower or "dir" in lower or "folder" in lower:
        return "filesystem location"
    if "url" in lower or "remote" in lower or "endpoint" in lower:
        return "network or routing value"
    if "cache" in lower or "history" in lower or "stack" in lower or "state" in lower:
        return "runtime state or cache"
    if "profile" in lower or "project" in lower or "category" in lower or "catalog" in lower or "section" in lower:
        return "portfolio data state"
    if "tool" in lower or "compiler" in lower or "language" in lower or "compile" in lower:
        return "compiler or language configuration"
    if "assistant" in lower or "ollama" in lower or "openai" in lower or "portfolioai" in lower or lower.endswith("ai"):
        return "AI configuration or conversation state"
    if "git" in lower or "publish" in lower or "auth" in lower or "credential" in lower or "branch" in lower:
        return "Git publishing and authorization state"
    if "element" in lower or "dialog" in lower or "grid" in lower or "input" in lower or "button" in lower or "section" in lower and file_name == "script.js":
        return "DOM reference or UI state"
    if "class" in lower or "style" in lower or "color" in lower or "template" in lower:
        return "appearance configuration"
    if "map" in init or "set(" in init:
        return "lookup collection"
    return "temporary calculation"


def explain_variable(variable: VariableDoc, file_name: str, top_level: bool = False, source: str = "", function_body: str = "") -> str:
    name = variable.name
    custom = TOP_LEVEL_VARIABLE_DETAILS.get(file_name, {}).get(name) if top_level else None
    category = variable_category(name, variable.initializer, file_name)
    shape = initializer_shape(variable.initializer)
    keys = object_keys(variable.initializer)
    if custom:
        base = custom
    else:
        init_text = variable.initializer.strip()
        if len(init_text) > 180:
            init_text = init_text[:177].rstrip() + "..."
        if shape == "no explicit initializer":
            base = f"{name} starts without an immediate value. In this file that usually means it is filled by a loop, branch, callback, or later assignment inside the same function. "
        elif shape == "object literal":
            base = f"{name} stores a structured object for {category}. "
        elif shape == "array literal":
            base = f"{name} stores a list for {category}. "
        elif shape == "DOM reference":
            base = f"{name} stores a DOM element reference. "
        elif shape == "path-derived value":
            base = f"{name} stores a resolved filesystem or route value. "
        elif shape == "awaited asynchronous result":
            base = f"{name} stores the completed result of an awaited operation. "
        elif shape == "network result":
            base = f"{name} stores data returned from a network call. "
        else:
            base = f"{name} stores a {shape} for {category}. "
        if init_text:
            base += f"It is initialized from {init_text}. "
    detail_parts: list[str] = []
    lower = name.lower()
    if "root" in lower or "path" in lower or "dir" in lower:
        detail_parts.append("It controls a boundary or destination for file work, so an incorrect value can redirect uploads, previews, compiler artifacts, or published assets.")
    elif "profile" in lower or "project" in lower or "category" in lower or "catalog" in lower:
        detail_parts.append("It carries portfolio content, so rendering, search, AI context, and publishing expect its shape to stay consistent.")
    elif "tool" in lower or "compiler" in lower or "language" in lower or "compile" in lower:
        detail_parts.append("It supports the compile workspace by describing language rules, executable locations, tool status, or compiler output.")
    elif "assistant" in lower or "ollama" in lower or "openai" in lower or "portfolioai" in lower or lower.endswith("ai"):
        detail_parts.append("It influences assistant behavior: conversation memory, source display, model prompts, backend routing, or context selection.")
    elif "git" in lower or "publish" in lower or "auth" in lower or "credential" in lower or "branch" in lower:
        detail_parts.append("It affects publishing, where mistakes can change the public website or expose the wrong files.")
    elif "dialog" in lower or "grid" in lower or "input" in lower or "button" in lower or "element" in lower:
        detail_parts.append("It anchors UI behavior to a specific browser element or window state.")
    elif name.startswith("destructured"):
        detail_parts.append("It unpacks API handles from another object, giving the rest of the file short direct names for those APIs.")
    if keys:
        detail_parts.append(f"Visible object fields include {', '.join(keys)}. Those fields form the contract consumed by related functions.")
    if top_level and source:
        refs = function_references_for_variable(name, source)
        if refs:
            detail_parts.append(f"Functions that reference it include {', '.join(refs)}.")
    if not top_level and function_body:
        uses = local_variable_usage(name, function_body)
        if uses:
            detail_parts.append("Within the function it is " + "; ".join(uses) + ".")
    return paragraph(base.rstrip() + " " + " ".join(detail_parts))


def add_variable_section(doc: Document, guide: dict) -> None:
    file_name = guide["file"]
    if file_name not in {"server.mjs", "script.js"}:
        return
    source = (REPO / file_name).read_text(encoding="utf-8")
    variables = top_level_variables(source)
    doc.add_heading("Chapter 2: Top-Level Objects And Variables", level=1)
    doc.add_paragraph(
        "This chapter explains the long-lived objects and variables before the function walkthrough. Think of these top-level values as the workshop benches, labeled drawers, shared measuring tools, and safety rails of the file. A function can walk into the workshop, use these shared objects, and leave without recreating the whole environment from scratch. They are explained by meaning and use, not by line number."
    )
    grouped: dict[str, list[VariableDoc]] = {}
    for variable in variables:
        category = variable_category(variable.name, variable.initializer, file_name)
        grouped.setdefault(category, []).append(variable)
    for category, items in grouped.items():
        doc.add_heading(category.title(), level=2)
        for variable in items:
            doc.add_heading(variable.name, level=3)
            doc.add_paragraph(explain_variable(variable, file_name, top_level=True, source=source))
            add_code_block(doc, variable.source, language="javascript", title=f"Source excerpt for {variable.name}")


def parse_parameters(params: str) -> list[str]:
    if not params:
        return []
    if params.startswith("{") and "}" in params:
        inner = params[1 : params.find("}")]
        return split_top_level_commas(inner)
    return split_top_level_commas(params)


def parameter_label(param: str) -> str:
    cleaned = param.strip().lstrip("...")
    cleaned = cleaned.split("=")[0].strip()
    cleaned = cleaned.replace("{", "").replace("}", "").strip()
    return cleaned or param.strip()


def explain_parameter(param: str) -> str:
    label = parameter_label(param)
    lower = label.lower()
    default_note = ""
    if "=" in param:
        default_note = " The default value is intentional: it lets the function fail softly or produce a predictable empty result when the caller omits that field."
    if label in {"request", "response"}:
        return paragraph(f"{label} is the Node HTTP object passed into the handler. The function reads request details or writes response details through this object instead of depending on browser globals.{default_note}")
    if "project" in lower:
        return paragraph(f"{label} identifies or carries the project being built, rendered, compiled, searched, or published. It links this function back to one portfolio project instead of letting work spill across unrelated projects.{default_note}")
    if "language" in lower:
        return paragraph(f"{label} names the programming or hardware-description language. The function uses it to choose file extensions, highlighting rules, compilers, simulators, labels, or output behavior.{default_note}")
    if "file" in lower or "path" in lower or "url" in lower or "remote" in lower:
        return paragraph(f"{label} points at an external resource, local file, route, Git remote, or public link. The function normally normalizes or validates it before using it so the app does not treat unsafe or malformed paths as trusted input.{default_note}")
    if "code" in lower or "source" in lower:
        return paragraph(f"{label} is source text supplied by the builder, project catalog, GitHub, or an uploaded file. The function inspects, formats, saves, compiles, highlights, or extracts meaning from this text.{default_note}")
    if "query" in lower or "question" in lower or "intent" in lower:
        return paragraph(f"{label} is user language: a search query, AI question, or classified assistant intent. It drives scoring, context selection, routing, or AI response behavior.{default_note}")
    if "options" in lower or "config" in lower or "payload" in lower or "context" in lower or "data" in lower:
        return paragraph(f"{label} is a structured object rather than one small value. It lets the caller pass several related settings or pieces of state together so the function can make one coordinated decision.{default_note}")
    if "element" in lower or "dialog" in lower or "node" in lower or "root" in lower:
        return paragraph(f"{label} is a DOM object or parsed content node. The function reads from it, writes into it, or uses it as the current place where UI or project content should be rendered.{default_note}")
    if "value" in lower or "text" in lower or "title" in lower or "name" in lower or "label" in lower:
        return paragraph(f"{label} is human-facing text. The function cleans, limits, displays, compares, or turns it into a safe identifier so the UI and generated site remain readable.{default_note}")
    return paragraph(f"{label} is the caller-supplied input used by this function. The function interprets it according to its local responsibility, often converting it into safer, more predictable data before returning or rendering anything.{default_note}")


SERVER_IMPORTANT: dict[str, list[str]] = {
    "compileAndRunCode": [
        paragraph("This is the central Compile Code brain. It receives the project id, active source file, language, action requested by the user, optional stdin, and any workspace files that belong to the project. From that one payload it decides whether the user is only beautifying code, saving code, compiling, building, running, simulating HDL, or reading waveform output. It is async because every serious action here touches the filesystem, discovers tools, writes temporary workspaces, runs external compiler processes, and waits for terminal output."),
        paragraph("The function is needed because the builder treats each project like a small IDE workspace. JavaScript, Python, Java, C, C++, Verilog, SystemVerilog, LTspice, and HTML all have different compile/run rules. This function keeps those rules behind one API endpoint so the frontend can ask for a compile operation without knowing every compiler command. Without it, Compile Code would degrade into a text box with no reliable build, run, simulate, cache, or output behavior."),
        paragraph("Important internal helpers include ensureRunSource, append, and missing. ensureRunSource writes the active source into the compile workspace before a run. append converts command results into terminal-style output. missing returns a helpful response when a required tool is unavailable. Those small local functions matter because they keep the larger language branches readable."),
    ],
    "handleApi": [
        paragraph("handleApi is the backend router. It receives the parsed URL and decides which local API feature should run: project catalog reads, uploads, draft saves, compile actions, compiler installation, target authentication, load from target, apply to site, update checks, security reports, AI calls, or builder download reports. In a framework such as Express this would be spread across route handlers. Here it is one explicit switch-like controller."),
        paragraph("The function protects the boundary between a browser click and local machine power. Before write operations, it checks whether the request is local. It reads request JSON only when needed, calls the specific operation function, and returns a normalized JSON response. Without handleApi, the server could still serve files, but the builder would have no organized command surface."),
    ],
    "publishSiteChanges": [
        paragraph("publishSiteChanges is the final publishing operation behind Apply to site. It assumes authorization has been proven or supplied, synchronizes publishable files into the portfolio mirror, checks Git status, stages only approved paths, creates a commit when there are changes, and pushes to the selected branch. It is async because Git commands and file synchronization are external operations that take time and can fail."),
        paragraph("This function is important because publishing is the one place where local draft work becomes public website content. It must therefore be strict. If it staged arbitrary files or pushed without checking state, builder internals, private drafts, or stale branches could leak to the website. The function works with assertPublishAccess, syncPortfolioPublishFiles, getGitStatus, runPublishGit, and stageablePublishPaths."),
    ],
    "assertPublishAccess": [
        paragraph("assertPublishAccess is the gatekeeper that decides whether the current machine and target are allowed to change the website. It checks the configured publishing target, compares cached authorization against the current remote and branch, respects the once-per-day authentication window, and can force a fresh verification when needed. It returns the access object that publishing functions use as proof."),
        paragraph("Without this function, Apply to site would either ask for GitHub authorization too often or push too freely. The cache logic is what makes the app usable day to day while still preventing an unauthenticated installation from silently modifying Maurice's website."),
    ],
    "authenticateGitHubForTarget": [
        paragraph("authenticateGitHubForTarget is the interactive setup path for a publishing target. It validates the target remote, initializes or fixes the local Git repository when necessary, stores credentials when credentials are supplied, verifies write access, saves the target configuration, and records the authorization cache. It is the correct first step before Load from target or Apply to site."),
        paragraph("The function exists because merely typing a GitHub URL does not prove the builder can read from or write to that repository. Authentication must be completed, verified, and remembered before the target is considered usable."),
    ],
    "syncFromPublishTarget": [
        paragraph("syncFromPublishTarget is the import path. After a target is authenticated, this function pulls or fetches the latest public portfolio files from the selected repository and copies compatible website assets back into the local builder workspace. It is what lets a fresh machine reconstruct the current portfolio state from the website repository."),
        paragraph("This function is deliberately separate from authentication. Authenticate target proves access. Load from target reads the target. Keeping those concepts separate makes the UI clearer and reduces accidental overwrites."),
    ],
    "downloadAndLaunchAppUpdate": [
        paragraph("downloadAndLaunchAppUpdate is the updater handoff. It checks release information, downloads the installer, writes a PowerShell script that can wait for the current process to exit, launches the installer or updater quietly, and schedules the current app to close. It must be async because it performs network download, file writes, and process launches."),
        paragraph("The function is needed because a running Electron app cannot safely replace all of its own files while those files are in use. The handoff script becomes a temporary helper that finishes the update after the current process gets out of the way."),
    ],
    "handlePortfolioAi": [
        paragraph("handlePortfolioAi is the backend side of Ask My Portfolio. It reads the visitor's question, conversation history, intent, and context, enriches that context when public sources are allowed, then chooses an AI provider path. It can call OpenAI when configured, call Ollama locally when available, or fall back to rule-based answers when no model is available."),
        paragraph("The function matters because the public script should not hold private API keys or perform privileged source fetching directly. It also gives the assistant one place to enforce prompt rules about portfolio-specific answers, general engineering answers, public source usage, and uncertainty."),
    ],
    "runProcess": [
        paragraph("runProcess is the safe wrapper around external commands. Compilers, Git, Winget, Java, vvp, and installer helpers all eventually need child processes. This function starts the command with controlled options, captures stdout and stderr, enforces a timeout, kills process trees on Windows when necessary, and resolves to a structured result instead of letting raw process behavior leak through the API."),
        paragraph("Without this wrapper, every compiler and Git function would need to repeat timeout logic, stdout/stderr collection, and error handling. More importantly, failures would be inconsistent and the frontend terminal would receive unreliable output."),
    ],
    "parseVcdScopeText": [
        paragraph("parseVcdScopeText reads VCD waveform text from HDL simulation and converts it into signals, time points, values, and scope metadata that a frontend scope viewer can draw. VCD files are compact event logs, not friendly UI data. This parser turns symbol changes over time into understandable signal histories."),
        paragraph("The function is essential for the SystemVerilog/Verilog simulator feature. A simulation that produces only text is useful, but a scope view requires structured waveform data. Without this parser, the scope icon would have no signal timeline to show."),
    ],
}


SCRIPT_IMPORTANT: dict[str, list[str]] = {
    "loadProjectCatalog": [
        paragraph("loadProjectCatalog is the public website startup pipeline. It chooses embedded preview data when the builder is previewing, otherwise it fetches projects.json from the published site. After that it hydrates the major runtime state: categories, projects, profile, site content, fun facts, custom sections, text styles, and rich text variants. Then it renders the visible page and restores any deep-linked project window."),
        paragraph("This function is async because fetching projects.json can take time and can fail. It is the reason the static HTML shell becomes a real portfolio. Without it, the page would show placeholders but no saved projects, profile details, resume, fun facts, project categories, search index, or assistant context."),
    ],
    "renderProfile": [
        paragraph("renderProfile applies owner identity to the public page. It updates the displayed name, tagline, profile photo, brand/title text, email link, phone link, GitHub/resume links, contact cards, resume viewer, metadata hints, and social/contact links. It is the bridge between saved profile data and the visible recruiter-facing contact section."),
        paragraph("The function is important because the same website shell can serve Maurice's portfolio or a fresh user's portfolio. Profile data cannot be hardcoded if the builder is intended to be reusable. Without renderProfile, the page would either remain generic or require manual HTML editing for every owner."),
    ],
    "renderRichContent": [
        paragraph("renderRichContent is the renderer for builder-created rich content. It accepts blocks representing paragraphs, images, formulas, files, links, code blocks, captions, and styled text, then converts those blocks into safe public HTML. It decides whether an image should appear inline, whether a file should be downloadable, how captions attach, and how text styles survive into the website."),
        paragraph("This function matters because the builder is not just saving plain text. It supports formatted explanations, pasted images, formulas, code, and files. Without renderRichContent, much of what the builder collects would be flattened or lost when the portfolio is published."),
    ],
    "sanitizeRichInlineHtml": [
        paragraph("sanitizeRichInlineHtml is the safety filter for rich inline HTML. Builder editors may store spans, links, bold text, colors, and other inline markup, but the public website should not blindly inject arbitrary HTML. This function removes unsafe tags and attributes while preserving the formatting the portfolio actually needs."),
        paragraph("The function prevents a content field from becoming a script injection surface. It works with renderRichContent and renderRichFieldContent, which need formatted HTML but should only receive controlled markup."),
    ],
    "tokenizedCodeHtml": [
        paragraph("tokenizedCodeHtml is the lightweight syntax highlighter used by the public site. It escapes the code first, then wraps comments, strings, numbers, keywords, preprocessor directives, and other recognized tokens in spans with CSS classes. It supports the major languages that the builder exposes, including C, C++, SystemVerilog, LTspice-oriented text, Java, JavaScript, Python, and HTML-like code."),
        paragraph("The function exists so code added to a project reads like code rather than plain paragraph text. Without it, source examples in the portfolio would be visually flat, harder to scan, and less convincing to technical recruiters."),
    ],
    "answerAssistantQuestion": [
        paragraph("answerAssistantQuestion is the public chatbot orchestrator. It receives the user's typed question, appends it to chat history, classifies the intent, gathers relevant portfolio context, creates a pending assistant answer, calls the backend AI when appropriate, falls back to local answers when needed, renders sources only when relevant, and restores focus to the input."),
        paragraph("This function is async because the AI backend may need network time and because context gathering can include indexed file text. It is the reason Ask My Portfolio behaves like a conversation instead of a static search form."),
    ],
    "assistantContextForQuestion": [
        paragraph("assistantContextForQuestion builds the packet of portfolio evidence sent to the AI backend. It uses the question intent, search results, named project matches, profile links, project summaries, and available evidence to decide what the model should see. It is selective by design so a generic greeting does not receive a pile of unrelated project context."),
        paragraph("Without this function, the assistant would either know too little about the portfolio or always overload the model with the same generic links. This is the function that makes the AI choose context based on the conversation."),
    ],
    "assistantSourcesForDisplay": [
        paragraph("assistantSourcesForDisplay chooses which portfolio links appear under an assistant answer. The user complained earlier that related links were generic; this function is the fix-point for that behavior. It scores source relation to the prompt, respects named projects and topic tokens, and avoids showing source buttons that do not belong to the answer context."),
        paragraph("The function matters because sources are a promise. A displayed link should help verify or explore the answer. If this function is too loose, the assistant looks random even when the text answer is good."),
    ],
    "openParsedSection": [
        paragraph("openParsedSection opens a project section or nested subsection as a full-window public view. It finds the project, finds the section, walks the requested subsection path, applies the selected appearance template, renders the overview and child sections, updates navigation buttons, and syncs browser history."),
        paragraph("This function is the heart of the website's project-window behavior. Without it, clicking Design, Simulation, Documentation, or any user-created section would either scroll the page awkwardly or render content in the wrong place."),
    ],
    "ensureSectionDialog": [
        paragraph("ensureSectionDialog creates and wires the reusable full-window dialog used for project sections. It builds the top controls, back/forward behavior, close behavior, content container, and event delegation for child cards. It creates the window once and then later calls reuse it."),
        paragraph("The function exists so every project section window behaves consistently. If each click created its own markup by hand, close buttons, back buttons, keyboard behavior, and nested navigation would drift apart."),
    ],
    "renderProjects": [
        paragraph("renderProjects rebuilds the visible project directory from current state. It checks the active category filter and search query, groups visible projects by category, creates the category sections, and updates the summary count. It is called after catalog load, filter changes, and search changes."),
        paragraph("The function is the public directory refresh mechanism. Without it, filters and searches might update state internally but the page would not visually change."),
    ],
    "goToSearchResult": [
        paragraph("goToSearchResult turns a search dropdown item into navigation. Depending on what the result represents, it may open a project window, jump to a main page section, highlight text on the current page, or open a file/link. This makes search feel like a search engine instead of a passive list."),
        paragraph("The function is necessary because search results can point to many different kinds of things: project titles, section overviews, files, custom sections, contact links, or page text. One navigation rule would not work for all of them."),
    ],
}


def classify_function(file_name: str, func: FunctionDoc) -> tuple[str, str]:
    name = func.name
    lower = name.lower()
    body = func.body.lower()
    if file_name == "server.mjs":
        if name in {"securityHeaders", "sendJson", "handleApi", "readRequestJson", "isLocalRequest"}:
            return "HTTP API surface", "This function belongs to the local HTTP/API layer. It either shapes responses, reads requests, or decides whether a request is allowed to perform local work."
        if "compile" in lower or "hdl" in lower or "vcd" in lower or "tool" in lower or "process" in lower or "source" in lower or "java" in lower or "cfamily" in lower:
            return "Compiler and IDE workspace", "This function supports the Compile Code workspace: detecting languages, writing files, finding tools, running compilers, parsing terminal output, or preparing simulation data."
        if "github" in lower or "source" in lower and "fetch" in lower or "portfolio" in lower and "context" in lower:
            return "Public source and AI context", "This function helps the assistant read public sources, GitHub repositories, local site files, or source snippets so answers can be based on actual portfolio evidence."
        if "ollama" in lower or "openai" in lower or "portfolioai" in lower or lower.endswith("ai") or "conversation" in lower or "answer" in lower:
            return "AI backend", "This function is part of the portfolio AI path. It prepares messages, reads model responses, maintains conversation shape, or creates a fallback answer."
        if "publish" in lower or "git" in lower or "remote" in lower or "branch" in lower or "credential" in lower or "target" in lower or "auth" in lower or "domain" in lower:
            return "Publishing and GitHub authorization", "This function protects and performs website publishing. It validates Git remotes, checks branches, caches authorization, stages approved files, or imports the target site."
        if "update" in lower or "release" in lower or "security" in lower or "version" in lower:
            return "Application update and reporting", "This function supports release checks, installer handoff, version comparison, or security/download reporting."
        if "safe" in lower or "resolveinside" in lower or "path" in lower or "file" in lower:
            return "Filesystem safety", "This helper prevents unsafe names and paths from escaping the intended workspace, portfolio mirror, or compile workspace."
        return "Backend utility", "This is a backend helper that normalizes data or supports a larger server operation."
    if "assistant" in lower:
        return "Portfolio AI frontend", "This function belongs to the public chatbot. It classifies questions, gathers context, renders answers, or manages the chat UI."
    if "search" in lower or "score" in lower or "query" in lower or "match" in lower or "filter" in lower:
        return "Search and filtering", "This function builds, scores, filters, highlights, or navigates search results so the public website behaves like a real searchable portfolio."
    if "section" in lower or "dialog" in lower or "route" in lower or "history" in lower or "path" in lower or "node" in lower:
        return "Project windows and navigation", "This function supports full-window project sections, nested subsections, browser history, back/forward behavior, or content tree traversal."
    if "rich" in lower or "render" in lower or "html" in lower or "inline" in lower or "math" in lower or "code" in lower:
        return "Rich content rendering", "This function turns saved builder content into public HTML: styled text, formulas, code blocks, images, files, links, and safe markup."
    if "profile" in lower or "site" in lower or "fun" in lower or "contact" in lower or "mail" in lower or "phone" in lower:
        return "Profile and site content", "This function renders or normalizes owner identity, front-page text, contact links, fun facts, or custom main-page sections."
    if "template" in lower or "responsive" in lower or "style" in lower or "color" in lower or "font" in lower or "crop" in lower:
        return "Appearance and responsive presentation", "This function applies templates, visual style, font/color/crop settings, or responsive layout decisions."
    if "download" in lower or "resource" in lower or "file" in lower or "link" in lower or "url" in lower:
        return "Resources and links", "This function decides whether content is a website link, local download, uploaded file, or resource card."
    if "project" in lower or "category" in lower:
        return "Project directory", "This function normalizes, flattens, filters, or renders projects and categories."
    return "Public runtime utility", "This helper supports the public site runtime by normalizing values or preparing data for another renderer."


def function_calls(func: FunctionDoc, all_names: set[str]) -> list[str]:
    calls: list[str] = []
    for name in sorted(all_names, key=lambda item: func.body.find(item)):
        if name == func.name:
            continue
        if re.search(rf"\b{name}\s*\(", func.body):
            calls.append(name)
        if len(calls) >= 8:
            break
    return calls


def return_description(func: FunctionDoc) -> str:
    body = func.body
    if re.search(r"\breturn\s+await\b", body):
        return "It returns the awaited result of another asynchronous operation, so its caller receives the completed value rather than a pending Promise."
    if re.search(r"\breturn\s+\{", body):
        return "It returns a structured object. That object is the contract the next layer uses instead of trying to interpret raw strings or side effects."
    if re.search(r"\breturn\s+\[", body):
        return "It returns an array. The caller usually iterates over that list to render entries, process files, or choose candidates."
    if re.search(r"\breturn\s+`", body) or re.search(r"\breturn\s+['\"]", body):
        return "It returns a string that is already normalized for display, command output, a URL, a CSS value, or a file/path label."
    if re.search(r"\breturn\s+", body):
        return "It returns a computed value used immediately by a caller. The important point is that the caller does not repeat this rule elsewhere."
    return "It mostly works through side effects, such as updating the DOM, writing files, sending an HTTP response, changing history, or mutating controlled runtime state."


def operation_description(func: FunctionDoc) -> str:
    signals = []
    body = func.body
    if "await " in body:
        signals.append("waits for asynchronous work")
    if "fetch(" in body:
        signals.append("calls a network resource")
    if "runProcess(" in body or "execFile" in body or "spawn(" in body:
        signals.append("runs an external command")
    if any(token in body for token in ["readFile", "writeFile", "mkdir", "rm(", "copyFile", "readdir"]):
        signals.append("reads or writes files")
    if "querySelector" in body or "innerHTML" in body or "classList" in body:
        signals.append("updates browser DOM")
    if "history." in body or "URLSearchParams" in body:
        signals.append("updates browser route/history")
    if "localStorage" in body:
        signals.append("uses browser local storage")
    if "setTimeout" in body:
        signals.append("uses a timer to avoid blocking or to wait for another process")
    if not signals:
        return "The body is mostly deterministic transformation logic: it reads its inputs, normalizes or scores them, and returns a value the caller can trust."
    return "Inside the body, it " + ", ".join(signals[:-1]) + (" and " if len(signals) > 1 else "") + signals[-1] + "."


def body_detail_notes(func: FunctionDoc) -> str:
    body = func.body
    notes: list[str] = []
    if "JSON.parse" in body:
        notes.append("parses JSON rather than treating incoming text as already trusted data")
    if "JSON.stringify" in body:
        notes.append("serializes structured data before sending or saving it")
    if "path.resolve" in body or "path.join" in body or "path.relative" in body:
        notes.append("uses Node path utilities so Windows path separators and relative paths are handled consistently")
    if "throw new Error" in body or "throw publishAccessError" in body:
        notes.append("throws explicit errors when a required safety or compatibility condition fails")
    if "response.writeHead" in body or "response.end" in body:
        notes.append("writes the HTTP response directly")
    if "document.createElement" in body:
        notes.append("creates DOM elements instead of relying only on static HTML")
    if "innerHTML" in body:
        notes.append("writes generated markup into the page, which is why escaping and sanitizing helpers around it matter")
    if "classList" in body:
        notes.append("toggles CSS classes to express UI state")
    if "addEventListener" in body:
        notes.append("attaches event handlers that turn user actions into behavior")
    if "URLSearchParams" in body:
        notes.append("reads or writes URL query parameters so browser history can represent the current view")
    if "localStorage" in body:
        notes.append("uses local browser storage for lightweight persistence")
    if "crypto.createHash" in body:
        notes.append("hashes sensitive scope information so authorization cache keys do not store raw secrets")
    if "fetch(" in body:
        notes.append("performs a fetch and therefore must handle network delay and failure")
    if "runProcess(" in body:
        notes.append("delegates command execution to runProcess so timeout and terminal capture remain consistent")
    if "await mkdir" in body or "mkdir(" in body:
        notes.append("creates folders before writing files so missing directories do not crash the workflow")
    if "await writeFile" in body or "writeFile(" in body:
        notes.append("writes data to disk as part of the operation")
    if "await readFile" in body or "readFile(" in body:
        notes.append("reads data from disk as part of the operation")
    if "history.pushState" in body or "history.replaceState" in body:
        notes.append("updates browser history so Back, Forward, and refresh can follow the same window state")
    if not notes:
        return ""
    return paragraph("; ".join(notes[:8]) + ".")


def failure_description(file_name: str, category: str, func: FunctionDoc) -> str:
    if "Publishing" in category:
        return paragraph(f"Without {func.name}, the publishing flow would lose one guardrail or one mechanical step. Depending on the function, that could mean pushing to the wrong branch, forgetting authorization, accepting a bad remote, staging unsafe files, or failing to import the latest website state.")
    if "Compiler" in category:
        return paragraph(f"Without {func.name}, the IDE-style compile workspace would become less reliable. The app might misidentify a language, lose a file role, fail to find a compiler, show confusing terminal output, skip waveform data, or run code in the wrong folder.")
    if "AI" in category or "assistant" in category:
        return paragraph(f"Without {func.name}, the assistant would either answer with less context, display unrelated links, lose conversation memory, fail to format messages cleanly, or expose model/provider details in the wrong layer.")
    if "Search" in category:
        return paragraph(f"Without {func.name}, search would become less like a search engine. Results could be missing, poorly ranked, not highlighted, or unable to navigate the visitor to the exact project, section, file, or phrase.")
    if "Project windows" in category:
        return paragraph(f"Without {func.name}, nested project navigation would break down. A visitor might click a section and see nothing, lose Back/Forward behavior, remain trapped in a dialog, or see content rendered in the wrong hierarchy.")
    if "Rich content" in category:
        return paragraph(f"Without {func.name}, builder-created content would not survive cleanly into the website. Formatting, images, formulas, links, code blocks, or sanitized HTML could be lost or rendered unsafely.")
    if "Profile" in category:
        return paragraph(f"Without {func.name}, owner-specific content would remain generic or malformed. Contact links, profile text, fun facts, and custom sections would not render with the intended public shape.")
    return paragraph(f"Without {func.name}, the specific rule it owns would disappear from the named workflow. Callers would either skip that behavior or reimplement it inconsistently, which is exactly how rendering, publishing, compiling, searching, or AI context drifts over time.")


def function_role_sentence(file_name: str, func: FunctionDoc, category: str, category_context: str) -> str:
    name = func.name
    lower = name.lower()
    subject = name
    if lower.startswith("normalize"):
        target = name[len("normalize") :] or "value"
        return paragraph(
            f"{subject} standardizes {target} data before another part of the app uses it. The point is not decoration; it prevents different callers from interpreting the same input differently. A normalizer usually accepts loose user, catalog, or file data and turns it into the one shape the rest of the subsystem expects."
        )
    if lower.startswith("render"):
        target = name[len("render") :] or "content"
        return paragraph(
            f"{subject} turns saved data for {target} into visible UI. It is a presentation function, but it also enforces public-site rules: empty content should not appear, rich content must be converted into safe HTML, and the generated markup must match the class names that styles.css expects."
        )
    if lower.startswith("fetch"):
        target = name[len("fetch") :] or "resource"
        return paragraph(
            f"{subject} retrieves {target} from outside the immediate function. It wraps network or repository access so callers do not need to know URL construction, headers, limits, or error behavior. This matters because external data is slow, failure-prone, and must be bounded before it is trusted."
        )
    if lower.startswith("read"):
        target = name[len("read") :] or "data"
        return paragraph(
            f"{subject} reads {target} from a request, file, cache, or generated artifact. It gives the caller parsed data rather than raw bytes or untrusted text, which keeps parsing rules centralized."
        )
    if lower.startswith("write"):
        target = name[len("write") :] or "data"
        return paragraph(
            f"{subject} writes {target} to a controlled destination. Write helpers matter because a wrong path, stale format, or missing directory can corrupt local drafts, compile workspaces, publishing state, or generated website files."
        )
    if lower.startswith("safe") or lower.startswith("validate") or lower.startswith("resolveinside"):
        return paragraph(
            f"{subject} is a guardrail function. It checks or transforms caller input before the system uses that input as a filename, path, remote, domain, credential, or public value. This keeps unsafe input from becoming filesystem access, Git access, or broken public links."
        )
    if lower.startswith("parse"):
        target = name[len("parse") :] or "input"
        return paragraph(
            f"{subject} interprets {target} text and returns structured meaning. Parsing functions are needed because the rest of the app should not repeatedly scan raw strings when it can work with a stable object, array, or normalized value."
        )
    if lower.startswith("detect") or lower.startswith("infer"):
        return paragraph(
            f"{subject} makes an educated classification from incomplete evidence. It looks at names, extensions, source text, repository state, or runtime state and returns the app's best decision so the next operation can choose the correct path."
        )
    if lower.startswith("score"):
        return paragraph(
            f"{subject} ranks a candidate against a question or search term. Scoring functions are what make search and AI source selection feel relevant instead of random."
        )
    if lower.startswith("collect") or lower.startswith("flatten"):
        return paragraph(
            f"{subject} gathers scattered nested data into a shape that another function can loop over. This is important because portfolio data is hierarchical while rendering, search, and AI context often need a flat list."
        )
    if lower.startswith("open"):
        return paragraph(
            f"{subject} changes UI state from a closed or hidden state into a visible working state. In this app, opening usually means finding the correct project or section data, preparing the window, rendering content, and synchronizing navigation."
        )
    if lower.startswith("close"):
        return paragraph(
            f"{subject} reverses an open window or route state. It does more than hide markup: it must also clean navigation state, restore the previous view when appropriate, and avoid trapping the visitor inside a dialog."
        )
    if lower.startswith("ensure"):
        return paragraph(
            f"{subject} creates or repairs something only when it is needed. This pattern avoids duplicate DOM nodes, duplicate repositories, duplicate handlers, or missing folders while still letting later code assume the resource exists."
        )
    if lower.startswith("sync"):
        return paragraph(
            f"{subject} reconciles two states that can drift apart. In this project, sync functions connect local drafts, route state, Git branches, publish mirrors, or frontend state to the current truth."
        )
    if lower.startswith("run"):
        return paragraph(
            f"{subject} executes an operation with side effects. It is separated from callers so process execution, Git execution, optional command behavior, or status collection remains consistent."
        )
    if lower.startswith("install"):
        return paragraph(
            f"{subject} installs or prepares external tooling. It belongs in the backend because installation touches the machine, not just the browser page."
        )
    if lower.startswith("assistant"):
        return paragraph(
            f"{subject} is one piece of the Ask My Portfolio conversation system. It helps decide what the user means, what evidence belongs in context, how answers should be displayed, or how the chat panel should behave."
        )
    if lower.startswith("project") or lower.startswith("category"):
        return paragraph(
            f"{subject} handles project-directory structure. It helps turn saved catalog data into visible categories, cards, filters, responsive layout, or searchable project text."
        )
    if lower.startswith("is") or lower.startswith("has") or lower.startswith("can"):
        return paragraph(
            f"{subject} is a decision predicate. It answers one yes/no question so the larger workflow can branch clearly instead of embedding the same condition in several places."
        )
    return paragraph(
        f"{subject} belongs to {category.lower()}. {category_context} In this role, it owns one named operation that later code depends on."
    )


def generated_function_paragraphs(file_name: str, func: FunctionDoc, all_names: set[str]) -> list[tuple[str, str]]:
    important = SERVER_IMPORTANT.get(func.name) if file_name == "server.mjs" else SCRIPT_IMPORTANT.get(func.name)
    category, category_context = classify_function(file_name, func)
    calls = function_calls(func, all_names)
    params = parse_parameters(func.params)
    paragraphs: list[tuple[str, str]] = []
    if important:
        paragraphs.extend(("Purpose and intuition", text) for text in important)
    else:
        paragraphs.append(("Purpose and intuition", function_role_sentence(file_name, func, category, category_context)))
    if params:
        parameter_text = " ".join(explain_parameter(param) for param in params)
    else:
        parameter_text = ""
    if parameter_text:
        paragraphs.append(("Parameters and data it receives", parameter_text))
    paragraphs.append(
        (
            "What it does",
            paragraph(
                f"{operation_description(func)} {return_description(func)}"
            ),
        )
    )
    details = body_detail_notes(func)
    if details:
        paragraphs.append(("Important body details", details))
    if calls:
        call_text = paragraph(
            f"It works with {', '.join(calls)}. Those calls show that {func.name} is not isolated; it is one piece of a larger workflow where helpers clean inputs, perform side effects, render results, or protect boundaries before the next step runs."
        )
    else:
        call_text = paragraph(
            f"No other named helper from this file is detected inside the body. That means {func.name} performs its rule directly from parameters, module state, standard APIs, or local calculations."
        )
    paragraphs.append(("What it calls or works with", call_text))
    paragraphs.append(("Why the file needs it", failure_description(file_name, category, func)))
    return paragraphs


SERVER_OVERVIEW = [
    (
        "Abstract",
        [
            paragraph(
                "server.mjs is the local nervous system of the OMB Portfolio Builder. It is not the public website and it is not the Electron shell. It is the private backend that runs on the owner's machine and gives the builder the powers a browser page should not have by itself: reading and writing local files, saving drafts, accepting uploads, running compilers, using Git, checking publishing authorization, importing a target site, launching installers, and relaying AI requests. The file is intentionally broad because the builder is local-first. Instead of depending on a hosted database and a remote application server, the desktop app starts this Node server beside the user-owned workspace and talks to it through local API endpoints."
            ),
            paragraph(
                "The most important idea is boundary control. The frontend can draw buttons and collect input, but it cannot safely decide where files are written, which Git repository receives a push, whether a command should run, or whether a request came from the local computer. server.mjs owns those decisions. It is the layer that turns a UI action into a guarded filesystem, compiler, Git, installer, or AI operation."
            ),
        ],
    ),
    (
        "Introduction",
        [
            paragraph(
                "When the builder starts, Electron loads a local page. That page needs catalog data, templates, upload handling, compile actions, publishing actions, and update checks. A static HTML page cannot perform those operations directly without either exposing dangerous privileges or running into browser security walls. server.mjs solves that by becoming a small local application server. It serves the builder files, exposes API endpoints under /api, and keeps privileged work in Node where the filesystem, process runner, and Git commands are available."
            ),
            paragraph(
                "A useful mental model is to treat this file like a workshop manager. The frontend is the workbench surface where the user sees projects, editors, buttons, and dialogs. server.mjs is the person behind the counter who knows where the tools are, where files belong, whether a customer has permission to ship something, and whether the requested operation is safe. If the frontend says 'save this PDF under this project,' server.mjs sanitizes the path and writes it under the correct docs folder. If the frontend says 'compile this C++ program,' server.mjs finds g++, creates a temporary run folder, captures terminal output, and returns a result object. If the frontend says 'apply to site,' server.mjs checks Git authorization first, then synchronizes publishable files and pushes them."
            ),
        ],
    ),
    (
        "Background Information",
        [
            paragraph(
                "The file uses built-in Node modules rather than a large web framework. createServer from node:http accepts requests, node:fs/promises handles files, node:child_process runs compilers and Git, node:path normalizes paths, node:crypto hashes authorization cache scopes, and node:os identifies the current machine. This keeps the backend predictable and package-light. The tradeoff is that routing and request handling are written manually. That is why handleApi is large: it plays the role that Express or another framework would normally play."
            ),
            paragraph(
                "There are three important roots. root is the builder workspace that contains the site files and local draft catalog. portfolioRoot is the publish mirror, which may be separate from root so generated website files can be pushed without mixing in builder-only files. compileRoot is where project-local code workspaces and temporary compiler run folders live. Most path-related functions exist to make sure a request stays inside one of these roots."
            ),
            paragraph(
                "The backend is also intentionally local-sensitive. Some endpoints are read-only, but write endpoints reject non-local callers. That is why functions such as isLocalRequest matter: the builder may be opened in a browser or accessed on a LAN during testing, but operations that change files, compile code, install tools, or publish should be performed only from the owner's computer."
            ),
        ],
    ),
    (
        "Purpose Of The File",
        [
            paragraph(
                "The purpose of server.mjs is to make the builder operational instead of decorative. Without it, template-preview.html and template-preview.js would still be able to render a screen, but they would not be able to reliably save drafts, upload project evidence, run compilers, authenticate GitHub targets, import an existing website, launch updates, or call private AI providers. The file therefore holds the application's real side effects."
            ),
            paragraph(
                "The file is also a protection layer. It validates filenames, normalizes repository URLs, keeps publish authentication scoped to a machine and target, prevents stageable publishing paths from wandering outside the publish mirror, blocks remote write requests, caps text fetched from external sources, and strips terminal output into UI-safe strings. Those decisions are not cosmetic. They prevent the portfolio builder from becoming a general-purpose remote command runner or a tool that accidentally pushes the wrong files."
            ),
        ],
    ),
    (
        "Primary Responsibilities",
        [
            paragraph(
                "The first responsibility is serving local files. The same server that handles /api requests also serves index.html, styles.css, script.js, template-preview.html, project catalogs, assets, and uploaded docs during local preview. The static-serving logic checks that the requested path stays inside the workspace before reading the file. That simple check is what prevents a URL from escaping into arbitrary folders on the machine."
            ),
            paragraph(
                "The second responsibility is compilation. Compile Code works because the backend can write source files to a project workspace, discover tools such as node, python, gcc, g++, javac, iverilog, vvp, and LTspice, run those tools with timeouts, collect stdout and stderr, cache successful builds, and return terminal output immediately to the frontend. The backend treats each language differently because JavaScript, Python, C, C++, Java, Verilog, SystemVerilog, LTspice, and HTML do not build the same way."
            ),
            paragraph(
                "The third responsibility is publishing. Publishing is not just 'write projects.json.' The backend must verify that the publish mirror is a Git repository, know its origin remote and branch, check compatibility with a static portfolio site, synchronize with the remote to avoid non-fast-forward errors, stage only approved publish paths, commit if there are changes, and push to the selected branch. This is why the publishing code is careful and verbose."
            ),
            paragraph(
                "The fourth responsibility is intelligence. For local AI, the backend can use an OpenAI API key or a local Ollama model. It also enriches portfolio context by reading public project links, same-site files, GitHub repositories, GitHub profiles, README files, and selected source files when the question calls for it. The browser cannot safely hold private API keys, so server.mjs keeps that path private."
            ),
            paragraph(
                "The fifth responsibility is application maintenance. The update functions check GitHub Releases, download an installer, write a PowerShell handoff script, close the running app, run the installer, and relaunch the updated executable. This is backend work because the browser cannot replace the running desktop app."
            ),
        ],
    ),
    (
        "Important Data Objects And Why They Matter",
        [
            paragraph(
                "compileLanguageProfiles is the backend's language map. It tells the compiler system what default filename to use, which file extensions belong to each language, what human label to show, which tools are required, and which Winget package can install the tool if automatic installation exists. Without this object, compileAndRunCode would need hardcoded language assumptions scattered throughout the file."
            ),
            paragraph(
                "compileToolCandidates is the tool-discovery map. It lists command names and likely Windows install paths for Git-adjacent and compiler-adjacent executables. findTool uses this object so the app can work even when a tool is installed but not perfectly exposed on PATH. The caches compileToolCache and compileToolVersionCache exist so repeated compiles do not waste time rediscovering the same executables."
            ),
            paragraph(
                "publishPaths is the publishing allowlist. It lists the site files and folders the backend is allowed to copy into the publish mirror and stage in Git. This is one of the most important safety objects in the file. Without an allowlist, an Apply to site operation could accidentally stage builder-only files, private drafts, temporary compile output, or unrelated local files."
            ),
            paragraph(
                "portfolioAiInstructions is the system behavior contract for the portfolio assistant. It tells the model how to distinguish general conversation, general engineering, general knowledge, and portfolio-specific questions. It also tells the model when to use context, when not to invent, and how to handle public GitHub source excerpts. That string is long because it is the policy boundary for the AI feature."
            ),
        ],
    ),
]


SCRIPT_OVERVIEW = [
    (
        "Abstract",
        [
            paragraph(
                "script.js is the public website runtime. It is the JavaScript that turns the static portfolio files into an interactive recruiter-facing site. The HTML page provides stable containers and the JSON catalog provides saved content, but script.js decides what appears, how projects are filtered, how rich text renders, how files become links, how project sections open into full-window views, how browser Back and Refresh interact with those views, how search behaves like a real search tool, and how the AI chat gathers useful context before asking the backend."
            ),
            paragraph(
                "The important distinction is that script.js is public and read-only. It should never expose builder-only editing controls. Its job is to present what has already been saved and parsed by the builder. That is why much of the file is about rendering, normalization, routing, search indexing, assistant context selection, and UI state rather than file writes."
            ),
        ],
    ),
    (
        "Introduction",
        [
            paragraph(
                "When a visitor opens the portfolio, index.html loads styles.css, electronics-search.js, and script.js. script.js immediately grabs the DOM elements it will control: the project grid, search input, filters, profile fields, resume area, contact area, AI assistant, and footer. Then it waits for portfolio data. The data can come from an embedded window.__PORTFOLIO_CATALOG__ object in previews or from projects.json on the published website. Once the catalog is loaded, script.js normalizes profile data, renders the hero and contact information, creates category filters, renders project cards, builds the search index, and restores any deep-linked project section."
            ),
            paragraph(
                "This file is the main reason the same portfolio can feel alive without requiring a database-backed website. It uses static files as the source of truth, then builds an application-like experience in the browser. That makes the public website easy to host on GitHub Pages or Cloudflare, while still giving visitors search, project windows, AI questions, and responsive navigation."
            ),
        ],
    ),
    (
        "Background Information",
        [
            paragraph(
                "The file works with a catalog generated by the builder parser. That catalog contains categories, projects, profile fields, site content, field styles, rich text blocks, fun facts, and optional site sections. The functions in script.js assume the catalog is already safe enough to display, but they still escape HTML where user-facing strings are inserted into the DOM. Rich text rendering is intentionally structured because portfolio content can include plain text, links, formulas, images, files, code blocks, captions, and nested project subsections."
            ),
            paragraph(
                "The website also has a windowing model. Project cards appear inside categories. A project section button opens a full-screen dialog. Inside that dialog, subsections can be clicked like nested folders. Browser history is updated with the project id, section index, and resource path so the browser Back button can move through the same hierarchy. This makes project content feel more like a navigable portfolio than a long flat webpage."
            ),
            paragraph(
                "Search is also broader than a normal text filter. The search index contains page sections, project names, category labels, tools, languages, rich text, files, URLs, captions, uploaded text files where fetchable, and electronics keyword expansions. The goal is that a recruiter can type an electronics phrase and land near the relevant artifact, even if the phrase appears inside a project section or file caption rather than on the top card."
            ),
            paragraph(
                "Because the website is static, script.js must do work that a database-backed application might normally ask a server to do. It must decide which categories exist, which projects belong to each category, which sections are empty, which files should render as downloads, which links should open externally, and which project paths should be represented in browser history. The price of static hosting is that the browser runtime has to be disciplined. The benefit is that the public site remains easy to deploy, cache, mirror, and serve from GitHub Pages or Cloudflare without a long-running custom web server."
            ),
            paragraph(
                "The file also sits between several other files. index.html gives it DOM anchors. styles.css gives visual meaning to the classes it writes. electronics-search.js can provide expanded electronics vocabulary. projects.json provides the portfolio catalog. The AI backend endpoint gives it remote intelligence when available. If one of those contracts changes, script.js is usually the first file that reveals the mismatch because rendering, search, or chat stops behaving correctly."
            ),
        ],
    ),
    (
        "Purpose Of The File",
        [
            paragraph(
                "The purpose of script.js is to convert portfolio data into a polished public experience. The file decides not only what to render, but also what not to render. Empty sections are hidden. Overview content appears as the first readable content when present. Links that look like websites open as websites rather than downloads. Local files become downloadable resources. Images render inline only when they are meant to be shown. AI links are only displayed when related to the question context."
            ),
            paragraph(
                "Without script.js, index.html would be mostly a static shell with placeholders. There would be no dynamic project directory, no generated categories, no rich project windows, no browser-aware section routing, no searchable uploaded file text, no AI context gathering, no responsive project rendering, and no dynamic profile/contact population."
            ),
        ],
    ),
    (
        "Primary Responsibilities",
        [
            paragraph(
                "The first responsibility is state hydration. loadProjectCatalog reads the portfolio catalog, normalizes the data, stores it in module-level variables, and triggers rendering. This is the startup path for the public website."
            ),
            paragraph(
                "The second responsibility is rendering. Functions such as renderProfile, renderSiteContent, renderFunFacts, renderSiteSections, projectCard, parsedSection, and parsedNodeContent turn structured catalog objects into actual HTML."
            ),
            paragraph(
                "The third responsibility is interaction. The dialog functions create the full-window project section view, the routing functions keep browser history synchronized, the search functions build and display selectable results, and the AI functions manage the chat panel."
            ),
            paragraph(
                "The fourth responsibility is interpretation. The assistant functions classify questions, gather relevant portfolio context, select public source links, decide whether a question is general or portfolio-specific, and prepare the request sent to the AI backend. This prevents the assistant from blindly dumping generic links for every question."
            ),
            paragraph(
                "The fifth responsibility is restraint. Public visitors should see only finished, parsed, saved portfolio output. The file therefore hides empty content, avoids exposing builder edit controls, limits which sources are shown with AI answers, and uses normalized link behavior so a website link does not become a fake downloadable file. This restraint is what separates the public website from the local builder application."
            ),
            paragraph(
                "The sixth responsibility is mobile and desktop continuity. The same project data must read cleanly on an iPhone and on a wide monitor. The responsive helper functions estimate project density and media presence, while the dialog and assistant functions keep windows and chat panels usable without letting one part of the page resize unrelated parts. script.js is therefore not only a renderer; it is also the public interaction policy for different devices."
            ),
        ],
    ),
    (
        "Important Data Objects And Why They Matter",
        [
            paragraph(
                "categories, projects, siteSections, profile, siteContent, funFacts, fieldStyles, and the rich variants of those fields are the hydrated catalog state. They are not hardcoded content. They are populated from projects.json or preview data, which is why the same script can run for Maurice's site and also for someone else's freshly installed builder output."
            ),
            paragraph(
                "searchableEntries is the in-memory search index. Every item in it has a title, type, context, text, optional project id, optional section index, optional URL, and normalized text. The search functions score and display these entries."
            ),
            paragraph(
                "assistantChatHistory and assistantSourceMap are the AI interface memory. assistantChatHistory keeps recent conversation so follow-up questions make sense. assistantSourceMap maps clickable source buttons in an answer back to portfolio search results so a visitor can jump to the relevant project or file."
            ),
            paragraph(
                "currentSearchResults and searchResultLimit are the active search session. They remember what the latest query produced and how many items the user wants to see. Without these values, pressing Enter, selecting a result, or changing the result limit would not know which result list is current."
            ),
            paragraph(
                "sectionRouteKey and sectionRouteParams are the browser-history contract for project windows. They keep the public URL connected to the currently opened project section and nested subsection. The values are small, but they are what let Refresh reopen the same project view and Back return to the previous view."
            ),
        ],
    ),
]


INDEX_OVERVIEW = [
    (
        "Abstract",
        [
            paragraph(
                "index.html is the public website shell. It is the first file a recruiter or visitor receives from the static host. It defines the document metadata, search-engine hints, social-card metadata, favicon and app icons, structured data, visible page sections, semantic landmarks, and the fixed DOM ids that script.js later fills with real portfolio content. The file is intentionally light on final personal content because the builder writes most custom data into projects.json and script.js renders it dynamically."
            ),
            paragraph(
                "The most important idea is that index.html is an interface contract. It tells the browser, search engines, social media previews, accessibility tools, styles.css, and script.js where things are and what each major region means. If script.js is the behavior layer, index.html is the skeleton that gives that behavior safe attachment points."
            ),
        ],
    ),
    (
        "Introduction",
        [
            paragraph(
                "The document begins with standard HTML setup: language, character set, viewport, description, robots instructions, author, and the AI endpoint metadata. These are not decorative tags. They affect how the page is indexed, how it scales on phones, how icons appear, and where the AI chat sends requests. After the metadata, index.html defines the visible page from top to bottom: header, hero, fun facts, builder download, summary metrics, projects, resume, dynamic sections, process, contact/AI, and footer."
            ),
            paragraph(
                "Most of the actual portfolio content is not typed directly in this file. That is deliberate. The builder is supposed to let users change names, images, resumes, project categories, links, and rich text without editing HTML manually. Therefore index.html supplies stable placeholders such as #hero-title, #project-grid, #resume, #dynamic-sections, #contact, and #ask-ai. script.js finds those placeholders and fills them from the catalog."
            ),
        ],
    ),
    (
        "Background Information",
        [
            paragraph(
                "A static portfolio can be hosted cheaply and reliably if the page can be built from HTML, CSS, JavaScript, and JSON. index.html is designed for that model. It avoids server-only assumptions and includes assets with relative paths. It loads electronics-search.js before script.js because script.js can use the electronics search vocabulary when building project search terms."
            ),
            paragraph(
                "The file also contains schema.org JSON-LD. That structured data gives search engines a machine-readable hint that the page is a ProfilePage about a Person. The builder later updates the visible content, but this baseline makes the page understandable even before dynamic content fully loads."
            ),
            paragraph(
                "The page uses semantic sections and labels because public portfolio viewers include recruiters, mobile users, browser accessibility tools, and search crawlers. A good static shell gives all of them landmarks. For example, the projects section is labeled by projects-title, the resume section is hidden until a resume URL exists, and the AI assistant has a labelled panel so it can be reached from navigation and screen readers."
            ),
            paragraph(
                "The file is intentionally not a finished personal biography. It is a reusable public shell. A fresh install should not hardcode Maurice-specific content, while Maurice's own built portfolio can still publish his profile through the catalog. That is why placeholder labels such as Portfolio Owner, Portfolio, and generic engineering copy exist here. The builder-generated catalog is the layer that specializes the site."
            ),
            paragraph(
                "The HTML also controls load order. CSS loads before the page renders so layout and colors are available immediately. electronics-search.js loads before script.js because script.js may call the electronics keyword helper while building the search index. script.js loads at the end of the body so the DOM elements already exist by the time it queries them. If script.js loaded too early without deferring, many document.querySelector calls would return null."
            ),
        ],
    ),
    (
        "Purpose Of The File",
        [
            paragraph(
                "The purpose of index.html is to guarantee that the public portfolio has a stable starting structure. Without it, script.js would have nowhere reliable to render projects, no header to navigate from, no AI panel to attach events to, no resume viewer to populate, no contact band to update, and no metadata to help browsers or search engines understand the page."
            ),
            paragraph(
                "It also separates permanent page structure from changeable portfolio content. The structure changes slowly; the content changes whenever the builder saves and publishes. That separation is what allows the builder to generate portfolio updates without rewriting the entire website architecture each time."
            ),
            paragraph(
                "The file is also the public accessibility foundation. Section headings, labels, aria attributes, hidden states, object fallback text, and semantic main/header/footer landmarks make the site more understandable to assistive technology. script.js can only enhance that foundation; it cannot rescue a shell that has no stable names or landmarks."
            ),
            paragraph(
                "Another purpose is failure tolerance. If projects.json fails to load, the page should still show a header, hero, search area, process section, and footer. If the resume is not configured, the resume section remains hidden. If the profile image is missing, the image elements stay hidden. This is what keeps a partially configured portfolio from looking broken."
            ),
        ],
    ),
    (
        "Primary Responsibilities",
        [
            paragraph(
                "The first responsibility is document identity. The head contains title, description, canonical URL, social tags, icons, AI endpoint, and structured data. These values determine how the page appears to crawlers, shared links, browser tabs, and mobile install surfaces."
            ),
            paragraph(
                "The second responsibility is visible layout landmarks. The body defines the header, main content regions, sections, and footer. styles.css controls how they look, and script.js controls most dynamic filling."
            ),
            paragraph(
                "The third responsibility is runtime hooks. Every id and class that script.js queries is part of a contract. Removing #project-grid, #project-search, #ask-ai, #contact, #resume, or #dynamic-sections would break the corresponding rendering logic."
            ),
            paragraph(
                "The fourth responsibility is graceful fallback. The page contains placeholder text and hidden sections so the website has a sensible shape even before projects.json finishes loading. If a profile image or resume does not exist, those areas remain hidden instead of showing broken controls."
            ),
            paragraph(
                "The fifth responsibility is hosting compatibility. The document uses ordinary relative paths and static assets. There is no server-side template language, no database call, and no build step required at request time. That is why the same page can be served from GitHub Pages, Cloudflare, or a local preview server."
            ),
            paragraph(
                "The sixth responsibility is public trust. The page exposes recruiter-facing areas only: projects, resume, process, contact, download builder, and AI assistant. It does not include local builder controls, Git publishing buttons, compiler controls, or editing widgets. That separation protects the private builder workflow from public visitors."
            ),
        ],
    ),
    (
        "Important Elements And Why They Matter",
        [
            paragraph(
                "The header is the main navigation surface. Its brand link returns to the top, the nav links jump to AI, projects, resume, process, and contact, and the avatar link can take the visitor to contact information. The brand image and text are later customized by script.js."
            ),
            paragraph(
                "The hero section is the first impression. It contains a hidden hero image, overlay shade, eyebrow, title, copy, project/resume buttons, GitHub button, and Ask AI jump link. These are all visible near the top because recruiter-facing content should not require deep scrolling."
            ),
            paragraph(
                "The project section is the core portfolio directory. It includes a search input, filter controls, and the empty project grid that script.js fills. The search input is deliberately placed in the section heading because search is a primary navigation method, not an afterthought."
            ),
            paragraph(
                "The contact band is last. It holds contact information and the AI assistant panel. The AI lives in this area because it is a helper for exploring the portfolio, while the contact section remains the final page region."
            ),
            paragraph(
                "The hidden attributes are part of the contract. They do not mean the feature is dead; they mean script.js will reveal the feature only when the catalog contains the required data. This prevents empty resume viewers, empty avatars, and profile images without sources from appearing to visitors."
            ),
            paragraph(
                "The IDs are also part of the contract. #project-grid, #project-search, #project-filters, #resume, #dynamic-sections, #contact, #ask-ai, #ai-assistant-form, #ai-assistant-input, and #ai-assistant-log are the anchors script.js expects. Renaming them would not only change HTML; it would break runtime behavior."
            ),
        ],
    ),
]


SERVER_GROUPS = [
    {
        "title": "Startup, roots, and response safety",
        "intro": paragraph("These functions and constants establish the backend's working environment. They define where files live, how JSON responses are sent, and what security headers are attached."),
        "items": [
            (
                "securityHeaders(extra)",
                paragraph("This function returns the baseline security headers added to local backend responses. The optional extra object lets a caller add or override headers without rewriting the common policy. Its purpose is to keep every response consistent: no MIME sniffing, strict referrer behavior, and browser permission limits. Without it, every endpoint would need to remember its own security headers, and missed endpoints could behave differently."),
            ),
            (
                "sendJson(response, status, data)",
                paragraph("This is the standard JSON response writer. It receives the Node response object, an HTTP status code, and any serializable data object. It writes headers, disables caching, stringifies the data with indentation, and closes the response. The function matters because every API endpoint should return JSON in the same shape and with the same cache rules. It is intentionally synchronous around response.end because once the data object is ready, the output should be immediate."),
            ),
            (
                "createServer request handler",
                paragraph("The final createServer block is the actual HTTP entry point. It first asks handleApi whether an /api request has been handled. If not, it treats the request as a static file request. The pathname is normalized, joined to root, and rejected if it escapes root. Then the file is read and sent with a MIME type from the types map. This is the simplest possible local server: API first, static files second, forbidden paths rejected, missing files return 404."),
            ),
        ],
    },
    {
        "title": "Language detection, filenames, and code formatting",
        "intro": paragraph("These helpers make Compile Code feel automatic. They normalize language names, infer languages from file names and source text, produce safe filenames, and apply simple beautification before code is saved or displayed."),
        "items": [
            (
                "normalizeCodeLanguage(value), languageFromFileName(fileName, code), and detectCodeLanguageFromSource(code, fileName)",
                paragraph("These functions form the language-detection chain. normalizeCodeLanguage turns aliases such as js, py, cpp, c++, sv, or system verilog into canonical ids. languageFromFileName prefers file extensions because .java, .sv, .py, and .cpp are strong evidence. detectCodeLanguageFromSource uses both filename and source clues, such as public class for Java or module syntax for HDL. The returned language id selects a compile profile, default filename, toolchain, highlighting rules, and build behavior."),
            ),
            (
                "safeCodeFileName(value, language)",
                paragraph("This function accepts a requested filename and a language, then returns a filename safe enough to write under the compile workspace. It strips dangerous path characters, keeps a language-appropriate extension, and falls back to the default filename from compileLanguageProfiles. Without it, a user-provided filename could accidentally create confusing paths or miss the extension needed by a compiler."),
            ),
            (
                "beautifyCode(code, language) and indentBraceCode(code)",
                paragraph("beautifyCode is a lightweight formatter used by the code editor. For brace-based languages it delegates to indentBraceCode, which adjusts indentation by watching braces and brackets. It is not a full compiler-grade formatter, but it gives pasted code a cleaner shape. The function exists because code that looks organized is easier to read before being appended to a project."),
            ),
        ],
    },
    {
        "title": "Tool discovery and process execution",
        "intro": paragraph("Compile Code, Git publishing, installers, and AI helpers all need controlled process execution. These functions find tools, run them safely, and return useful terminal text."),
        "items": [
            (
                "findExecutableUnder(folder, names, maxDepth) and findTool(toolName)",
                paragraph("findExecutableUnder recursively searches likely install folders for executable names. findTool uses compileToolCandidates first, then command lookup, then known package folders. The important parameter is toolName: it is the key into the candidate map and cache. The function is async because filesystem searches and command lookup can take time. It stores successful results in compileToolCache so repeated compiler runs do not lag."),
            ),
            (
                "runProcess(command, args, options)",
                paragraph("runProcess is the controlled command runner. command is the executable path or command name; args is the exact argument list; options can supply cwd, timeoutMs, env, and stdin input. The function spawns the process, collects stdout and stderr, kills the process on timeout, and resolves a structured object containing success state, output, exit code, elapsed time, and timeout status. Without this wrapper, every compiler and Git call would duplicate risky process code."),
            ),
            (
                "terminalLine, processTerminalText, replacePathReferences, and processTerminalTextWithPaths",
                paragraph("These helpers turn raw process results into readable terminal output. They combine stdout and stderr, include elapsed time and status, and replace long local paths with shorter display names. They do not compile anything themselves; they make compiler and Git results understandable in the builder UI."),
            ),
        ],
    },
    {
        "title": "Compile workspace and build orchestration",
        "intro": paragraph("This is the backend IDE layer. It turns editor state into real files, runs the correct language-specific toolchain, caches outputs, and returns terminal and waveform data."),
        "items": [
            (
                "saveCompileSource({ projectId, fileId, title, fileName, language, role, code, stdin })",
                paragraph("saveCompileSource writes the active source file and a metadata file under compileRoot. The object parameter is the frontend's description of the active file. projectId chooses the project workspace, fileId chooses the file folder, fileName chooses the source filename, language chooses the compile profile, role distinguishes HDL design/testbench files, code is the source text, and stdin is optional runtime input. The returned saved object tells the frontend exactly where the backend stored the source."),
            ),
            (
                "compileWorkspaceFilesFromPayload(payload, activeFileName, activeLanguage)",
                paragraph("This function normalizes the list of files sent by the frontend. It makes sure the active file is included, gives every file a safe id and filename, normalizes language and role, and removes duplicates. It matters because a project workspace is not just one file: C/C++ headers, Java classes, and HDL testbenches may all be needed together."),
            ),
            (
                "writeCompileWorkspaceSources(files, targetDir, options)",
                paragraph("This async function writes a normalized list of workspace files into a run or cache directory. It can filter by languages or extensions, and it returns written source records that include sourcePath and uniqueName. Many compilers need real files on disk rather than strings in memory, so this function is the bridge between editor state and compiler input."),
            ),
            (
                "compileAndRunCode(payload)",
                paragraph("compileAndRunCode is the central build engine. Its payload contains language, fileName, code, action, projectId, stdin, workspaceFiles, and forceRebuild. It detects the language, saves the active source, prepares run folders, selects the action, finds tools, runs language-specific compile/run/simulate logic, and returns ok, language, saved metadata, terminal text, and waveform data when available. It is async because every meaningful operation may touch disk or external processes. Without it, the Compile Code UI would be only a text editor with no real execution path."),
            ),
            (
                "installCompilerTools(language)",
                paragraph("installCompilerTools looks up the selected language profile and runs Winget packages when automatic installation is configured. Its input is the language id. Its output is a terminal-style result describing whether installation commands succeeded. It exists to reduce manual setup, but it intentionally returns a clear message for languages that do not have automatic installation configured."),
            ),
        ],
    },
    {
        "title": "C, C++, Java, and HDL build details",
        "intro": paragraph("These helpers customize compileAndRunCode for languages that need multiple files, binary artifacts, cached builds, testbenches, or waveform output."),
        "items": [
            (
                "cFamilyCompileProfile, cFamilyWorkspaceSources, cFamilyBinaryName, and cFamilyRunOutput",
                paragraph("These functions define how C and C++ behave. cFamilyCompileProfile chooses compiler flags, standard version, and whether a file is a header-only syntax check. cFamilyWorkspaceSources selects relevant .c/.cpp/.h/.hpp files. cFamilyBinaryName gives a safe executable name. cFamilyRunOutput formats runtime output. Together they let Compile Code distinguish syntax checking, building a workspace, and running a binary."),
            ),
            (
                "javaMainClassName(code, fileName)",
                paragraph("Java execution needs a class name, not only a filename. This function reads the code to find a public class first, then any class, then falls back to the filename. compileAndRunCode uses that class name to know which .class file to expect and what to pass to the java runtime."),
            ),
            (
                "hdlFilesFromPayload, inferCompileFileRole, normalizeCompileFileRole, hdlModuleNames, and hdlHasWaveDump",
                paragraph("These helpers make Verilog/SystemVerilog simulation sane. They identify design files versus testbenches, normalize the role labels, extract module names so a top module can be selected, and check whether a testbench contains $dumpfile and $dumpvars. Simulation needs stimulus and waveform dumping; these functions enforce that idea before vvp runs."),
            ),
            (
                "parseVcdScopeText(text, source), readHdlWaveform(cacheDir), and clearHdlWaveforms(cacheDir)",
                paragraph("parseVcdScopeText converts a VCD waveform file into structured signal data: timescale, max time, signal names, and value changes. readHdlWaveform finds a generated VCD file in the cache directory and parses it. clearHdlWaveforms removes stale VCD files before a fresh simulation run. These functions are why the builder can show a signal scope instead of only text output."),
            ),
        ],
    },
    {
        "title": "GitHub, public source, and AI context enrichment",
        "intro": paragraph("These functions let the assistant answer from project evidence, public GitHub links, public profile links, and fetched text without pretending to see private data."),
        "items": [
            (
                "sourceUrlAllowed, parseGitHubSourceUrl, fetchGitHubJson, and fetchLimitedText",
                paragraph("These helpers control external reads. sourceUrlAllowed rejects unsafe or unsupported URLs. parseGitHubSourceUrl understands GitHub profile and repository URLs. fetchGitHubJson calls GitHub's public API with appropriate headers. fetchLimitedText retrieves bounded text so a large public file cannot flood the AI context. Together they keep enrichment useful and limited."),
            ),
            (
                "fetchGitHubProfileSource, fetchGitHubRepositorySource, and fetchGitHubSourceText",
                paragraph("These functions are the GitHub-specific knowledge collectors. A profile source can yield profile metadata and repositories. A repository source can yield README text, repository metadata, and selected source files scored against the question. The selected snippets become sourceExcerpts that the AI can quote or explain. They exist so a question about code can reach public GitHub evidence rather than only project summaries."),
            ),
            (
                "enrichPortfolioContext(context)",
                paragraph("This function receives the context assembled by the frontend and expands it with fetched source excerpts when allowed. It reads sourceFetches, fetches safe public text, and returns a richer context object. The function is async because it may call GitHub, fetch same-site files, or read local files. Without it, the AI backend would only see the compact catalog context and would be weaker on uploaded files or public repositories."),
            ),
        ],
    },
    {
        "title": "Publishing target, authorization cache, and site push",
        "intro": paragraph("This group is the site-deployment guardrail. It validates the target, verifies access, caches successful authorization, imports from target, and applies generated site files."),
        "items": [
            (
                "validatePublishRemoteUrl, parseGitHubRemote, validateCredentialPair, and validateCustomDomain",
                paragraph("These functions validate user-provided publishing inputs. repository URLs must be plausible Git remotes, credentials must be paired when supplied, and custom domains must be simple valid domain names. They exist so configure/authenticate actions fail early with clear messages instead of half-writing a broken target."),
            ),
            (
                "resolveInsideRoot, resolveInsidePortfolioRoot, resolveInsideCompileRoot, and samePath",
                paragraph("These path helpers defend the workspace boundaries. They turn user-controlled segments into paths inside the intended root and prevent path traversal. They are boring by design, but without them uploads, compile files, and publish sync could write outside the intended folders."),
            ),
            (
                "configurePublishTarget(options) and authenticateGitHubForTarget(options)",
                paragraph("configurePublishTarget sets up the publish repository, remote, branch, optional custom domain, and stored credentials. authenticateGitHubForTarget goes further: it temporarily applies the target, checks Git availability, stores credentials if provided, synchronizes the remote branch, verifies write access, writes the auth cache, and only then keeps the target. This distinction matters: target setup should not be considered complete until write access is proven."),
            ),
            (
                "publishAuthCacheScope, writePublishAuthCache, publishAuthCacheIsFresh, and assertPublishAccess",
                paragraph("These functions control the 'do not ask every time' behavior. The cache stores remote, branch, repository, machine/user scope, checked time, expiration, and extended trust history. assertPublishAccess first validates repository shape, then reuses a fresh cache if the exact target and machine scope match; otherwise it performs a new write check. Without this group, the app would either annoy the owner constantly or trust stale authorization too broadly."),
            ),
            (
                "syncFromPublishTarget()",
                paragraph("syncFromPublishTarget authenticates, clones the selected branch into a temporary folder, checks for a compatible projects.json, backs up local compatible files, copies imported site assets into the builder workspace, updates the local draft, and syncs publish files. This is the 'load from target' behavior for moving to another machine or restoring from the live repository."),
            ),
            (
                "syncPortfolioPublishFiles(options) and publishSiteChanges(publishAccess)",
                paragraph("syncPortfolioPublishFiles copies only allowed publishPaths from the builder workspace to the publish mirror. publishSiteChanges then synchronizes the branch, bumps asset versions, stages allowed paths, commits changes when needed, and pushes. This pair turns local saved portfolio output into the public website while keeping builder-only files out of the publish target."),
            ),
        ],
    },
    {
        "title": "Updates, security reporting, AI, and API routing",
        "intro": paragraph("These functions make the backend responsible for application maintenance and the public assistant endpoint used during local preview."),
        "items": [
            (
                "getUpdateInfo(), getBuilderReleaseDownloadReport(), safeUpdateFileSegment(), powershellSingleQuoted(), and downloadAndLaunchAppUpdate()",
                paragraph("getUpdateInfo reads release information and decides whether a newer builder exists. downloadAndLaunchAppUpdate downloads the latest installer, writes a PowerShell handoff script, closes running builder processes, runs the installer, and relaunches the app. The helper functions sanitize version/file fragments and PowerShell strings so generated scripts are safer. This group exists because a running desktop app cannot replace itself like a normal webpage refresh."),
            ),
            (
                "getSecurityReport()",
                paragraph("getSecurityReport summarizes what the builder can and cannot know about visitors, downloads, public site headers, local-only protections, and authentication caching. It is intentionally honest: a static GitHub Pages site cannot identify every visitor without an analytics backend."),
            ),
            (
                "callOllamaPortfolioAi, extractOpenAiText, extractOllamaText, ruleBasedConversationAnswer, and handlePortfolioAi",
                paragraph("These functions make the local AI endpoint. handlePortfolioAi reads the question, conversation, intent, and context; enriches the context; chooses OpenAI if an API key exists; otherwise tries Ollama; otherwise returns a clear unavailable response. The extractor helpers normalize provider response shapes. ruleBasedConversationAnswer covers simple local fallback conversation. This keeps AI secrets on the backend side and gives the frontend one endpoint to call."),
            ),
            (
                "handleApi(request, response, url)",
                paragraph("handleApi is the router for every backend endpoint. It handles read endpoints such as catalog, templates, system check, update info, security report, and compiler tools. It rejects non-local write requests, then handles write operations such as AI, code beautify, code save, code compile, tool install, Git authentication, load from target, save draft, apply to site, and upload. It returns true when it handled a route so the static file server does not also try to serve it."),
            ),
        ],
    },
]


SCRIPT_GROUPS = [
    {
        "title": "Page state, profile, hero, and fun facts",
        "intro": paragraph("This group turns catalog-level identity and front-page content into visible website content."),
        "items": [
            (
                "normalize(value), normalizeFunFacts(value), and renderFunFacts()",
                paragraph("normalize gives the rest of the file a consistent lowercase comparison string. normalizeFunFacts accepts either an array or newline-separated text, trims it, removes empty lines, and caps the visible list. renderFunFacts decides whether the fun facts callout is hidden, then renders either rich blocks or plain short facts. Without this group, top-page personal statements would be inconsistent and could show empty containers."),
            ),
            (
                "normalizeSiteContent, normalizeProfile, renderSiteContent, and renderProfile",
                paragraph("These functions are the identity hydrators. normalizeSiteContent and normalizeProfile merge catalog data with safe defaults. renderSiteContent updates the hero eyebrow, title, copy, and brand text. renderProfile updates the contact band, links, email/phone behavior, resume section, hero image, brand image, profile photo, and footer owner. These functions make the same HTML shell become a personalized portfolio."),
            ),
            (
                "mailComposeLink(email) and phoneLink(phone)",
                paragraph("These two helpers convert contact values into browser actions. mailComposeLink creates a Gmail compose URL with the recipient filled in. phoneLink creates a tel: URL so phones and desktop calling apps can offer a call. They are separate functions because email and phone values must be encoded differently."),
            ),
        ],
    },
    {
        "title": "Dialog dragging, resizing, and window feel",
        "intro": paragraph("These functions manage draggable section dialogs on the public site. They turn a modal into a controlled full-window or movable document-style view."),
        "items": [
            (
                "clampSectionDialogPosition, anchorSectionDialog, and anchorSectionDialogForResize",
                paragraph("clampSectionDialogPosition keeps a dialog within the viewport. anchorSectionDialog converts the browser's current layout position into explicit left/top coordinates so dragging can begin smoothly. anchorSectionDialogForResize does the same before resizing. Their inputs are coordinates and the dialog element. Their side effect is CSS positioning."),
            ),
            (
                "beginSectionDialogDrag, moveSectionDialogDrag, endSectionDialogDrag, beginSectionDialogResize, moveSectionDialogResize, and endSectionDialogResize",
                paragraph("These functions form pointer-state machines. The begin functions remember the starting mouse/touch point and dialog rectangle. The move functions calculate deltas and update style. The end functions clear active state. They are separate because browser pointer events arrive over time, so the file must remember what operation is active between events."),
            ),
            (
                "updateSectionDialogMinimize, toggleSectionDialogMinimized, and enableSectionDialogDrag",
                paragraph("These functions keep the dialog's minimized state and interaction affordances in sync. The minimized state changes classes and visible behavior; enableSectionDialogDrag wires the required event handlers once. Without them, the section window would feel like a static modal instead of a controllable portfolio window."),
            ),
        ],
    },
    {
        "title": "Rich content, links, math, code, and images",
        "intro": paragraph("This group renders the mixed content created by the builder: text, formulas, links, images, captions, file downloads, and syntax-highlighted code."),
        "items": [
            (
                "escapeHtml, renderInlineMath, renderMultilineInlineText, sanitizeRichInlineHtml, and linkifyRichTextNodes",
                paragraph("These functions protect and enrich displayed text. escapeHtml prevents raw text from becoming HTML. renderInlineMath recognizes formula-like fragments and renders them cleanly. renderMultilineInlineText preserves intentional line breaks. sanitizeRichInlineHtml removes unsafe tags while keeping allowed inline formatting. linkifyRichTextNodes turns recognized URLs into links. Together they keep rich text useful without trusting arbitrary pasted HTML."),
            ),
            (
                "normalizeLinkTarget, looksLikeBareWebAddress, looksLikeWebOrContactText, isWebsiteLinkItem, linkAttributes, downloadAttribute, and resourceLink",
                paragraph("These functions decide whether a string is a web link, contact link, local download, or empty planned asset. normalizeLinkTarget is the main converter; linkAttributes adds target and rel for real web links; downloadAttribute adds download only for local file assets; resourceLink builds the final anchor or muted placeholder. This is the group that prevents LinkedIn or GitHub links from being treated like files."),
            ),
            (
                "normalizeCodeLanguage, detectCodeLanguage, tokenizedCodeHtml, and renderRichCodeBlock",
                paragraph("These functions make saved code look like code instead of ordinary paragraph text. detectCodeLanguage uses aliases and source clues. tokenizedCodeHtml highlights common language tokens. renderRichCodeBlock wraps the result in a presentable block with language metadata. The functions do not compile code; they render code evidence for public viewing."),
            ),
            (
                "richImageCropStyle, richImageDownloadLink, cleanRichImageTitle, renderRichContent, and renderRichFieldContent",
                paragraph("These functions render rich editor blocks from the builder. Images may have crop settings, captions, downloadable behavior, or visible inline behavior. renderRichContent loops through saved blocks and dispatches to text, image, formula, code, or file rendering. renderRichFieldContent applies the same idea to profile/front-page fields. This is where builder-created rich content becomes portfolio-visible content."),
            ),
        ],
    },
    {
        "title": "Project templates, categories, and responsive project cards",
        "intro": paragraph("This group maps saved project appearance and project data into project cards that adapt to the visitor's device."),
        "items": [
            (
                "canonicalTemplateId, projectTemplateId, projectTemplateClass, projectTemplateStyle, and applyProjectTemplateToElement",
                paragraph("These functions treat templates as appearance choices, not content templates. canonicalTemplateId handles legacy template names. projectTemplateId extracts the active appearance. projectTemplateClass and projectTemplateStyle turn that choice into CSS classes and variables. applyProjectTemplateToElement applies those variables to project cards or section windows. This keeps project visuals consistent between card view and opened section view."),
            ),
            (
                "hydrateProjectCategories, normalizeCategory, setActiveFilter, bindFilterButtons, renderCategoryFilters, categorySection, and projectCard",
                paragraph("These functions build the project directory. hydrateProjectCategories combines saved categories with projects so empty and custom categories behave correctly. renderCategoryFilters creates filter buttons. categorySection renders each category band. projectCard renders either the modern parsed project model or the older legacy project model. The result is a page organized by categories without hardcoding category names."),
            ),
            (
                "inferResponsiveProfile and projectResponsiveProfile",
                paragraph("These functions read the project structure and estimate whether the project is simple, balanced, or dense. They count visible sections, nested depth, media presence, and child nodes. The returned profile gives CSS guidance such as content width, card layout, media max width, and touch targets. This is why the website can render projects differently on phones and PCs without Codex manually rearranging each project."),
            ),
        ],
    },
    {
        "title": "Search engine behavior",
        "intro": paragraph("This group turns portfolio content and project evidence into a local searchable index with scoring, highlighting, and result navigation."),
        "items": [
            (
                "flattenSearchText, richTextTerms, addSearchEntry, searchSnippet, entryScore, and searchResultsFor",
                paragraph("These functions define a search entry and how it ranks. flattenSearchText collapses arrays, objects, and strings into searchable text. richTextTerms extracts text-like values from rich blocks. addSearchEntry stores normalized text. entryScore rewards exact title matches, starts-with matches, context hits, type hits, word hits, and useful kinds such as files or sections. searchResultsFor scores, sorts, and de-duplicates entries. The user experiences this as a real search dropdown rather than a simple project-name filter."),
            ),
            (
                "fileIsBrowserSearchable and indexSearchableFileText",
                paragraph("fileIsBrowserSearchable decides which uploaded or linked files can be fetched as text in the browser. indexSearchableFileText fetches those files, appends a bounded amount of file text to the entry, and refreshes the dropdown. This allows code, logs, CSV, Markdown, Verilog, and similar files to become searchable by content when the browser can access them."),
            ),
            (
                "addProjectSearchEntries, addSiteSectionSearchEntries, rebuildSearchIndex, renderSearchResults, updateSearchDropdown, and goToSearchResult",
                paragraph("These functions build and use the full index. Projects contribute titles, categories, tools, languages, sections, files, rich text, and electronics vocabulary. Site sections contribute profile/resume/contact/process content. renderSearchResults shows the selectable list; updateSearchDropdown highlights visible text; goToSearchResult scrolls, opens a section window, or opens a link depending on result kind."),
            ),
        ],
    },
    {
        "title": "Portfolio AI context and conversation",
        "intro": paragraph("This group decides what the public AI assistant should know for a question and how the chat should behave."),
        "items": [
            (
                "assistantQuestionIntent and its classifier helpers",
                paragraph("assistantQuestionIntent uses helpers such as assistantQuestionIsCasual, assistantQuestionHasPortfolioIntent, assistantQuestionLooksConceptual, and assistantQuestionIsEngineeringRelated. The output can be general_conversation, general_engineering, general_knowledge, or portfolio_specific. That intent decides whether the assistant should answer generally or gather portfolio context first. This prevents a question like 'what is embedded systems' from immediately dumping Maurice's project links."),
            ),
            (
                "assistantContextForQuestion, assistantProjectEvidence, assistantPublicProfileLinks, assistantSourcesForDisplay",
                paragraph("These functions assemble the request context. They identify named projects, collect matching search results, gather public profile links, include project evidence, and choose which source buttons should display. assistantSourcesForDisplay is deliberately strict so the assistant does not append random generic links."),
            ),
            (
                "assistantGeneralEngineeringAnswer and assistantLocalAnswer",
                paragraph("These provide local fallback intelligence. assistantGeneralEngineeringAnswer contains direct explanations for embedded systems, op amps, VCO/PWM, FPGA/ASIC, STM32, PCB, and power topics. assistantLocalAnswer handles greetings, identity questions, catalog-loading cases, no-match cases, and summary answers from local search results. This keeps the assistant conversational even when the backend AI is unavailable."),
            ),
            (
                "askRemoteAssistant and answerAssistantQuestion",
                paragraph("askRemoteAssistant sends question, context, conversation, intent, and web-search permission to the backend endpoint with a timeout. answerAssistantQuestion is the UI orchestrator: it appends the user message, remembers conversation, shows a pending assistant message, decides intent, gathers sources, calls the remote assistant, falls back locally if needed, replaces the pending answer, remembers the assistant response, and restores input focus. It is async because network AI calls and source fetches take time."),
            ),
            (
                "appendAssistantMessage, replaceAssistantMessage, renderAssistantAnswerContent, clearAssistantChat, and updateAssistantPanelGrowth",
                paragraph("These functions are the chat interface. renderAssistantAnswerContent converts ChatGPT-like plain text into paragraphs, bullets, and code blocks. append/replace message functions update the DOM. clearAssistantChat resets history and source maps. updateAssistantPanelGrowth grows the AI console upward on desktop without resizing the contact panel. This group makes the assistant feel like a clean chatbot instead of a static Q&A box."),
            ),
        ],
    },
    {
        "title": "Parsed project windows and browser routing",
        "intro": paragraph("This group turns nested project sections into clickable windows and makes browser Back/Forward participate."),
        "items": [
            (
                "nodeHasRenderableContent, sectionHasRenderableContent, nodeOverviewDetails, parsedNodeContent, parsedSectionContent, parsedSection, and projectCard",
                paragraph("These functions decide what exists on the public site. Empty nodes are not rendered. Overview content is shown automatically and collapsibly when it exists. Child sections become clickable cards. Files become inline file links. projectCard wraps all of that into the category card. This group enforces the rule that only valid created content appears."),
            ),
            (
                "pathToString, pathFromString, nodeAtPath, nodeChildren, sectionRouteState, sectionRouteUrl, sectionRouteFromLocation, and sameSectionRoute",
                paragraph("These are routing helpers for nested project windows. A path such as 0.2.1 means 'walk through child indexes to reach a subsection.' The route functions translate project id, section index, and resource path into URL/search/history state and back. They are necessary because the browser Back button must know which project window is currently open."),
            ),
            (
                "ensureSectionDialog, openParsedSection, closeSectionDialog, closeOrStepBackSectionDialog, applySectionRoute, and restoreSectionRouteAfterCatalogLoad",
                paragraph("ensureSectionDialog creates the reusable dialog. openParsedSection finds the project/section/node, applies the project appearance template, updates route state, writes the title, renders content, and opens the full-window dialog. close/step-back functions decide whether to go up one nested level or close to the main page. applySectionRoute and restoreSectionRouteAfterCatalogLoad connect this behavior to browser history and page refresh."),
            ),
            (
                "loadProjectCatalog",
                paragraph("loadProjectCatalog is the startup pipeline. It loads either embedded preview data or projects.json, stores catalog pieces in state, renders profile/site/facts/sections/filters/projects, rebuilds search, updates the search dropdown, and restores deep-linked section windows. Without this function, the static HTML shell would never become the actual portfolio."),
            ),
        ],
    },
]


INDEX_GROUPS = [
    {
        "title": "Head metadata and search identity",
        "intro": paragraph("The head makes the page understandable before JavaScript renders anything."),
        "items": [
            (
                "charset, viewport, description, robots, author, and canonical link",
                paragraph("The charset tells the browser to interpret the file as UTF-8. The viewport makes the page scale correctly on phones. The description gives search engines a concise page summary. robots tells crawlers that indexing is allowed. author is a generic placeholder updated by published content or metadata later. The canonical link is empty in the template because the final domain can vary."),
            ),
            (
                "portfolio-ai-endpoint meta tag",
                paragraph("This meta tag gives script.js a backend URL for the portfolio AI assistant. assistantEndpoint reads it before falling back to /api/portfolio-ai. Keeping the endpoint in metadata lets the same JavaScript bundle run locally and publicly while pointing to different AI backends."),
            ),
            (
                "Open Graph and Twitter card tags",
                paragraph("These tags control how the portfolio looks when shared on social platforms or messaging apps. The defaults are generic because the builder can later fill real owner, title, image, and description values."),
            ),
            (
                "favicon and touch-icon links",
                paragraph("These links tell browsers and mobile operating systems which icons to show in tabs, bookmarks, install prompts, and home-screen shortcuts. They point at the OMB asset set so the portfolio has a recognizable identity."),
            ),
            (
                "ProfilePage JSON-LD script",
                paragraph("This structured data gives search engines a machine-readable summary of the page as a profile page about a person. The values are conservative placeholders. The important part is the shape: Person, name, image, URL, sameAs, and knowsAbout. It gives crawlers context even though the visible profile is populated dynamically."),
            ),
        ],
    },
    {
        "title": "Header and navigation contract",
        "intro": paragraph("The header gives visitors immediate navigation and gives script.js stable places to insert brand/profile visuals."),
        "items": [
            (
                "header#top and brand link",
                paragraph("The header has id top so Back to top links can target it. The brand anchor points to #top and contains the brand image, OMB engraving, title, and subtitle. script.js later replaces brand text and image from profile data. If these classes disappeared, renderProfile would lose its brand hooks."),
            ),
            (
                "site-nav links",
                paragraph("The navigation links jump to Ask AI, Projects, Resume, Process, and Contact. These anchors match section ids later in the file. They are simple hash links so the page works on static hosting and does not need a router."),
            ),
            (
                "header-avatar",
                paragraph("The avatar link starts hidden because not every portfolio has a profile image. renderProfile reveals it and sets the image when profile data provides one. Its href points to contact so the image acts as a fast route to contact information."),
            ),
        ],
    },
    {
        "title": "Hero, fun facts, download, and metrics",
        "intro": paragraph("These are top-of-page sections that help a visitor orient quickly."),
        "items": [
            (
                "hero section",
                paragraph("The hero contains a hidden image, shade overlay, eyebrow, h1 title, copy, action buttons, optional GitHub button, and Ask AI jump link. script.js updates the text and images. The section is labelled by hero-title, which improves accessibility and gives the document a clear main heading."),
            ),
            (
                "fun-facts-callout",
                paragraph("This hidden section becomes visible only when the catalog includes fun facts or rich fun fact blocks. It sits near the top so personality appears before deep project scrolling."),
            ),
            (
                "builder-download-section",
                paragraph("This section exposes the latest Windows installer link. It is public because visitors may want to clone/install the builder app. The href points to the stable latest release asset. The download attribute suggests download behavior, but the browser and GitHub ultimately decide how the file is handled."),
            ),
            (
                "summary-band metrics",
                paragraph("The metrics display project count, category count, and artifact promise. script.js updates the first two after loading projects.json. The third is static because the builder portfolio consistently supports source, diagrams, and reports."),
            ),
        ],
    },
    {
        "title": "Projects, resume, dynamic sections, process, contact, and AI",
        "intro": paragraph("These sections are the public portfolio body. Most content is populated dynamically."),
        "items": [
            (
                "projects section",
                paragraph("This is the main project directory. It has a heading, search input, filter bar, and empty project-grid. script.js builds filter buttons, project categories, project cards, parsed section buttons, search results, and nested section windows from this area."),
            ),
            (
                "resume section",
                paragraph("The resume section starts hidden because a fresh portfolio may not have a resume. renderProfile reveals it, sets the PDF object data, and updates Open/Download links when profile.resumeUrl exists. The object element gives an in-page PDF viewer when the browser supports it."),
            ),
            (
                "dynamic-sections container",
                paragraph("This div is where user-created main portfolio sections render. The builder can add personal life, professional profile, social links, hobbies, or other lighter sections without editing index.html."),
            ),
            (
                "process section",
                paragraph("The process section is a stable explanatory band describing problem/constraints, build artifacts, and results/reflection. It gives the page a professional engineering narrative even before specific custom sections are added."),
            ),
            (
                "contact-band and ask-ai panel",
                paragraph("The contact band is intentionally last. It holds profile photo, contact details, external links, and the AI assistant. The AI form, log, input, status, and clear button have fixed ids because script.js attaches the chatbot behavior directly to them."),
            ),
            (
                "footer and script loading",
                paragraph("The footer provides Back to top and copyright year. The scripts load electronics-search.js before script.js so the search vocabulary is available when script.js builds the index. Version query strings help browsers fetch updated CSS and JS after publishing."),
            ),
        ],
    },
]


GUIDES = [
    {
        "file": "server.mjs",
        "title": "Curated File Guide: server.mjs",
        "subtitle": "The local backend that saves drafts, runs compilers, checks GitHub publishing access, imports/pushes site files, handles updates, and relays AI.",
        "overview": SERVER_OVERVIEW,
        "groups": SERVER_GROUPS,
    },
    {
        "file": "script.js",
        "title": "Curated File Guide: script.js",
        "subtitle": "The public website runtime that renders the portfolio, search, project windows, browser history, rich content, links, and AI chat behavior.",
        "overview": SCRIPT_OVERVIEW,
        "groups": SCRIPT_GROUPS,
    },
    {
        "file": "index.html",
        "title": "Curated File Guide: index.html",
        "subtitle": "The public portfolio shell: metadata, semantic page structure, runtime hooks, visible sections, and script/style entry points.",
        "overview": INDEX_OVERVIEW,
        "groups": INDEX_GROUPS,
    },
]


def setup_document(title: str, subtitle: str) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)
    styles["Normal"].paragraph_format.first_line_indent = Inches(0.2)
    styles["Normal"].paragraph_format.line_spacing = 1.12
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Heading 1"].font.name = "Arial"
    styles["Heading 1"].font.size = Pt(18)
    styles["Heading 1"].font.color.rgb = RGBColor(10, 49, 80)
    styles["Heading 1"].paragraph_format.first_line_indent = Inches(0)
    styles["Heading 1"].paragraph_format.space_before = Pt(14)
    styles["Heading 1"].paragraph_format.space_after = Pt(8)
    styles["Heading 2"].font.name = "Arial"
    styles["Heading 2"].font.size = Pt(14)
    styles["Heading 2"].font.color.rgb = RGBColor(13, 83, 122)
    styles["Heading 2"].paragraph_format.first_line_indent = Inches(0)
    styles["Heading 2"].paragraph_format.space_before = Pt(10)
    styles["Heading 2"].paragraph_format.space_after = Pt(5)
    styles["Heading 3"].font.name = "Arial"
    styles["Heading 3"].font.size = Pt(11)
    styles["Heading 3"].font.color.rgb = RGBColor(31, 41, 55)
    styles["Heading 3"].paragraph_format.first_line_indent = Inches(0)
    styles["Heading 3"].paragraph_format.space_before = Pt(8)
    styles["Heading 3"].paragraph_format.space_after = Pt(3)
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_p.paragraph_format.first_line_indent = Inches(0)
    run = title_p.add_run(title)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(8, 35, 61)
    subtitle_p = doc.add_paragraph()
    subtitle_p.paragraph_format.first_line_indent = Inches(0)
    subtitle_run = subtitle_p.add_run(subtitle)
    subtitle_run.italic = True
    subtitle_run.font.color.rgb = RGBColor(71, 85, 105)
    return doc


def add_body_paragraph(doc: Document, text: str, lead: str = "", italic: bool = False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(0.2)
    if lead:
        lead_run = p.add_run(lead)
        lead_run.bold = True
        lead_run.font.color.rgb = RGBColor(15, 82, 110)
    run = p.add_run(text)
    run.italic = italic


def add_analogy_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(0)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    set_paragraph_shading(p, "EFF6FF")
    lead = p.add_run("Analogy: ")
    lead.bold = True
    lead.font.color.rgb = RGBColor(30, 64, 175)
    body = p.add_run(text)
    body.italic = True
    body.font.color.rgb = RGBColor(30, 41, 59)


def add_overview(doc: Document, guide: dict) -> None:
    doc.add_heading("Before The Code: File Overview", level=1)
    for heading, paragraphs in guide["overview"]:
        doc.add_heading(heading, level=2)
        for text in paragraphs:
            doc.add_paragraph(text)
    doc.add_page_break()


def add_server_overview_chapter(doc: Document, guide: dict) -> None:
    doc.add_heading("Chapter 1: Server Abstract, Background, And Design Purpose", level=1)
    add_body_paragraph(
        doc,
        paragraph(
            "This chapter explains server.mjs as a complete backend file before isolating any single function. The useful first question is not 'what does line one do?' but 'what job does this file perform inside the whole application?' The answer is that server.mjs is the private local backend for the OMB Portfolio Builder. It receives browser requests from the builder UI, performs privileged work on the owner's machine, and returns structured results that the UI can display."
        ),
    )
    add_body_paragraph(
        doc,
        paragraph(
            "The file exists because a portfolio builder needs powers that a normal webpage should not have. A browser page can collect a project title, display an editor, and send a request. It should not directly write arbitrary files, run compilers, launch installers, push to GitHub, or store publishing authorization. server.mjs becomes the controlled layer between the friendly interface and those dangerous machine-level operations."
        ),
    )
    add_analogy_paragraph(
        doc,
        "Think of the builder UI as the front desk of a hardware lab. The visitor can ask for a test, a file save, a compile, or a website push. server.mjs is the lab technician behind the counter. It knows which cabinet the files belong in, which tools are allowed, which requests are local, and which operations need authentication before anything public changes.",
    )
    for heading, paragraphs in guide["overview"]:
        doc.add_heading(heading, level=2)
        for text in paragraphs:
            add_body_paragraph(doc, text)
    doc.add_page_break()


def add_server_syntax_chapter(doc: Document) -> None:
    doc.add_heading("Chapter 2: JavaScript And Node Syntax Foundations", level=1)
    add_body_paragraph(
        doc,
        paragraph(
            "This chapter is intentionally general. It explains the JavaScript and Node ideas needed to read the later server chapters, without tying the explanation to any one server.mjs function. Once these ideas are in place, the later chapters can focus on behavior instead of repeating syntax notes."
        ),
    )
    doc.add_heading("2.1 Modules And Toolboxes", level=2)
    add_body_paragraph(
        doc,
        paragraph(
            "Modern Node files can use ES module import statements. An import is a declaration that says, 'this file needs a tool that lives somewhere else.' The imported tool might come from Node itself, such as an HTTP server, path helper, cryptographic hash function, filesystem API, or child-process runner. In a local backend, these imports are the toolboxes that make the file more powerful than browser JavaScript."
        ),
    )
    add_code_block(
        doc,
        "import { usefulTool } from \"node:some-module\";\n\nconst result = usefulTool(\"input\");",
        language="javascript",
        title="Generic ES module import pattern",
    )
    add_body_paragraph(
        doc,
        paragraph(
            "The analogy is a lab bench. The bench itself is the file. The import statements are labeled instruments placed on the bench before work begins. Later functions do not need to manufacture a multimeter, oscilloscope, or power supply every time they need one; they reach for the prepared instrument."
        ),
    )

    doc.add_heading("2.2 Objects, Arrays, And Destructuring", level=2)
    add_body_paragraph(
        doc,
        paragraph(
            "JavaScript objects store named fields. They are common in this backend because API requests, compiler profiles, Git status reports, update reports, and AI context packets all carry several related values at once. Destructuring is simply a shortcut for pulling named fields out of an object so the function can work with them directly."
        ),
    )
    add_code_block(
        doc,
        "const profile = { language: \"C\", compiler: \"gcc\", timeoutMs: 30000 };\nconst { language, compiler } = profile;\n\nconst files = [\"main.c\", \"filter.h\", \"filter.c\"];",
        language="javascript",
        title="Generic object, destructuring, and array pattern",
    )
    add_body_paragraph(
        doc,
        paragraph(
            "A good circuit analogy is a connector with labeled pins. The object is the connector body. Each property name is a pin label. Destructuring is grabbing only the pins you need for the next operation."
        ),
    )

    doc.add_heading("2.3 Async Work And Await", level=2)
    add_body_paragraph(
        doc,
        paragraph(
            "Node backends constantly wait for slow work: reading a file, writing a file, calling GitHub, running Git, launching a compiler, or waiting for a model response. JavaScript represents those delayed results with Promises. The async keyword lets a function use await inside its body. await pauses that function until the operation finishes, while the rest of Node can continue handling other work."
        ),
    )
    add_code_block(
        doc,
        "async function saveText(filePath, text) {\n  await writeFile(filePath, text, \"utf8\");\n  return { ok: true, filePath };\n}",
        language="javascript",
        title="Generic asynchronous file operation",
    )
    add_body_paragraph(
        doc,
        paragraph(
            "The lab analogy is waiting for a soldering iron to heat. The instruction is still part of the procedure, but the technician does not pretend the result exists before the tool is ready."
        ),
    )

    doc.add_heading("2.4 HTTP Requests And Responses", level=2)
    add_body_paragraph(
        doc,
        paragraph(
            "A local backend receives HTTP requests. Each request has a method, such as GET or POST, and a path, such as /api/catalog. A response has a status code, headers, and a body. Headers describe how the browser should interpret the body. The body carries the useful information, often JSON."
        ),
    )
    add_code_block(
        doc,
        "if (request.method === \"GET\" && url.pathname === \"/api/example\") {\n  response.writeHead(200, { \"Content-Type\": \"application/json\" });\n  response.end(JSON.stringify({ ok: true }));\n}",
        language="javascript",
        title="Generic request and JSON response pattern",
    )
    add_body_paragraph(
        doc,
        paragraph(
            "Read routing code like a switchboard. The method and path select the circuit. The selected branch decides which subsystem receives the request and what response returns to the browser."
        ),
    )

    doc.add_heading("2.5 Paths, Processes, And Guardrails", level=2)
    add_body_paragraph(
        doc,
        paragraph(
            "A backend that writes files must treat paths carefully. User input should never be allowed to decide an unrestricted disk location. The normal pattern is to start from a trusted root folder, combine safe path segments, normalize the result, and reject anything that escapes the root."
        ),
    )
    add_code_block(
        doc,
        "const target = path.resolve(rootFolder, safeSegment);\nif (!target.startsWith(rootFolder)) {\n  throw new Error(\"Path escaped the workspace.\");\n}",
        language="javascript",
        title="Generic path-boundary pattern",
    )
    add_body_paragraph(
        doc,
        paragraph(
            "Process execution has a similar rule. A compiler or Git command should run with a specific executable, exact arguments, a working folder, a timeout, and captured output. That keeps the backend from becoming an open-ended command prompt."
        ),
    )
    doc.add_page_break()


def add_foundation_chapter(doc: Document, guide: dict) -> None:
    file_name = guide["file"]
    source = (REPO / file_name).read_text(encoding="utf-8")
    doc.add_heading("Chapter 1: Syntax, Runtime, And Framework Foundations", level=1)

    if file_name == "server.mjs":
        add_body_paragraph(
            doc,
            paragraph(
                "server.mjs is an ES module running on Node.js. That matters before any individual function is discussed, because the file is not browser JavaScript. It can read and write files, start child processes, run Git, call compilers, listen on a local port, and keep private API keys on the machine. The browser side asks for work; this file performs the work."
            ),
            lead="Runtime context: ",
        )
        add_body_paragraph(
            doc,
            paragraph(
                "The file intentionally does not use Express or another web framework. Instead, Node's createServer receives each request and the file routes it manually. That is why request parsing, security headers, JSON responses, local-write checks, and static-file serving all appear in this file."
            ),
            lead="Framework choice: ",
        )
        add_analogy_paragraph(
            doc,
            "This is like wiring a circuit by hand instead of plugging into a premade development board: it takes more explicit routing, but every connection is visible and owned by the file.",
        )
        import_block = "\n".join(line for line in source.splitlines() if line.startswith("import "))
        if import_block:
            add_code_block(doc, import_block, language="javascript", title="Node module imports used by server.mjs")
        add_body_paragraph(
            doc,
            paragraph(
                "The imports are the backend toolboxes. node:http creates the local HTTP server. node:fs/promises gives Promise-based file access. node:child_process runs Git, compilers, simulators, and installers. node:path and node:url make paths predictable on Windows. node:crypto supports hashing for authorization scopes. node:os lets the backend reason about the current machine. Once you understand these modules, the file reads less like random utility code and more like a small operating station for the builder."
            ),
            lead="Module map: ",
        )

    elif file_name == "script.js":
        add_body_paragraph(
            doc,
            paragraph(
                "script.js runs in the public website visitor's browser. It cannot run Git or touch the local filesystem. Its power comes from the DOM, browser history, fetch, forms, dialogs, and event listeners."
            ),
            lead="Runtime context: ",
        )
        add_analogy_paragraph(
            doc,
            "index.html creates the switches and displays, styles.css paints them, and script.js is the wiring harness that makes clicks, searches, project windows, and chat messages do something.",
        )
        add_body_paragraph(
            doc,
            paragraph(
                "The first part of the file captures DOM anchors with document.querySelector. Those anchors are not decorative variables; they are the live handles to the project grid, search box, filters, profile fields, contact area, resume panel, AI assistant, and footer. Later functions use those handles to populate the page after projects.json loads."
            ),
            lead="DOM anchors: ",
        )
        add_code_block(doc, "\n".join(source.splitlines()[:90]), language="javascript", title="Opening DOM anchors and public runtime state in script.js")

    else:
        add_body_paragraph(
            doc,
            paragraph(
                "index.html is the public shell. It is not a passive text file; it is the structural contract that script.js and styles.css depend on. Tags create landmarks, ids create runtime hooks, classes create styling hooks, and metadata tells browsers and search engines how to treat the site."
            ),
            lead="Markup contract: ",
        )
        add_analogy_paragraph(
            doc,
            "Read the HTML like a circuit board. Sections are major components, ids are test points, classes are labels and routes for styling, and script tags connect the runtime logic.",
        )
        add_body_paragraph(
            doc,
            paragraph(
                "If a hook such as #project-grid or #ai-assistant-input changes, the JavaScript connection that depends on it can break even if the page still looks valid."
            ),
            lead="Why hooks matter: ",
        )
        add_code_block(doc, "\n".join(source.splitlines()[:120]), language="html", title="Opening HTML structure, metadata, and runtime hooks")

    doc.add_heading("Function Syntax Rules Used Later", level=2)
    add_body_paragraph(
        doc,
        paragraph(
            "The later chapters include function entry points and exact source blocks, but they do not re-explain the same syntax over and over. Read async as a marker that the function may wait for file work, network work, Git, compilers, or AI responses. Read parameters as the caller's contract with the function. Read default values after = as fallback behavior. Read destructured parameters in braces as one object being unpacked into named fields. Read const, let, and var as places where the function or file stores information while it works."
        ),
        lead="Reading rule: ",
    )
    add_code_block(
        doc,
        "async function exampleOperation({ projectId, fileName = \"main.c\" }, options = {}) {\n  const target = resolveInsideCompileRoot(projectId, fileName);\n  return await runProcess(\"tool\", [target], options);\n}",
        language="javascript",
        title="Compact syntax example used as a reference for later chapters",
    )
    add_body_paragraph(
        doc,
        paragraph(
            "After this chapter, a function section will only explain what is specific to that function: why it exists, what data it receives, what exact code implements it, what helpers it calls, and what local objects or variables carry the work."
        ),
        lead="How to use later chapters: ",
    )
    doc.add_page_break()


SERVER_STATE_TEXT = {
    "root": paragraph("root is the trusted builder workspace. Almost every local file path eventually relates back to this folder, so it acts like the mechanical chassis of the backend. A request may name a file, but the backend decides whether that name can safely resolve under root."),
    "portfolioRoot": paragraph("portfolioRoot is the publishing mirror. It can be the same folder as root, but the newer architecture lets it be separate so public website files can be synchronized without mixing in private builder-only files."),
    "compileRoot": paragraph("compileRoot is the code-workspace area. Compile Code writes sources, cache folders, run folders, generated binaries, Java classes, HDL simulations, and VCD files here instead of scattering temporary build material through the portfolio."),
    "types": paragraph("types is the small static-file MIME table. It tells the browser whether a served file is JavaScript, CSS, JSON, an image, a PDF, or a generic download. Without it, local preview could serve files as the wrong content type."),
    "draftPath": paragraph("draftPath points at the private local draft catalog. Save draft writes here so the builder can preserve unfinished work without immediately changing the public projects.json."),
    "catalogPath": paragraph("catalogPath points at the publish-ready catalog. Apply to site writes here before the publishing path copies and commits website files."),
    "publishPaths": paragraph("publishPaths is the publishing allowlist. It is one of the most important safety objects in the file because it names exactly what can leave the builder workspace and enter the public website repository."),
    "publishAuthCachePath": paragraph("publishAuthCachePath stores the short-lived authorization record. The backend uses it to remember that the current machine, repository, branch, and workspace were recently verified instead of asking GitHub every time the app opens."),
    "publishAuthCacheTtlMs": paragraph("publishAuthCacheTtlMs expresses the normal one-day authorization window requested for publishing. It keeps the app convenient while still forcing periodic verification."),
    "publishAuthExtendedThreshold": paragraph("publishAuthExtendedThreshold is the rule that turns repeated successful authorizations into longer trust. In this app, repeated success over the recent window can stretch trust to thirty days."),
    "gitCandidates": paragraph("gitCandidates is the Git discovery list. The backend tries environment paths, normal Program Files locations, and per-user Git installs because Windows machines do not all put Git in the same place."),
    "compileLanguageProfiles": paragraph("compileLanguageProfiles is the compile workspace rulebook. Each language entry defines default filenames, file extensions, display labels, required tools, and optional Winget package ids. Compile, build, run, simulate, and install-tools operations all consult this object."),
    "compileToolCandidates": paragraph("compileToolCandidates is the tool-location map. It gives findTool realistic places to search for gcc, g++, javac, java, node, python, iverilog, vvp, and LTspice before declaring a compiler missing."),
    "portfolioAiInstructions": paragraph("portfolioAiInstructions is the policy prompt for Ask My Portfolio. It tells the model how to separate general questions from portfolio-specific questions, how to use public source excerpts, when to cite GitHub code, and what not to invent."),
}


SERVER_FUNCTION_NARRATIVES = {
    "securityHeaders": [
        paragraph("securityHeaders is a small function, but it sits at the outer skin of the backend. The function returns a headers object that every local response can reuse. It begins with fixed browser-safety headers, then merges the caller's extra headers at the end."),
        paragraph("Reading the body is straightforward: the returned object disables MIME sniffing, limits referrer leakage, blocks browser permissions that the builder does not need, and keeps the opener policy same-origin. The spread of extra is intentionally last so a caller can add Content-Type, Cache-Control, or another endpoint-specific header without rebuilding the safety baseline."),
        paragraph("This helper is not included because the syntax is difficult. It is included because repeated headers are easy to forget. Centralizing them makes local API responses and static-file responses behave consistently."),
    ],
    "sendJson": [
        paragraph("sendJson is the backend's normal API response writer. The caller gives it the Node response object, a status code, and a data object. The function writes headers first, then serializes the data as formatted JSON and ends the response."),
        paragraph("The important sequence is header policy before body output. It calls securityHeaders, adds application/json, disables caching, and only then sends JSON.stringify(data, null, 2). The indentation is not required by machines, but it makes API responses readable during debugging."),
        paragraph("When later endpoint branches call sendJson, they do not need to remember the JSON content type or cache rule. That keeps handleApi focused on routing decisions instead of response formatting."),
    ],
    "readRequestJson": [
        paragraph("readRequestJson turns a POST body stream into a parsed object. Node gives request data in chunks; the function collects those chunks, rejects bodies above a fixed size, and parses the final string as JSON."),
        paragraph("The chronological flow matters. It starts with an empty body, appends each incoming chunk, throws if the body grows past sixty megabytes, and waits for the stream to finish. Only after the stream is complete does it call JSON.parse. That order prevents partially received JSON from being treated as a real command."),
        paragraph("This function is the front door for almost every write action. Save draft, upload, compile, authenticate target, apply to site, and AI all depend on request data arriving as a safe object instead of an uncontrolled stream."),
    ],
    "resolveInsideRoot": [
        paragraph("resolveInsideRoot converts caller-supplied path segments into a path under the builder workspace. It resolves the final path, compares it against root, and rejects the request if the path escapes the workspace."),
        paragraph("The key detail is that path.normalize and path.join collapse the requested segments before comparison. That means dangerous segments such as '..' are interpreted before the safety check. The function then returns the normalized path only if it still belongs under root."),
        paragraph("This is the filesystem equivalent of a current limit on a bench supply. The UI can request a destination, but this function keeps the request inside the safe operating region."),
    ],
    "resolveInsidePortfolioRoot": [
        paragraph("resolveInsidePortfolioRoot performs the same boundary check for the publish mirror. It exists separately from resolveInsideRoot because the public portfolio workspace can live outside the builder workspace."),
        paragraph("The function first builds the candidate path from portfolioRoot and the supplied segments. If the resolved path does not start inside portfolioRoot, it throws. Publishing then receives a usable path or no path at all; there is no middle state where a bad path is silently accepted."),
    ],
    "resolveInsideCompileRoot": [
        paragraph("resolveInsideCompileRoot applies the boundary rule to code workspaces. Compile Code writes source files, run folders, binary artifacts, Java classes, and waveform files, so it needs the same path discipline as uploads and publishing."),
        paragraph("This guard is especially important because filenames and project ids can originate from user-edited project data. The function ensures a source filename becomes a file inside compileRoot, not an accidental write to another part of the computer."),
    ],
    "runProcess": [
        paragraph("runProcess is the command execution wrapper. The backend uses it for compilers, simulators, Git, Java, Python, Node, and installer-related commands. The function begins by choosing a timeout and working directory, then starts the child process with shell disabled and the environment explicitly composed."),
        paragraph("After spawn, the function becomes a recorder. It accumulates stdout and stderr as data arrives, writes stdin if the caller supplied input, and ends stdin so tools waiting for input do not hang forever. A timer watches the process. On timeout, Windows process trees are killed through taskkill before the child receives SIGKILL."),
        paragraph("The result object is the value of the wrapper. Instead of making every caller interpret Node child-process events, the function returns ok, code, stdout, stderr, timedOut, and elapsedMs. That shape is why Compile Code and Git publishing can show terminal output consistently."),
    ],
    "findTool": [
        paragraph("findTool locates a compiler, runtime, simulator, or installer helper. It starts with cached answers so repeated compile runs do not rescan the machine. If no cache entry exists, it reads the candidate list for the requested tool name."),
        paragraph("The function walks through realistic possibilities in order: explicit executable paths, commands available on PATH, and known installation folders. When it finds something usable, it stores that answer in compileToolCache and returns it. When it fails, it caches an empty string so the same missing tool does not cause repeated expensive searches."),
        paragraph("This is why the builder can run on different Windows machines without hardcoding one install layout. The compiler feature asks for gcc or iverilog; findTool translates that request into an actual executable path if one exists."),
    ],
    "normalizeCodeLanguage": [
        paragraph("normalizeCodeLanguage accepts loose human or UI language names and turns them into canonical backend ids. It trims the value, lowercases it, removes separator noise, and checks an alias map."),
        paragraph("The alias table is the center of the function. Values such as c++, cpp, sv, system verilog, py, js, node, spice, and cir collapse into the ids used by compileLanguageProfiles. If the cleaned value already names a known profile, the function keeps it. Otherwise JavaScript becomes the fallback."),
        paragraph("This function prevents every compile branch from having to understand every spelling a user might type. The rest of the backend can switch on c, cpp, verilog, systemverilog, java, javascript, python, html, or ltspice."),
    ],
    "compileWorkspaceFilesFromPayload": [
        paragraph("compileWorkspaceFilesFromPayload turns frontend editor state into a clean list of files for a project workspace. It starts with payload.workspaceFiles, then defines an inner addFile helper that normalizes each file before placing it in a Map."),
        paragraph("The helper ignores empty code, detects or normalizes the language, chooses a safe filename, creates a safe id, normalizes HDL roles, and stores the result by id. After processing the incoming workspace files, it adds the active file as well. That last step matters because the editor's active file might not already be present in the workspace list."),
        paragraph("The returned array is sorted by filename. Compile Code can then write and compile a stable set of files rather than trying to reason about raw frontend payloads."),
    ],
    "writeCompileWorkspaceSources": [
        paragraph("writeCompileWorkspaceSources is where normalized workspace objects become real files. It creates the target directory, tracks duplicate filenames, filters by language or extension when requested, writes each source file, and returns records that include uniqueName and sourcePath."),
        paragraph("Duplicate handling is the most practical part of the body. If two workspace entries claim the same filename, the second one receives a suffix before its extension. That avoids overwriting one source with another while still preserving a readable name in compiler output."),
        paragraph("Compilers do not compile rich editor blocks. They compile files. This function is the bridge between project editor state and the physical source files a toolchain can consume."),
    ],
    "parseVcdScopeText": [
        paragraph("parseVcdScopeText converts a Verilog VCD waveform file into data a scope viewer can draw. It reads the text line by line and keeps separate state for scopes, signal definitions, timescale text, current time, and the largest observed time."),
        paragraph("The first part of the loop handles metadata. Timescale lines are collected until their end marker. Scope lines push and pop hierarchical names. $var lines register signal codes, widths, and references, but the function caps the number of signals so a huge waveform does not overwhelm the UI."),
        paragraph("After $enddefinitions, the function switches from definitions to signal changes. Lines beginning with # update the current simulation time. Scalar and vector value-change lines are normalized, matched to their signal code, and appended only when the value actually changes. The returned object contains the waveform source name, timescale, max time, and a trimmed signal list."),
    ],
    "compileAndRunCode": [
        paragraph("compileAndRunCode is the backend IDE engine. It starts by deciding the language from the requested language or detected source. Then it chooses the language profile, produces a safe active filename, saves the active source, builds a normalized workspace list, and prepares local helpers for run directories, terminal text, and missing-tool checks."),
        paragraph("The body then becomes a language dispatcher. HTML receives a lightweight tag-balance validation. JavaScript and Python perform syntax checks first, then run the active file when the requested action is run. C and C++ select a compiler profile, write workspace sources, compute a cache key, compile with warnings and debug flags, and run the generated binary only when requested. Java compiles into a cache directory and runs the detected class name. Verilog and SystemVerilog require a testbench and waveform dump before simulation, compile sources through iverilog, run vvp, and parse VCD output for the scope. LTspice runs the simulator in batch mode and includes the log file if one appears."),
        paragraph("The large size of this function is intentional because it is the point where one frontend command becomes several very different toolchain behaviors. The function keeps the UI simple: the frontend sends one payload with action, code, language, files, project id, and stdin; the backend returns saved metadata, terminal text, ok status, and waveform data when HDL simulation succeeds."),
    ],
    "installCompilerTools": [
        paragraph("installCompilerTools is the automatic setup path for language toolchains that have Winget packages configured. It normalizes the selected language, reads the language profile, and stops early with a clear message when no automatic install rule exists."),
        paragraph("When a Winget list exists, the function runs each package install with source and package agreements accepted. It appends each command result to a terminal array and returns ok only if every package command succeeds. This gives the frontend a conventional installer-like log instead of a vague success or failure."),
    ],
    "validatePublishRemoteUrl": [
        paragraph("validatePublishRemoteUrl turns a typed repository target into either a clean remote URL or a clear error. It accepts blank input when the caller is not setting a remote, but when a value exists it must match a GitHub HTTPS URL, SSH URL, or a plausible .git remote."),
        paragraph("The function is deliberately early in the publish flow. A malformed repository URL should fail before Git remotes are changed, credentials are stored, or files are copied."),
    ],
    "getPublishTargetInfo": [
        paragraph("getPublishTargetInfo asks the local publish workspace what repository it is connected to. It checks whether the folder is a Git worktree, reads the origin remote and current branch when available, parses GitHub owner/repo names, and checks for a CNAME file."),
        paragraph("The result is a status object rather than a thrown error in normal missing-target cases. That lets the builder show a useful target panel even when no repository has been associated yet."),
    ],
    "configurePublishTarget": [
        paragraph("configurePublishTarget is the lower-level target setup helper. It reads repositoryUrl, customDomain, authUsername, and authPassword from the options object, validates each value, ensures the portfolio folder is a Git repository, and sets the origin remote when a remote was supplied."),
        paragraph("After the remote is set, the function detects the remote default branch and checks out that branch locally. Then it writes or removes CNAME based on the custom-domain input, stores credentials if provided, and returns the current target information. It configures the workspace, but authentication is handled by a stricter flow elsewhere."),
    ],
    "writePublishAuthCache": [
        paragraph("writePublishAuthCache records a successful publishing authorization for the current target. It reads the previous cache first so it can preserve successful history when the repository, branch, and remote match."),
        paragraph("The function converts recent successes into timestamps, counts successes in the last week, and decides whether extended trust applies. It then writes branch, remote, repository, machine scope, checked time, expiration time, history, trust days, and extended-trust status to the cache file. Finally it tries to chmod the file to owner-only permissions."),
        paragraph("This is the backend implementation of the convenience rule: normal success lasts one day, but repeated successful authorizations can be remembered longer on the same machine and workspace."),
    ],
    "assertPublishAccess": [
        paragraph("assertPublishAccess is the gate in front of Apply to site. It first proves the publish workspace is a Git worktree, then reads origin and branch. If either is missing, it throws a publishAccessError with details the UI can display."),
        paragraph("After the basic Git shape is known, the function optionally checks that the workspace looks like a compatible static portfolio. Then it builds an access object containing branch, remote display URL, repository identity, and authorization status. A fresh cache can satisfy the check immediately if it matches the same repository, branch, machine scope, and expiration window."),
        paragraph("When the cache is absent or stale, the function synchronizes the branch from remote, makes sure a local HEAD exists for write checks, verifies write access, writes a new authorization cache, and returns a richer access object. This is the reason Apply to site can be strict without forcing a GitHub login on every launch."),
    ],
    "syncFromPublishTarget": [
        paragraph("syncFromPublishTarget is the Load from target workflow. It starts by calling assertPublishAccess with compatibility disabled, because the point is to inspect the target repository and decide whether it contains compatible files."),
        paragraph("The function clones the target branch into a temporary folder, collects available import paths, and requires projects.json. Before replacing local files, it creates a timestamped backup under .omb-backups. Then it copies compatible assets, docs, Backgrounds, CNAME, robots, sitemap, and the catalog into the local builder workspace."),
        paragraph("The imported projects.json is also written to the local draft path, and syncPortfolioPublishFiles refreshes the publish mirror. The finally block removes the temporary clone, so the machine does not accumulate stale target copies after imports."),
    ],
    "authenticateGitHubForTarget": [
        paragraph("authenticateGitHubForTarget is the interactive trust-building path. It validates the repository URL and optional domain, validates credentials, requires a remote URL, ensures the publish folder is a Git repository, and captures the current target state so it can roll back on failure."),
        paragraph("Inside the try block, it sets the origin remote, checks Git/Git Credential Manager health, detects the target branch, and checks out that branch. If a matching authorization cache is already fresh and no new credentials were supplied, it can return cached success after writing the custom domain."),
        paragraph("If credentials are needed, the function either stores the supplied username/token pair or launches Git Credential Manager's GitHub login. It then forces assertPublishAccess, writes the custom domain, returns target/system details, and restores the previous target state if any error interrupts the sequence."),
    ],
    "syncPortfolioPublishFiles": [
        paragraph("syncPortfolioPublishFiles copies publishable website files from the builder workspace to the portfolio mirror. When root and portfolioRoot are the same folder, there is nothing to copy and the function reports that the workspace is not separated."),
        paragraph("In a separated setup, the function loops only through publishPaths. Existing source paths are copied after the destination is cleared. Missing paths can be removed from the mirror when removeMissing is true. That behavior keeps the public mirror aligned with the current builder output without staging arbitrary folders."),
    ],
    "publishSiteChanges": [
        paragraph("publishSiteChanges is the final Apply to site operation. It starts with a publishAccess object if one was already created, otherwise it calls assertPublishAccess itself. Then it synchronizes the remote branch, copies publishable files into the portfolio mirror, and bumps published asset versions."),
        paragraph("The function asks stageablePublishPaths which allowlisted paths actually exist, stages those paths with git add, checks status only for those paths, and commits only when there are changes. The commit message is date-based because the action is a generated portfolio update, not a hand-authored feature commit."),
        paragraph("Finally it pushes to the current branch and returns a public-facing publish report: workspace path, sync details, branch, whether a commit happened, commit output, push status, and push output."),
    ],
    "fetchGitHubRepositorySource": [
        paragraph("fetchGitHubRepositorySource is the public-code collector for Ask My Portfolio. It starts by reading repository metadata through GitHub's API, chooses a branch, and derives limits for how many files and how much text can enter the AI context."),
        paragraph("If the URL points directly to a file, the function fetches that raw file and returns a source excerpt immediately. Otherwise it reads the repository tree, filters to text-like files, scores files against the visitor's question, pulls README text, lists candidate files, and fetches selected source files. The final text is clamped so a repository cannot flood the prompt."),
        paragraph("This function is why the assistant can answer code questions from public GitHub repositories while still respecting boundaries. It fetches public evidence; it does not claim private access."),
    ],
    "enrichPortfolioContext": [
        paragraph("enrichPortfolioContext receives the context packet prepared by the public site and adds safe fetched excerpts. It copies up to ten sourceFetches, attaches the current question to each one, and runs fetchSourceText for each source."),
        paragraph("The returned context keeps everything from the original packet and adds sourceExcerpts plus a sourceFetchPolicy string. That policy is important because the model receives not only evidence but also the rule explaining what kind of evidence was fetched and what kind was only represented by metadata."),
    ],
    "callOllamaPortfolioAi": [
        paragraph("callOllamaPortfolioAi is the local-model fallback. It chooses the Ollama host and model from environment variables, creates an AbortController, and sets a timeout so a missing or slow local model does not freeze the API request."),
        paragraph("The function sends a chat request containing the portfolio AI instructions, question, intent, conversation, context, and web-search flag. It reads the response JSON, extracts text through extractOllamaText, and returns a normalized ok/error object. If the request fails or times out, the caller receives a clean failure instead of an unhandled network exception."),
    ],
    "handlePortfolioAi": [
        paragraph("handlePortfolioAi is the backend API endpoint for Ask My Portfolio. It reads the request JSON, clamps the question, normalizes conversation history, validates the intent, and rejects empty questions before doing any model work."),
        paragraph("The next step enriches portfolio context. After context is prepared, the function chooses a provider. If OPENAI_API_KEY is missing, it tries Ollama and returns a 503 only when the local model path also fails. If OpenAI is configured, it builds a Responses API payload with developer instructions, visitor question, intent, conversation JSON, portfolio context JSON, model settings, and optional web search."),
        paragraph("The function also handles fallback model selection for default OpenAI model errors. A successful response returns answer text, model, and web-search status. A failed response returns the provider error through sendJson so the frontend can show a clear AI-backend problem."),
    ],
    "getUpdateInfo": [
        paragraph("getUpdateInfo checks whether a newer builder release exists. It reads the current version from the app environment or package.json, calls the GitHub latest-release API, extracts the release tag, finds installer and portable assets, and compares versions."),
        paragraph("The function also checks blockedAppUpdateVersions. A release can exist but be deliberately skipped if a known updater issue would make automatic installation unsafe. The returned object gives the UI enough information to show up-to-date, update-available, blocked, or error states."),
    ],
    "getSecurityReport": [
        paragraph("getSecurityReport produces the builder's security visibility report. It runs release-download counts, authorization cache reads, and target info reads in parallel with Promise.allSettled so one failed source does not destroy the whole report."),
        paragraph("The returned report is intentionally honest. GitHub release assets can reveal download counts, and the local cache can reveal publishing-auth state, but a static GitHub Pages website cannot identify every visitor or provide raw IP logs by itself. That is why the report recommends Cloudflare analytics, Logpush, Analytics Engine, or a Worker-backed endpoint for IP-level logging."),
    ],
    "downloadAndLaunchAppUpdate": [
        paragraph("downloadAndLaunchAppUpdate performs the in-app updater handoff. It starts by calling getUpdateInfo and refuses to continue when the current version is already current, the release is blocked, no installer is available, or the OS is not Windows."),
        paragraph("After those checks, it downloads the installer into a temp update folder, rejects suspiciously tiny downloads, builds candidate relaunch paths, and writes both a PowerShell launcher script and a small CMD wrapper. The PowerShell script waits for the current process, stops lingering builder processes, runs the installer, retries with elevation if necessary, waits for the installed executable, and relaunches it."),
        paragraph("The backend launches the CMD wrapper detached, emits an update-started event, and exits the current process after a short delay. That sequence lets the app replace itself without asking the running executable to overwrite files that are still in use."),
    ],
    "handleApi": [
        paragraph("handleApi is the local API router. It receives the parsed URL and checks specific method/path combinations in order. GET routes handle catalog reads, template reads, publish-target status, system checks, app updates, security reports, and compiler-tool status."),
        paragraph("After the read-only routes, the function rejects non-POST requests that were not handled. Then it applies the most important boundary rule in the API: POST write actions must come from the local machine. Only after that local check does it route AI requests, code beautify/save/compile/install actions, Git installation, app update installation, GitHub authentication, load from target, save draft, apply catalog, and uploads."),
        paragraph("The function is large because it is replacing a framework router. The chronological ordering is the safety design: cheap read routes first, local-only write gate second, then each command branch wrapped in endpoint-specific try/catch logic so the frontend receives JSON rather than a crashed server."),
    ],
}


SERVER_TEXTBOOK_CHAPTERS = [
    {
        "title": "Chapter 3: Runtime Boundaries, Response Shape, And File Roots",
        "intro": [
            paragraph("The first backend concern is boundary control. Before the app can compile code or push a website, it must know which folders are trusted, how responses are shaped, and how requests become safe data."),
        ],
        "state": ["root", "portfolioRoot", "compileRoot", "types", "draftPath", "catalogPath", "publishPaths"],
        "functions": ["securityHeaders", "sendJson", "readRequestJson", "resolveInsideRoot", "resolveInsidePortfolioRoot", "resolveInsideCompileRoot"],
    },
    {
        "title": "Chapter 4: Compile Code Workspace And Tool Execution",
        "intro": [
            paragraph("This chapter follows the backend IDE path. The UI sends code and project workspace files; server.mjs turns that material into safe filenames, real source files, compiler invocations, terminal output, cached artifacts, and HDL scope data."),
        ],
        "state": ["compileLanguageProfiles", "compileToolCandidates"],
        "functions": ["normalizeCodeLanguage", "findTool", "runProcess", "compileWorkspaceFilesFromPayload", "writeCompileWorkspaceSources", "parseVcdScopeText", "compileAndRunCode", "installCompilerTools"],
    },
    {
        "title": "Chapter 5: Publishing, Authentication, And Loading From Target",
        "intro": [
            paragraph("Publishing is where local work becomes public website content. The functions in this chapter guard repository setup, daily authorization caching, target import, publish-mirror synchronization, and final Git push."),
        ],
        "state": ["publishAuthCachePath", "publishAuthCacheTtlMs", "publishAuthExtendedThreshold", "gitCandidates"],
        "functions": ["validatePublishRemoteUrl", "getPublishTargetInfo", "configurePublishTarget", "writePublishAuthCache", "assertPublishAccess", "syncFromPublishTarget", "authenticateGitHubForTarget", "syncPortfolioPublishFiles", "publishSiteChanges"],
    },
    {
        "title": "Chapter 6: AI Context, Public Source Fetching, And Model Routing",
        "intro": [
            paragraph("Ask My Portfolio is stronger when it can separate a greeting from a portfolio question, fetch safe public source evidence, and choose an AI backend without exposing private keys to the browser."),
        ],
        "state": ["portfolioAiInstructions"],
        "functions": ["fetchGitHubRepositorySource", "enrichPortfolioContext", "callOllamaPortfolioAi", "handlePortfolioAi"],
    },
    {
        "title": "Chapter 7: Updates, Security Reporting, And The API Switchboard",
        "intro": [
            paragraph("The final server responsibilities are maintenance and routing. The app needs update checks, security visibility, and one central API switchboard that maps browser actions to backend functions."),
        ],
        "state": [],
        "functions": ["getUpdateInfo", "getSecurityReport", "downloadAndLaunchAppUpdate", "handleApi"],
    },
]


def server_variable_lookup(source: str) -> dict[str, VariableDoc]:
    return {variable.name: variable for variable in top_level_variables(source)}


def server_function_lookup(source: str) -> dict[str, FunctionDoc]:
    return {function.name: function for function in function_ranges(source)}


def add_server_state_excerpt(doc: Document, variable: VariableDoc, text: str) -> None:
    doc.add_heading(variable.name, level=3)
    add_body_paragraph(doc, text)
    add_code_block(doc, variable.source, language="javascript", title=f"Server state excerpt: {variable.name}")


def add_server_function_walkthrough(doc: Document, function: FunctionDoc) -> None:
    doc.add_heading(function.name, level=3)
    paragraphs = SERVER_FUNCTION_NARRATIVES.get(function.name, [])
    if paragraphs:
        add_body_paragraph(doc, paragraphs[0])
    add_code_block(doc, function.source, language="javascript", title=f"Full source for {function.name}")
    for text in paragraphs[1:]:
        add_body_paragraph(doc, text)


def add_server_textbook_chapters(doc: Document) -> None:
    source = (REPO / "server.mjs").read_text(encoding="utf-8")
    variables = server_variable_lookup(source)
    functions = server_function_lookup(source)
    for chapter in SERVER_TEXTBOOK_CHAPTERS:
        doc.add_heading(chapter["title"], level=1)
        for text in chapter["intro"]:
            add_body_paragraph(doc, text)
        if chapter["state"]:
            doc.add_heading("Important working state", level=2)
            add_body_paragraph(
                doc,
                paragraph(
                    "These objects and constants are shown before the functions because they are the shared state the functions read from or protect. They are not syntax trivia; they explain what the backend remembers while it runs."
                ),
            )
            for variable_name in chapter["state"]:
                variable = variables.get(variable_name)
                if variable:
                    add_server_state_excerpt(doc, variable, SERVER_STATE_TEXT.get(variable_name, explain_variable(variable, "server.mjs", top_level=True, source=source)))
        doc.add_heading("Function walkthrough", level=2)
        for function_name in chapter["functions"]:
            function = functions.get(function_name)
            if function:
                add_server_function_walkthrough(doc, function)
        doc.add_page_break()


def function_fact_map(file_name: str, func: FunctionDoc, all_names: set[str]) -> dict[str, list[str]]:
    facts: dict[str, list[str]] = {}
    for heading, text in generated_function_paragraphs(file_name, func, all_names):
        facts.setdefault(heading, []).append(text)
    return facts


def add_syntax_block(doc: Document, signature: str) -> None:
    intro = doc.add_paragraph()
    intro.paragraph_format.first_line_indent = Inches(0)
    intro_run = intro.add_run("Function entry point:")
    intro_run.bold = True
    intro_run.font.color.rgb = RGBColor(15, 82, 110)
    block = doc.add_paragraph()
    block.paragraph_format.left_indent = Inches(0.24)
    block.paragraph_format.first_line_indent = Inches(0)
    run = block.add_run(f"{signature} {{ ... }}")
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(22, 78, 99)


def phrase_variant(seed: str, options: list[str]) -> str:
    return options[sum(ord(char) for char in seed) % len(options)]


def add_textbook_function(doc: Document, file_name: str, function: FunctionDoc, all_names: set[str]) -> None:
    facts = function_fact_map(file_name, function, all_names)
    purpose = " ".join(facts.get("Purpose and intuition", []))
    parameters = " ".join(facts.get("Parameters and data it receives", []))
    what = " ".join(facts.get("What it does", []))
    body_details = " ".join(facts.get("Important body details", []))
    collaborators = " ".join(facts.get("What it calls or works with", []))
    why = " ".join(facts.get("Why the file needs it", []))

    doc.add_heading(function.signature, level=3)
    if purpose:
        add_body_paragraph(
            doc,
            paragraph(
                phrase_variant(
                    function.name + "intro",
                    [
                        f"Begin with the role of {function.name}. ",
                        f"To understand {function.name}, start with the problem it owns. ",
                        f"Read {function.name} first as a named idea, not as syntax. ",
                        f"The best entry point for {function.name} is its responsibility in the surrounding workflow. ",
                    ],
                )
                + purpose
            ),
            lead="Role: ",
        )
    add_syntax_block(doc, function.signature)
    if parameters:
        add_body_paragraph(
            doc,
            paragraph(
                f"The function-specific inputs are worth reading before the body. {parameters}"
            ),
            lead="Inputs: ",
        )
    add_body_paragraph(
        doc,
        paragraph(
            "Now place that explanation beside the actual source. Read the code slowly: the colored keywords show control flow, the teal function name marks the entry point, the orange sections are string data, and the muted text marks comments. The goal is not to memorize the code, but to see how the written implementation matches the responsibility described above."
        ),
        lead="Reading the source: ",
        italic=True,
    )
    add_code_block(doc, function.source, language="javascript", title=f"Exact source block for {function.name}")
    implementation_text = paragraph(
        phrase_variant(
            function.name + "impl",
            [
                "With the signature understood, the implementation can be read as a sequence of decisions and side effects. ",
                "After the syntax, move into the body and follow the work in the order the function performs it. ",
                "The body is where the contract above becomes behavior. ",
                "Once the inputs are clear, the implementation shows how the function turns those inputs into useful program state. ",
            ],
        )
        + what
    )
    if body_details:
        implementation_text = paragraph(
            implementation_text
            + " A close reading of the body shows these implementation details: "
            + body_details
        )
    add_body_paragraph(doc, implementation_text, lead="Implementation: ")
    if collaborators:
        add_body_paragraph(
            doc,
            paragraph(
                phrase_variant(
                    function.name + "calls",
                    [
                        "Next, trace the helper calls that surround the function. ",
                        "The function also has to be read through the helpers it calls. ",
                        "Its connection to the rest of the file matters as much as its local code. ",
                        "The surrounding workflow becomes clearer when you notice its collaborators. ",
                    ],
                )
                + collaborators
            ),
            lead="Collaborators: ",
        )
    if why:
        add_body_paragraph(
            doc,
            paragraph(
                phrase_variant(
                    function.name + "why",
                    [
                        "This is why the function is worth separating instead of burying the logic somewhere else. ",
                        "The practical importance of the function is easiest to see by imagining it removed. ",
                        "The application keeps this behavior isolated because other features rely on it being consistent. ",
                        "This function earns its place in the file because it protects one behavior from being rewritten differently elsewhere. ",
                    ],
                )
                + why
            ),
            lead="Why it matters: ",
        )

    locals_found = local_variables(function)
    if not locals_found:
        add_body_paragraph(
            doc,
            paragraph(
                f"Inside {function.name}, there are no local const, let, or var values to discuss. The function mainly works from its parameters, module-level state, direct expressions, or helpers it calls."
            ),
            lead="Local values: ",
        )
        return

    add_body_paragraph(
        doc,
        paragraph(
            phrase_variant(
                function.name + "locals",
                [
                    f"The local objects and variables inside {function.name} show how the implementation holds temporary work while it runs. ",
                    f"The easiest way to follow the body of {function.name} is to watch its local values. ",
                    f"The local state inside {function.name} explains how the larger task is split into smaller steps. ",
                    f"After the main flow, the working values inside {function.name} reveal what the function must remember while it runs. ",
                ],
            )
            + "They are included here because they explain the implementation rather than merely naming syntax."
        ),
        lead="Local values: ",
    )
    for variable in locals_found:
        bullet = doc.add_paragraph(style=None)
        bullet.paragraph_format.left_indent = Inches(0.22)
        bullet.paragraph_format.first_line_indent = Inches(-0.12)
        bullet.add_run(f"{variable.name}: ").bold = True
        bullet.add_run(explain_variable(variable, file_name, top_level=False, function_body=function.body))


def add_function_walkthrough(doc: Document, guide: dict) -> None:
    file_name = guide["file"]
    if file_name in {"server.mjs", "script.js"}:
        doc.add_heading("Chapter 3: Function Walkthrough With Source Blocks", level=1)
        source = (REPO / file_name).read_text(encoding="utf-8")
        functions = function_ranges(source)
        all_names = {function.name for function in functions}
        doc.add_paragraph(
            "This walkthrough is function-by-function. Each section now follows a textbook rhythm: first the idea of the function, then the syntax, then how to read that syntax, then the implementation and the local objects that make the implementation work. Line numbers are deliberately avoided because the source will keep changing."
        )
        grouped: dict[str, list[FunctionDoc]] = {}
        for function in functions:
            category, _context = classify_function(file_name, function)
            grouped.setdefault(category, []).append(function)
        for category, items in grouped.items():
            doc.add_heading(category, level=2)
            doc.add_paragraph(
                f"The functions in this group all support {category.lower()}. Read them as a small subsystem: the smaller helpers prepare data and the larger functions coordinate side effects or rendering."
            )
            for function in items:
                add_textbook_function(doc, file_name, function, all_names)
        return

    doc.add_heading("Chapter 2: Markup Walkthrough", level=1)
    doc.add_paragraph(
        "This file is markup rather than a JavaScript function module, so the walkthrough is procedural. Each section explains a structural object in the document and why the public runtime depends on it."
    )
    for group in guide["groups"]:
        doc.add_heading(group["title"], level=2)
        doc.add_paragraph(group["intro"])
        for name, explanation in group["items"]:
            doc.add_heading(name, level=3)
            doc.add_paragraph(explanation)
    html_source = (REPO / file_name).read_text(encoding="utf-8")
    doc.add_heading("Chapter 3: Complete Markup Source", level=1)
    doc.add_paragraph(
        "The exact source below is included so the explanation can be read beside the real document shell. The colored tag names show HTML structure, purple attribute names show the hooks used by CSS and JavaScript, and orange strings show values such as IDs, classes, URLs, and labels."
    )
    add_code_block(doc, html_source, language="html", title="Exact source block for index.html")


def add_guide_content(doc: Document, guide: dict) -> None:
    if guide["file"] == "server.mjs":
        add_server_overview_chapter(doc, guide)
        add_server_syntax_chapter(doc)
        add_server_textbook_chapters(doc)
        return
    add_overview(doc, guide)
    add_foundation_chapter(doc, guide)
    add_variable_section(doc, guide)
    add_function_walkthrough(doc, guide)


def write_guide_docx(guide: dict, output_path: Path) -> None:
    doc = setup_document(guide["title"], guide["subtitle"])
    add_guide_content(doc, guide)
    doc.save(output_path)


def convert_docx_to_pdf(docx_path: Path, output_dir: Path) -> Path | None:
    soffice = Path(r"C:\Program Files\LibreOffice\program\soffice.com")
    if not soffice.exists():
        return None
    output_dir.mkdir(parents=True, exist_ok=True)
    profile = Path(tempfile.mkdtemp(prefix="omb-lo-profile-"))
    try:
        result = subprocess.run(
            [
                str(soffice),
                f"-env:UserInstallation=file:///{str(profile).replace(chr(92), '/')}",
                "--invisible",
                "--headless",
                "--norestore",
                "--convert-to",
                "pdf",
                "--outdir",
                str(output_dir),
                str(docx_path),
            ],
            check=False,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
        )
        pdf_path = output_dir / f"{docx_path.stem}.pdf"
        if pdf_path.exists():
            return pdf_path
        if result.returncode != 0:
            raise RuntimeError((result.stderr or result.stdout or "LibreOffice PDF conversion failed.").strip())
        return None
    finally:
        shutil.rmtree(profile, ignore_errors=True)


def clean_output() -> None:
    def remove_child(child: Path) -> None:
        def clear_readonly(function, item_path, excinfo):
            try:
                Path(item_path).chmod(0o700)
                function(item_path)
            except Exception:
                raise excinfo

        if child.is_dir():
            shutil.rmtree(child, onexc=clear_readonly)
        else:
            child.chmod(0o700)
            child.unlink(missing_ok=True)

    for folder in [DOCX_DIR, PDF_DIR]:
        folder.mkdir(parents=True, exist_ok=True)
        for child in list(folder.iterdir()):
            remove_child(child)


def main() -> None:
    clean_output()
    master = setup_document(
        "OMB Portfolio Builder Curated File Guides",
        "Book-style explanations for selected important files. This first curated pass covers server.mjs, script.js, and index.html.",
    )
    master.add_heading("How This Curated Set Is Different", level=1)
    master.add_paragraph(
        "The generated source references are useful as raw maps, but they are not how a beginner learns a system. This curated set starts each file with a long overview, then explains code by responsibility. It avoids line-number references because source lines will move as the application evolves."
    )
    master.add_paragraph(
        "The first covered files are server.mjs, script.js, and index.html because they form the local backend, public runtime, and public shell of the portfolio website."
    )
    master.add_page_break()

    for guide in GUIDES:
        docx_path = DOCX_DIR / f"{guide['file']}.docx"
        write_guide_docx(guide, docx_path)
        convert_docx_to_pdf(docx_path, PDF_DIR)

        master.add_heading(guide["title"], level=1)
        master.add_paragraph(guide["subtitle"])
        add_guide_content(master, guide)
        master.add_page_break()

    master.save(MASTER_DOCX)
    convert_docx_to_pdf(MASTER_DOCX, PDF_DIR)
    generated_master_pdf = PDF_DIR / MASTER_PDF.name
    if generated_master_pdf.exists():
        shutil.copyfile(generated_master_pdf, MASTER_PDF)
        generated_master_pdf.unlink()


if __name__ == "__main__":
    main()
